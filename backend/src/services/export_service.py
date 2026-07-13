"""导出服务"""
from __future__ import annotations

from dataclasses import dataclass

import json
import logging

import math

import os

import sqlite3

import tempfile

from pathlib import Path

from typing import Any, Dict, List, Optional, Tuple

import html

from src.perf import perf_span, record_counter
from src.schemas.form import normalize_annotation_key, parse_annotation_positions



logger = logging.getLogger(__name__)
_TOC_FALLBACK_PAGE_NUMBER = "1"


# Task 4.5: 导出错误异常类（确保返回稳定 JSON：detail + code）
class ExportError(Exception):
    """项目导出错误，携带 detail + code + status_code"""

    def __init__(self, message: str, code: str, status_code: int = 400):
        self.message = message
        self.code = code
        self.status_code = status_code
        super().__init__(message)


_EXPORT_ERROR_CODES = {
    "SCHEMA_INCOMPATIBLE": "EXPORT_SCHEMA_INCOMPATIBLE",
    "DATA_INCOMPATIBLE": "EXPORT_DATA_INCOMPATIBLE",
}


def _validate_form_field_schema(db_path: str) -> None:
    """验证 form_field 表结构兼容性。

    检查：
    1. form_field 表是否存在
    2. 是否有 legacy sort_order 列（说明迁移未完成）
    3. order_index 列是否存在
    4. field_definition_id 是否有 NULL 值（历史坏数据）

    若不兼容则抛出 ExportError。
    """
    conn = sqlite3.connect(db_path)
    try:
        # 检查 form_field 表是否存在
        cursor = conn.execute(
            "SELECT name FROM sqlite_master WHERE type='table' AND name='form_field'"
        )
        if not cursor.fetchone():
            # 表不存在，无需验证（可能是空项目）
            return

        # 获取列信息
        cursor = conn.execute("PRAGMA table_info(form_field)")
        columns = {row[1]: row for row in cursor.fetchall()}

        # 检查 legacy sort_order 列存在（说明迁移未完成）
        if "sort_order" in columns:
            raise ExportError(
                "数据库 form_field 表存在 legacy 'sort_order' 列，未完成迁移。请运行最新版本完成迁移后再导出。",
                _EXPORT_ERROR_CODES["SCHEMA_INCOMPATIBLE"],
            )

        # 检查 order_index 列不存在（不兼容）
        if "order_index" not in columns:
            raise ExportError(
                "数据库 form_field 表缺少 'order_index' 列，结构不兼容。",
                _EXPORT_ERROR_CODES["SCHEMA_INCOMPATIBLE"],
            )

        # 检查 field_definition_id 有 NULL 值（历史坏数据）
        # 排除 is_log_row=1 的记录（日志行允许 field_definition_id 为 NULL）
        cursor = conn.execute(
            "SELECT COUNT(*) FROM form_field WHERE field_definition_id IS NULL AND (is_log_row IS NULL OR is_log_row = 0)"
        )
        null_count = cursor.fetchone()[0]
        if null_count > 0:
            raise ExportError(
                f"数据库 form_field 表有 {null_count} 条记录的 field_definition_id 为 NULL，数据不兼容。",
                _EXPORT_ERROR_CODES["DATA_INCOMPATIBLE"],
            )

    finally:
        conn.close()



from docx import Document

from docx.shared import Pt, RGBColor, Inches, Cm

from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_LINE_SPACING,
    WD_TAB_ALIGNMENT,
    WD_TAB_LEADER,
)

from docx.enum.section import WD_SECTION, WD_ORIENT

from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE

from docx.oxml.ns import qn

from docx.oxml import OxmlElement

from docx.enum.style import WD_STYLE_TYPE

from sqlalchemy.orm import Session



from src.models import Project

from src.repositories.project_repository import ProjectRepository

from src.schemas.form import (
    ANNOTATION_FORM_KEY,
    ANNOTATION_POSITION_MAX_Y,
    ANNOTATION_POSITION_MIN_Y,
)
from src.schemas.project import normalize_screening_number_format

from src.services.field_rendering import (
    build_field_control_weight,
    build_inline_column_demands,
    build_inline_table_model,
    extract_default_lines,
    resolve_checkbox_label,
)

from src.services.width_planning import (

    compute_text_weight,

    compute_choice_atom_weight,

    plan_inline_table_width,

    plan_unified_table_width,

    plan_normal_table_width,

    compute_choice_trailing_fill_char_count,
    compute_horizontal_choice_trailing_fill_chars,
    compute_fill_line_char_count,

)


# 标签字号档位 -> Word 磅值；默认档位沿用 10.5pt
DEFAULT_LABEL_FONT_PT = 10.5

_LABEL_FONT_SIZE_PT = {"large": 12.0, "small": 9.0}
ACRF_ANNOTATION_FONT_SIZE_PT = 8.0
ACRF_ANNOTATION_HEIGHT_CM = 0.7
ACRF_ANNOTATION_PADDING_X_EMU = 22860
ACRF_ANNOTATION_PADDING_Y_EMU = 18000
ACRF_ANNOTATION_BORDER_WIDTH_EMU = 12700
ACRF_ANNOTATION_BOX_WIDTH_MAX_CM = 4.6
ACRF_ANNOTATION_DEFAULT_VERTICAL_OFFSET_EMU = -120000
ACRF_ANNOTATION_EMU_PER_01CM = int(Cm(0.01))


def resolve_label_font_pt(form_field) -> float:
    """返回字段标签的 Word 磅值；未知或缺省档位回退到默认 10.5pt。"""
    key = getattr(form_field, "label_font_size", None)
    return _LABEL_FONT_SIZE_PT.get(key, DEFAULT_LABEL_FONT_PT)


def resolve_label_bold(form_field) -> bool:
    """返回字段标签是否加粗；NULL/缺省视为加粗以兼容旧数据。"""
    return getattr(form_field, "label_bold", None) != 0


@dataclass(frozen=True)

class LayoutDecision:

    """表单布局决策（内部数据结构，不持久化）。"""



    mode: str  # "legacy" | "mixed_landscape" | "unified_landscape"

    column_count: int  # N 列数（仅 unified 有意义）

    label_span: int  # label 区合并列数

    value_span: int  # value 区合并列数

    force_landscape: bool = False  # paper_orientation='landscape' 强制覆写：legacy 模式下切横向

    force_portrait: bool = False   # paper_orientation='portrait' 强制覆写：legacy 模式下抑制 inline 宽表自动切横向





@dataclass(frozen=True)

class Segment:

    """统一横向布局的字段片段（内部数据结构，不持久化）。"""



    type: str  # "regular_field" | "full_row" | "inline_block"

    fields: list





class ExportService:

    """导出服务类"""



    # 字体常量

    FONT_EAST_ASIA = "SimSun"  # 宋体

    FONT_ASCII = "Times New Roman"

    PORTRAIT_CONTENT_WIDTH_CM = 14.66

    LANDSCAPE_CONTENT_WIDTH_CM = 23.36
    FORM_TABLE_ROW_HEIGHT_CM = 1
    SINGLE_LINE_HEIGHT_PT = 15.6
    CELL_VPAD_PT = (FORM_TABLE_ROW_HEIGHT_CM * 28.3465 - SINGLE_LINE_HEIGHT_PT) / 2

    # 纵向选项之间的段前间距（pt）。跨栈契约：与前端 main.css
    # `.choice-group--vertical .choice-atom + .choice-atom { margin-top }` 同值，
    # 保证 Word 预览与导出文档的纵向选项间距一致。
    VERTICAL_OPTION_GAP_PT = 3



    def __init__(self, session: Session):

        self.session = session

        self.project_repo = ProjectRepository(session)

        self._toc_entries: list[tuple[str, int, str]] = []

        self._toc_field_paragraph: Any = None

        self._toc_bookmark_counter: int = 0

        # 标题文本 -> 该条目页码占位 w:t 元素列表，供服务器侧写死真实页码

        self._toc_pageref_values: Dict[str, list] = {}
        self._annotation_docpr_counter: int = 0
        self._column_width_overrides: Dict[Any, Any] = {}
        self._current_annotation_offsets: Dict[str, int] = {}



    def _apply_exact_line_spacing(self, paragraph: Any) -> None:

        """将单元格段落行距固定到 Word 网格单行高度。"""

        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY

        paragraph.paragraph_format.line_spacing = Pt(self.SINGLE_LINE_HEIGHT_PT)



    def _disable_snap_to_grid(self, paragraph: Any) -> None:
        """关闭该段落的行网格吸附（snapToGrid=0）。

        文档启用了 w:docGrid（type=lines，linePitch=312，即 15.6pt 行网格），
        而 Word 默认 snapToGrid=1 会把段落的 space_before 吸附到整行网格。
        纵向选项首项 space_before=0（正好落在网格线上）与其余项 space_before=3pt
        （被吸附到下一条网格线）由此在 Word 中呈现为“首项到第二项间距偏大”，
        而段落里存储的间距其实是一致的 3pt。

        显式写入 snapToGrid=0，让 EXACTLY 15.6pt 行距与精确段前间距被原样呈现，
        使同一单元格内每个纵向选项之间的间距保持一致。
        """
        pPr = paragraph._p.get_or_add_pPr()
        snap = pPr.find(qn("w:snapToGrid"))
        if snap is None:
            snap = OxmlElement("w:snapToGrid")
            # snapToGrid 在 CT_PPr 中须排在以下元素之前；按合法顺序插入，
            # 无论段落是否已有 spacing/ind/jc 等都能得到合法 XML。
            pPr.insert_element_before(
                snap,
                "w:spacing", "w:ind", "w:contextualSpacing", "w:mirrorIndents",
                "w:suppressOverlap", "w:jc", "w:textDirection", "w:textAlignment",
                "w:textboxTightWrap", "w:outlineLvl", "w:divId", "w:cnfStyle",
                "w:rPr", "w:sectPr", "w:pPrChange",
            )
        snap.set(qn("w:val"), "0")



    def _enable_update_fields_on_open(self, doc: Document) -> None:
        """启用 Word 打开文档时更新域，用于刷新预渲染目录的 PAGEREF 页码。

        目录条目已在导出阶段预渲染，updateFields 只承担"页码精确化"提示，
        不再负责生成整份目录。按 OOXML CT_Settings 顺序，w:updateFields 必须
        位于 compat/rsids/mathPr 等元素之前，故插入到首个此类锚点之前；
        无锚点时再追加到末尾。
        """
        settings = doc.settings.element
        update_fields = settings.find(qn("w:updateFields"))
        if update_fields is None:
            update_fields = OxmlElement("w:updateFields")
            anchor = None
            for tag in ("w:hdrShapeDefaults", "w:footnotePr", "w:endnotePr",
                        "w:compat", "w:rsids", "w:mathPr"):
                anchor = settings.find(qn(tag))
                if anchor is not None:
                    break
            if anchor is not None:
                anchor.addprevious(update_fields)
            else:
                settings.append(update_fields)
        update_fields.set(qn("w:val"), "true")



    def _apply_cell_paragraph_metrics(
        self,
        paragraph: Any,
        *,
        space_before: bool = True,
        space_after: bool = True,
    ) -> None:

        """应用单行 1cm 所需的单元格段落间距与固定行距。"""

        if space_before:

            paragraph.paragraph_format.space_before = Pt(self.CELL_VPAD_PT)

        if space_after:

            paragraph.paragraph_format.space_after = Pt(self.CELL_VPAD_PT)

        self._apply_exact_line_spacing(paragraph)



    def export_project_to_word(
        self,
        project_id: int,
        output_path: str,
        column_width_overrides: Optional[Dict] = None,
        bake_toc_page_numbers: bool = False,
        annotated: bool = False,
    ) -> bool:

        """导出项目到 Word 文档

        Args:
            project_id: 项目 ID
            output_path: 输出文件路径
            column_width_overrides: 列宽覆盖参数，支持两种格式：
                { "inline:fieldIds=1,2,3": [0.5, 0.3, 0.2] }
                { "form_id": { "normal": [0.3, 0.7], "inline": [...], "unified": [...] } }
                fraction 值为 0.0~1.0，表示该列占总宽度的比例
            bake_toc_page_numbers: 是否在导出后用 LibreOffice 渲染算出真实页码写回
                目录（仅服务器侧、LibreOffice 可用时生效；失败自动回退 Word 更新域）。
        """

        try:

            # 重置目录收集状态，避免实例复用时跨导出累积

            self._toc_entries = []

            self._toc_field_paragraph = None

            self._toc_bookmark_counter = 0

            self._toc_pageref_values = {}
            self._annotation_docpr_counter = 0

            # 一次性 eager load 完整关系树，消除导出链路上的 N+1 查询

            with perf_span("project_tree_load"):
                project = self.project_repo.get_with_full_tree(project_id)

            if not project:

                return False

            forms_count = len(getattr(project, "forms", []) or [])
            fields_count = sum(len(getattr(form, "form_fields", []) or []) for form in getattr(project, "forms", []) or [])
            record_counter("forms_count", forms_count)
            record_counter("fields_count", fields_count)

            # 存储列宽覆盖供后续使用
            self._column_width_overrides = column_width_overrides or {}

            # 创建 Word 文档
            with perf_span("docx_generate"):
                doc = Document()

                # 统一文档字体和样式

                self._apply_document_style(doc)

                self._enable_update_fields_on_open(doc)



            # 设置页面边距和页眉页脚距离

            section = doc.sections[0]

            section.top_margin = Cm(2.54)

            section.bottom_margin = Cm(2.54)

            section.left_margin = Cm(3.17)

            section.right_margin = Cm(3.17)

            section.header_distance = Cm(1.5)

            section.footer_distance = Cm(1.3)



            # 设置纸张大小为A4

            section.page_width = Cm(21)

            section.page_height = Cm(29.7)



            # 设置文档网格（每行37字符，每页44行）

            sectPr = section._sectPr

            docGrid = sectPr.find(qn('w:docGrid'))

            if docGrid is None:

                docGrid = OxmlElement('w:docGrid')

                sectPr.append(docGrid)

            docGrid.set(qn('w:type'), 'lines')  # 只指定行网格

            docGrid.set(qn('w:linePitch'), '312')  # 15.6磅 = 312 twips (1磅=20 twips)

            docGrid.set(qn('w:charSpace'), '220')  # 11磅 = 220 twips



            # 1. 添加封面页

            self._add_cover_page(doc, project)



            # 2. 设置页眉页脚

            self._setup_header_footer(doc, project)



            # 3. 添加目录（占位）

            self._add_toc_placeholder(doc)

            self._switch_section(doc, WD_ORIENT.LANDSCAPE, project)



            # 4. 添加访视流程图

            self._add_visit_flow_diagram(doc, project)

            self._switch_section(doc, WD_ORIENT.PORTRAIT, project)



            # 5. 添加表单内容

            self._add_forms_content(doc, project, annotated=annotated)

            # 6. 写入 TOC 预渲染条目（依赖 _add_toc_heading 收集结果）

            self._populate_toc(doc)



            # 保存文档

            with perf_span("file_response_prepare"):
                doc.save(output_path)

            # 服务器侧用 LibreOffice 渲染算出真实页码写回目录（失败自动回退）

            if bake_toc_page_numbers:
                self._bake_toc_page_numbers(doc, output_path)

            return True



        except Exception:

            logger.exception("导出失败 project_id=%s", project_id)

            return False



    @staticmethod

    def _validate_output(output_path: str) -> tuple[bool, str]:

        """验证导出文件是否为有效且结构完整的 docx。"""

        try:

            if os.path.getsize(output_path) <= 0:

                return False, "导出文件为 0 字节"

        except OSError as exc:

            return False, f"无法读取导出文件大小: {exc}"



        try:

            doc = Document(output_path)

        except Exception as exc:

            return False, f"导出文件不是有效的 docx: {exc}"



        if len(doc.tables) < 3:

            return False, "导出文档結構不完整：至少需要封面表、訪視圖表和至少一個表單內容表格"



        return True, ""



    def _add_cover_para(self, doc: Document, text: str, size: Optional[float] = None, *, bold: bool = True, line_spacing: Optional[float] = None):

        """封面专用段落：居中，指定字号（None 表示继承样式）。返回段落对象供调用方设置额外格式。"""

        para = doc.add_paragraph()

        run = para.add_run(text)

        self._set_run_font(run, size=Pt(size) if size is not None else None, bold=bold)

        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if line_spacing is not None:

            para.paragraph_format.line_spacing = line_spacing

        return para



    def _add_cover_page(self, doc: Document, project: Project):

        """添加封面页。"""

        trial_name = project.trial_name or "[请设置试验名称]"

        ver = project.crf_version or "[版本号]"

        date_str = (

            project.crf_version_date.strftime("%Y-%m-%d")

            if project.crf_version_date and hasattr(project.crf_version_date, "strftime")

            else "[日期]"

        )

        sponsor = (project.sponsor or "").strip()

        dmu = (project.data_management_unit or "").strip()



        self._add_cover_para(doc, trial_name, 18)

        self._add_cover_para(doc, "", 15, bold=False, line_spacing=1.5)

        self._add_cover_para(doc, "", 15, bold=False, line_spacing=1.5)

        self._add_cover_para(doc, "Draft CRF（建库用）", 15)

        self._add_cover_para(doc, f"版本号及日期：{ver}/{date_str}")

        self._add_cover_para(doc, "", 15, bold=False, line_spacing=1.5)



        table = doc.add_table(rows=3, cols=2)

        table.autofit = False

        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 显式固定表格总宽度，防止自动撑满页面

        tbl_pr = table._tbl.tblPr

        tbl_w = tbl_pr.find(qn('w:tblW'))

        if tbl_w is None:

            tbl_w = OxmlElement('w:tblW')

            tbl_pr.append(tbl_w)

        tbl_w.set(qn('w:w'), '3855')   # 3.2cm + 3.6cm ≈ 3855 DXA

        tbl_w.set(qn('w:type'), 'dxa')

        table.columns[0].width = Cm(3.2)

        table.columns[1].width = Cm(3.6)

        cover_rows = [

            ("方案编号", project.protocol_number or ""),

            ("中心编号", "|__|__|"),

            ("筛选号", normalize_screening_number_format(project.screening_number_format) or "S|__|__||__|__|__|"),

        ]



        for row_idx, (label, value) in enumerate(cover_rows):

            row = table.rows[row_idx]

            self._apply_exact_row_height(row)

            left_cell = table.cell(row_idx, 0)

            right_cell = table.cell(row_idx, 1)

            left_para = left_cell.paragraphs[0]

            right_para = right_cell.paragraphs[0]

            left_run = left_para.add_run(label)

            right_run = right_para.add_run(value)

            self._set_run_font(left_run, size=Pt(10), bold=True)

            self._set_run_font(right_run, size=Pt(10), bold=True)

            for cell in (left_cell, right_cell):

                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                self._remove_cell_borders(cell)

                for cp in cell.paragraphs:

                    self._apply_cell_paragraph_metrics(cp)



        self._add_cover_para(doc, "", 15, bold=False, line_spacing=1.5)

        self._add_cover_para(doc, "", 15, bold=False, line_spacing=1.5)

        if sponsor:

            p = self._add_cover_para(doc, f"申办方：{sponsor}", 15, bold=True)

            p.paragraph_format.space_before = Pt(7.8)

            p.paragraph_format.space_after = Pt(7.8)

        if dmu:

            p = self._add_cover_para(doc, f"数据管理单位：{dmu}", 15, bold=True)

            p.paragraph_format.space_before = Pt(7.8)

            p.paragraph_format.space_after = Pt(7.8)

        doc.add_page_break()



    def _setup_header_footer(self, doc: Document, project: Project):

        """设置页眉页脚。"""

        for section in doc.sections:

            self._apply_header_to_section(section, project)

            self._apply_footer_to_section(section)



    def _apply_header_to_section(self, section, project: Project):

        """为指定 section 设置页眉。"""

        header = section.header

        header.is_linked_to_previous = False

        first_paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()

        first_paragraph.clear()

        p = first_paragraph

        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        p.paragraph_format.space_after = Pt(8)

        p.paragraph_format.line_spacing = 1.0



        if project.company_logo_path:

            from src.config import get_config

            logo_path = Path(get_config().upload_path) / "logos" / project.company_logo_path

            if logo_path.exists():

                try:

                    run_img = p.add_run()

                    picture = run_img.add_picture(str(logo_path), height=Inches(0.4))

                    self._make_picture_float(picture)

                except Exception:

                    pass



        version_parts = []

        if project.crf_version:

            version_parts.append(project.crf_version)

        if project.crf_version_date:

            date_str = (

                project.crf_version_date.strftime("%Y%m%d")

                if hasattr(project.crf_version_date, "strftime")

                else str(project.crf_version_date)

            )

            version_parts.append(date_str)

        if version_parts:

            run = p.add_run(f"版本号/日期：{'/'.join(version_parts)}")

            self._set_run_font(run, size=Pt(9))



    def _apply_footer_to_section(self, section):

        """为指定 section 设置页脚。"""

        footer = section.footer

        footer.is_linked_to_previous = False

        p_footer = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()

        p_footer.clear()

        p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER



        prefix = p_footer.add_run("第 ")

        self._set_run_font(prefix, size=Pt(9))



        run_page = p_footer.add_run()

        fld_char_begin = OxmlElement('w:fldChar')

        fld_char_begin.set(qn('w:fldCharType'), 'begin')

        run_page._r.append(fld_char_begin)

        instr_text = OxmlElement('w:instrText')

        instr_text.set(qn('xml:space'), 'preserve')

        instr_text.text = "PAGE"

        run_page._r.append(instr_text)

        fld_char_end = OxmlElement('w:fldChar')

        fld_char_end.set(qn('w:fldCharType'), 'end')

        run_page._r.append(fld_char_end)

        self._set_run_font(run_page, size=Pt(9))



        middle = p_footer.add_run(" 页 / 共 ")

        self._set_run_font(middle, size=Pt(9))



        run_total = p_footer.add_run()

        fld_char_begin_total = OxmlElement('w:fldChar')

        fld_char_begin_total.set(qn('w:fldCharType'), 'begin')

        run_total._r.append(fld_char_begin_total)

        instr_text_total = OxmlElement('w:instrText')

        instr_text_total.set(qn('xml:space'), 'preserve')

        instr_text_total.text = "NUMPAGES"

        run_total._r.append(instr_text_total)

        fld_char_end_total = OxmlElement('w:fldChar')

        fld_char_end_total.set(qn('w:fldCharType'), 'end')

        run_total._r.append(fld_char_end_total)

        self._set_run_font(run_total, size=Pt(9))



        suffix = p_footer.add_run(" 页")

        self._set_run_font(suffix, size=Pt(9))



    def _add_toc_placeholder(self, doc: Document):

        """添加"目录"标题段并记录锚点。

        TOC 域指令（begin/instrText/separate）与预渲染条目由 ``_populate_toc``
        在正文渲染完成后写入：域起始合入第一条条目、域 end 合入最后一条条目。
        标题段后不再插入空行或空的域壳段，故标题与首条目录之间无空行。
        """

        # 目录标题：宋体、小四(12pt)、加粗、居中

        p_title = doc.add_paragraph()

        run_title = p_title.add_run("目录")

        self._set_run_font(run_title, size=Pt(12), bold=True)

        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 记录锚点：预渲染条目紧接标题段之后插入（中间不留空行）

        self._toc_field_paragraph = p_title



    def _add_toc_heading(
        self,
        doc: Document,
        text: str,
        level: int = 1,
        *,
        form_domain: str | None = None,
        annotated: bool = False,
        annotation_delta_y_01cm: int = 0,
    ):

        """添加标题段落并注册 TOC 条目。

        为标题插入唯一名 ``_Toc%08d`` 书签，并记录标题文本、层级与书签名
        到 ``_toc_entries`` 供 ``_populate_toc`` 生成预渲染条目。
        """

        self._toc_bookmark_counter += 1

        bookmark_name = f"_Toc{self._toc_bookmark_counter:08d}"

        heading_para = doc.add_heading(text, level=level)

        # 在标题段落中插入 bookmarkStart / bookmarkEnd。
        # OOXML 要求 w:pPr 排在 EG_PContent 之前，因此 bookmarkStart
        # 须插在 pPr 之后（若有 pPr 则 index=1，否则 index=0）。

        p_element = heading_para._p

        bm_start = OxmlElement("w:bookmarkStart")

        bm_start.set(qn("w:id"), str(self._toc_bookmark_counter))

        bm_start.set(qn("w:name"), bookmark_name)

        bm_end = OxmlElement("w:bookmarkEnd")

        bm_end.set(qn("w:id"), str(self._toc_bookmark_counter))

        pPr = p_element.find(qn("w:pPr"))

        insert_idx = 1 if pPr is not None else 0

        p_element.insert(insert_idx, bm_start)

        # bookmarkEnd 紧跟在最后一个 run/rPr 等内容元素之后

        last_content = None

        for child in p_element:

            if child.tag not in (qn("w:bookmarkStart"), qn("w:pPr")):

                last_content = child

        if last_content is not None:

            last_content.addnext(bm_end)

        else:

            p_element.append(bm_end)

        for run in heading_para.runs:

            self._set_run_font(run)

        if annotated and form_domain:

            self._add_oid_annotation_box(
                heading_para,
                form_domain,
                delta_y_01cm=annotation_delta_y_01cm,
            )

        self._toc_entries.append((text, level, bookmark_name))

    def _next_annotation_docpr_id(self) -> int:

        self._annotation_docpr_counter += 1

        return self._annotation_docpr_counter

    def _normalize_annotation_text(self, text: str | None) -> str:

        return " ".join(str(text or "").split())

    def _estimate_annotation_width_cm(self, text: str) -> float:

        weighted_chars = sum(2 if ord(char) > 127 else 1 for char in text)

        return min(
            ACRF_ANNOTATION_BOX_WIDTH_MAX_CM,
            max(0.9, 0.45 + weighted_chars * 0.20),
        )

    def _clamp_annotation_delta_y(self, value: Any) -> int:

        if isinstance(value, bool) or not isinstance(value, int):

            return 0

        return max(ANNOTATION_POSITION_MIN_Y, min(ANNOTATION_POSITION_MAX_Y, value))

    def _annotation_delta_y_to_emu(self, value: int) -> int:

        return self._clamp_annotation_delta_y(value) * ACRF_ANNOTATION_EMU_PER_01CM

    def _load_annotation_offsets(self, raw_value: Any) -> Dict[str, int]:

        try:

            positions = parse_annotation_positions(raw_value)

        except ValueError as exc:

            raise ExportError(
                f"表单 annotation_positions 数据非法: {exc}",
                _EXPORT_ERROR_CODES["DATA_INCOMPATIBLE"],
            ) from exc

        if positions is None:

            return {}

        return {
            key: self._clamp_annotation_delta_y(position.y)
            for key, position in positions.items()
        }

    @staticmethod
    def _annotation_text_key(value: Any) -> str:

        return normalize_annotation_key(value)

    def _field_annotation_text(self, field_def: Any) -> str:
        if getattr(field_def, "field_type", None) == "标签":
            return ""

        return self._annotation_text_key(getattr(field_def, "variable_name", None))

    def _annotation_delta_y_for_key(self, key: str | None) -> int:

        normalized_key = self._annotation_text_key(key)
        if not normalized_key:

            return 0

        return self._current_annotation_offsets.get(normalized_key, 0)

    def _add_oid_annotation_box(
        self,
        anchor_paragraph: Any,
        text: str,
        *,
        delta_y_01cm: int = 0,
    ) -> None:

        from docx.oxml import parse_xml



        display_text = self._normalize_annotation_text(text)

        if not display_text:

            return



        box_width_emu = int(Cm(self._estimate_annotation_width_cm(display_text)))

        box_height_emu = int(Cm(ACRF_ANNOTATION_HEIGHT_CM))
        pos_offset_emu = (
            ACRF_ANNOTATION_DEFAULT_VERTICAL_OFFSET_EMU
            + self._annotation_delta_y_to_emu(delta_y_01cm)
        )

        docpr_id = self._next_annotation_docpr_id()

        escaped_text = html.escape(display_text, quote=True)

        anchor_xml = f"""
        <w:drawing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
        <wp:anchor xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
                   xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
                   xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                   xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
                   distT="0" distB="0" distL="0" distR="0" simplePos="0"
                   relativeHeight="251658240" behindDoc="0" locked="0"
                   layoutInCell="1" allowOverlap="1">
            <wp:simplePos x="0" y="0"/>
            <wp:positionH relativeFrom="column">
                <wp:align>right</wp:align>
            </wp:positionH>
            <wp:positionV relativeFrom="paragraph">
                <wp:posOffset>{pos_offset_emu}</wp:posOffset>
            </wp:positionV>
            <wp:extent cx="{box_width_emu}" cy="{box_height_emu}"/>
            <wp:effectExtent l="0" t="0" r="0" b="0"/>
            <wp:wrapNone/>
            <wp:docPr id="{docpr_id}" name="Annotation {docpr_id}" descr="{escaped_text}"/>
            <wp:cNvGraphicFramePr/>
            <a:graphic>
                <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
                    <wps:wsp>
                        <wps:cNvSpPr txBox="1"/>
                        <wps:spPr>
                            <a:xfrm>
                                <a:off x="0" y="0"/>
                                <a:ext cx="{box_width_emu}" cy="{box_height_emu}"/>
                            </a:xfrm>
                            <a:prstGeom prst="rect">
                                <a:avLst/>
                            </a:prstGeom>
                            <a:solidFill>
                                <a:srgbClr val="FFF2F2"/>
                            </a:solidFill>
                            <a:ln w="{ACRF_ANNOTATION_BORDER_WIDTH_EMU}">
                                <a:solidFill>
                                    <a:srgbClr val="C00000"/>
                                </a:solidFill>
                            </a:ln>
                        </wps:spPr>
                        <wps:txbx>
                            <w:txbxContent>
                                <w:p>
                                    <w:pPr>
                                        <w:spacing w:before="0" w:after="0" w:line="240" w:lineRule="auto"/>
                                        <w:jc w:val="center"/>
                                    </w:pPr>
                                    <w:r>
                                        <w:rPr>
                                            <w:rFonts w:ascii="{self.FONT_ASCII}" w:hAnsi="{self.FONT_ASCII}" w:eastAsia="{self.FONT_EAST_ASIA}"/>
                                            <w:color w:val="C00000"/>
                                            <w:sz w:val="{int(ACRF_ANNOTATION_FONT_SIZE_PT * 2)}"/>
                                            <w:szCs w:val="{int(ACRF_ANNOTATION_FONT_SIZE_PT * 2)}"/>
                                        </w:rPr>
                                        <w:t>{escaped_text}</w:t>
                                    </w:r>
                                </w:p>
                            </w:txbxContent>
                        </wps:txbx>
                        <wps:bodyPr wrap="none" lIns="{ACRF_ANNOTATION_PADDING_X_EMU}" tIns="{ACRF_ANNOTATION_PADDING_Y_EMU}" rIns="{ACRF_ANNOTATION_PADDING_X_EMU}" bIns="{ACRF_ANNOTATION_PADDING_Y_EMU}" anchor="ctr"/>
                    </wps:wsp>
                </a:graphicData>
            </a:graphic>
        </wp:anchor>
        </w:drawing>
        """

        anchor_run = anchor_paragraph.add_run()

        anchor_run._r.append(parse_xml(anchor_xml))



    def _apply_raw_run_font(self, rpr, size_half_pt: int = 21) -> None:

        """为裸 ``w:r`` 的 ``rPr`` 写入中英文字体与字号（单位：半点，21=10.5pt）。

        预渲染目录条目用裸 OxmlElement 构建，无法复用 ``_set_run_font``；显式写入
        rFonts/sz 保证目录与正文同为宋体，避免回退默认字体导致"排版不像目录"。
        """

        rFonts = OxmlElement("w:rFonts")

        rFonts.set(qn("w:ascii"), self.FONT_ASCII)

        rFonts.set(qn("w:hAnsi"), self.FONT_ASCII)

        rFonts.set(qn("w:eastAsia"), self.FONT_EAST_ASIA)

        rpr.append(rFonts)

        sz = OxmlElement("w:sz")

        sz.set(qn("w:val"), str(size_half_pt))

        rpr.append(sz)

        szCs = OxmlElement("w:szCs")

        szCs.set(qn("w:val"), str(size_half_pt))

        rpr.append(szCs)



    def _ensure_toc_styles(self, doc: Document) -> None:

        """确保 styles.xml 定义 TOC1/TOC2/TOC3 段落样式（幂等）。

        预渲染条目引用这些样式；缺失时 Word 回退默认字体，使目录与正文宋体
        不一致。注入后预渲染条目与 Word 更新域后重生成的目录均使用宋体。
        """

        styles_el = doc.styles.element

        existing = {

            s.get(qn("w:styleId"))

            for s in styles_el.findall(qn("w:style"))

        }

        for level in (1, 2, 3):

            style_id = f"TOC{level}"

            if style_id in existing:

                continue

            style = OxmlElement("w:style")

            style.set(qn("w:type"), "paragraph")

            style.set(qn("w:styleId"), style_id)

            name = OxmlElement("w:name")

            name.set(qn("w:val"), f"toc {level}")

            style.append(name)

            based_on = OxmlElement("w:basedOn")

            based_on.set(qn("w:val"), "Normal")

            style.append(based_on)

            next_style = OxmlElement("w:next")

            next_style.set(qn("w:val"), "Normal")

            style.append(next_style)

            ui_priority = OxmlElement("w:uiPriority")

            ui_priority.set(qn("w:val"), "39")

            style.append(ui_priority)

            pPr = OxmlElement("w:pPr")

            spacing = OxmlElement("w:spacing")

            spacing.set(qn("w:after"), "100")

            pPr.append(spacing)

            if level > 1:

                ind = OxmlElement("w:ind")

                ind.set(qn("w:left"), str((level - 1) * 420))

                pPr.append(ind)

            tabs = OxmlElement("w:tabs")

            tab = OxmlElement("w:tab")

            tab.set(qn("w:val"), "right")

            tab.set(qn("w:leader"), "dot")

            tab.set(qn("w:pos"), "8306")

            tabs.append(tab)

            pPr.append(tabs)

            style.append(pPr)

            rPr = OxmlElement("w:rPr")

            self._apply_raw_run_font(rPr)

            style.append(rPr)

            styles_el.append(style)



    def _build_toc_entry(self, heading_text: str, level: int, bookmark_name: str, with_field_start: bool):

        """构建一条预渲染目录条目段落（``w:p``）。

        ``with_field_start=True`` 时在条目最前合入外层 TOC 域起始
        （begin(dirty) → instrText → separate），使第一条条目成为域结果起点，
        标题段与目录之间不留空行；外层 ``end`` 由调用方合入最后一条条目。
        """

        entry = OxmlElement("w:p")

        pPr = OxmlElement("w:pPr")

        pStyle = OxmlElement("w:pStyle")

        pStyle.set(qn("w:val"), f"TOC{level}")

        pPr.append(pStyle)

        tabs = OxmlElement("w:tabs")

        tab_el = OxmlElement("w:tab")

        tab_el.set(qn("w:val"), "right")

        tab_el.set(qn("w:leader"), "dot")

        tab_el.set(qn("w:pos"), "8306")

        tabs.append(tab_el)

        pPr.append(tabs)

        indent = OxmlElement("w:ind")

        indent.set(qn("w:left"), str((level - 1) * 440))

        pPr.append(indent)

        entry.append(pPr)

        # 外层 TOC 域起始合入首条条目：begin(dirty) → instrText → separate

        if with_field_start:

            toc_begin = OxmlElement("w:r")

            toc_begin_fc = OxmlElement("w:fldChar")

            toc_begin_fc.set(qn("w:fldCharType"), "begin")

            toc_begin_fc.set(qn("w:dirty"), "1")

            toc_begin.append(toc_begin_fc)

            entry.append(toc_begin)

            toc_instr_run = OxmlElement("w:r")

            toc_instr = OxmlElement("w:instrText")

            toc_instr.set(qn("xml:space"), "preserve")

            toc_instr.text = 'TOC \\o "1-3" \\h \\z \\u'

            toc_instr_run.append(toc_instr)

            entry.append(toc_instr_run)

            toc_sep = OxmlElement("w:r")

            toc_sep_fc = OxmlElement("w:fldChar")

            toc_sep_fc.set(qn("w:fldCharType"), "separate")

            toc_sep.append(toc_sep_fc)

            entry.append(toc_sep)

        # 书签超链接（点击跳转），文本写入宋体保持与正文一致

        hl = OxmlElement("w:hyperlink")

        hl.set(qn("w:anchor"), bookmark_name)

        hl.set(qn("w:history"), "1")

        hl_run = OxmlElement("w:r")

        hl_rPr = OxmlElement("w:rPr")

        self._apply_raw_run_font(hl_rPr)

        hl_run.append(hl_rPr)

        hl_t = OxmlElement("w:t")

        hl_t.set(qn("xml:space"), "preserve")

        hl_t.text = heading_text

        hl_run.append(hl_t)

        hl.append(hl_run)

        entry.append(hl)

        # 右对齐点引导 tab 分隔符

        tab_run = OxmlElement("w:r")

        tab_run.append(OxmlElement("w:tab"))

        entry.append(tab_run)

        # PAGEREF 域：先写入非空回退页码，Word 更新域后可校正

        begin_run = OxmlElement("w:r")

        begin_fc = OxmlElement("w:fldChar")

        begin_fc.set(qn("w:fldCharType"), "begin")

        begin_fc.set(qn("w:dirty"), "1")

        begin_run.append(begin_fc)

        entry.append(begin_run)

        instr_run = OxmlElement("w:r")

        instr = OxmlElement("w:instrText")

        instr.set(qn("xml:space"), "preserve")

        instr.text = f"PAGEREF {bookmark_name} \\h"

        instr_run.append(instr)

        entry.append(instr_run)

        sep_run = OxmlElement("w:r")

        sep_fc = OxmlElement("w:fldChar")

        sep_fc.set(qn("w:fldCharType"), "separate")

        sep_run.append(sep_fc)

        entry.append(sep_run)

        val_run = OxmlElement("w:r")

        val_rPr = OxmlElement("w:rPr")

        self._apply_raw_run_font(val_rPr)

        val_run.append(val_rPr)

        val_t = OxmlElement("w:t")

        val_t.set(qn("xml:space"), "preserve")

        val_t.text = _TOC_FALLBACK_PAGE_NUMBER

        val_run.append(val_t)

        entry.append(val_run)

        # 记录页码占位元素，供服务器侧 LibreOffice 渲染后写死真实页码

        self._toc_pageref_values.setdefault(heading_text, []).append(val_t)

        end_run = OxmlElement("w:r")

        end_fc = OxmlElement("w:fldChar")

        end_fc.set(qn("w:fldCharType"), "end")

        end_run.append(end_fc)

        entry.append(end_run)

        return entry



    def _populate_toc(self, doc: Document):

        """在"目录"标题之后写入预渲染目录条目，并以外层 TOC 域收尾。

        需在所有 ``_add_toc_heading`` 调用完成后执行（即 ``_add_forms_content``
        之后）。第一条条目合入 TOC 域起始（begin/instrText/separate），最后一条
        条目合入外层 ``end``，使所有条目成为 ``separate`` 与 ``end`` 之间的域结果：
        零点击即可见条目、Word 更新域时整体替换不重复，且"目录"标题与条目间无空行。
        """

        anchor = self._toc_field_paragraph

        if anchor is None or not self._toc_entries:

            return

        # 确保 TOC1/2/3 样式存在，使预渲染条目与重生成目录字体一致（宋体）

        self._ensure_toc_styles(doc)

        # 条目紧接"目录"标题段插入；首条合入 TOC 域起始，逐条 addnext 顺序排列

        anchor_el = anchor._p

        last_entry_el = None

        for index, (heading_text, level, bookmark_name) in enumerate(self._toc_entries):

            entry_el = self._build_toc_entry(

                heading_text, level, bookmark_name, with_field_start=(index == 0)

            )

            anchor_el.addnext(entry_el)

            anchor_el = entry_el

            last_entry_el = entry_el

        # 外层 TOC 域 end 合入最后一条条目末尾，闭合 separate→end 域结果

        toc_end_run = OxmlElement("w:r")

        toc_end_fc = OxmlElement("w:fldChar")

        toc_end_fc.set(qn("w:fldCharType"), "end")

        toc_end_run.append(toc_end_fc)

        last_entry_el.append(toc_end_run)



    def _bake_toc_page_numbers(self, doc: Document, output_path: str) -> None:

        """服务器侧用 LibreOffice 渲染算出真实页码并写回目录 PAGEREF 占位。

        LibreOffice 可用且渲染成功时写入真实页码；否则保留非空回退页码。
        PAGEREF 域与 ``updateFields=true`` 始终保留，Word 仍可按自身分页再校正
        （LibreOffice 与 Word 分页可能差一页）。
        """

        if not self._toc_pageref_values:

            return

        from src.services import toc_pagination

        with perf_span("toc_page_bake"):

            pages = toc_pagination.compute_heading_pages(output_path)

        if not pages:

            logger.warning(
                "未取得真实目录页码，保留非空回退页码 output_path=%s",
                output_path,
            )
            return

        updated = False
        missing_headings: list[str] = []

        for heading_text, value_elements in self._toc_pageref_values.items():

            page = pages.get(heading_text)

            if page is None:

                missing_headings.append(heading_text)
                continue

            for val_t in value_elements:

                val_t.text = str(page)

                updated = True

        if missing_headings:

            logger.warning(
                "部分目录页码未取得，保留非空回退页码 missing_count=%s total=%s",
                len(missing_headings),
                len(self._toc_pageref_values),
            )

        if updated:

            doc.save(output_path)



    def _add_visit_flow_diagram(self, doc: Document, project: Project):

        """添加访视流程图"""

        self._add_toc_heading(doc, "表单访视分布图", level=1)



        all_forms = {}

        visit_form_map = {}

        for visit in project.visits:

            for visit_form in visit.visit_forms:

                if visit_form.form:

                    if visit_form.form.id not in all_forms:

                        all_forms[visit_form.form.id] = visit_form.form

                    visit_form_map[(visit.id, visit_form.form.id)] = True



        sorted_forms = sorted(

            all_forms.values(),

            key=lambda f: (f.order_index if f.order_index is not None else 999999, f.id),

        )

        visits = sorted(project.visits, key=lambda v: (v.sequence, v.id))



        row_count = len(sorted_forms) + 1 if sorted_forms else 1

        col_count = len(visits) + 1 if visits else 1

        table = doc.add_table(rows=row_count, cols=col_count)

        self._apply_grid_table_style(table)
        for row in table.rows:
            self._apply_exact_row_height(row)

        header_tr = table.rows[0]._tr
        tr_pr = header_tr.trPr
        if tr_pr is None:
            tr_pr = OxmlElement('w:trPr')
            header_tr.insert(0, tr_pr)
        tbl_header = tr_pr.find(qn('w:tblHeader'))
        if tbl_header is None:
            tbl_header = OxmlElement('w:tblHeader')
            tr_pr.append(tbl_header)
        tbl_header.set(qn('w:val'), 'true')



        header_cell_00 = table.rows[0].cells[0]

        header_para_00 = header_cell_00.paragraphs[0]

        header_run_00 = header_para_00.add_run("访视名称")

        self._set_run_font(header_run_00, size=Pt(10.5), bold=True)

        header_para_00.alignment = WD_ALIGN_PARAGRAPH.CENTER

        header_cell_00.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        self._apply_cell_shading(header_cell_00, 'A5C9EB')



        for col_idx, visit in enumerate(visits, start=1):

            header_cell = table.rows[0].cells[col_idx]

            header_para = header_cell.paragraphs[0]

            header_run = header_para.add_run(visit.name)

            self._set_run_font(header_run, size=Pt(10.5), bold=True)

            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            header_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            self._apply_cell_shading(header_cell, 'A5C9EB')



        for row_idx, form in enumerate(sorted_forms, start=1):

            name_cell = table.rows[row_idx].cells[0]

            name_para = name_cell.paragraphs[0]

            name_run = name_para.add_run(form.name)

            self._set_run_font(name_run, size=Pt(10.5), bold=True)

            name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

            name_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER



            for col_idx, visit in enumerate(visits, start=1):

                cross_cell = table.rows[row_idx].cells[col_idx]

                cross_para = cross_cell.paragraphs[0]

                if (visit.id, form.id) in visit_form_map:

                    cross_run = cross_para.add_run("×")

                    self._set_run_font(cross_run, size=Pt(10.5), bold=True)

                    cross_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cross_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER



        for row in table.rows:

            for cell in row.cells:

                for paragraph in cell.paragraphs:

                    paragraph.style = 'VisitFlow'



    def _add_forms_content(self, doc: Document, project: Project, *, annotated: bool = False):

        """添加表单内容（支持横向表格渲染与统一横向布局）。"""

        if not project.forms:

            self._build_form_table(doc, [], form_id=None, annotated=annotated)

            return



        sorted_forms = sorted(

            project.forms,

            key=lambda f: (f.order_index if f.order_index is not None else 999999, f.id),

        )

        total_forms = len(sorted_forms)

        sorted_visits = sorted(project.visits, key=lambda v: (v.sequence, v.id))
        form_to_visits: Dict[int, List] = {}
        for visit in sorted_visits:
            for visit_form in visit.visit_forms:
                form_to_visits.setdefault(visit_form.form_id, []).append(visit)



        for idx, form in enumerate(sorted_forms, start=1):
            if annotated:
                self._current_annotation_offsets = self._load_annotation_offsets(
                    getattr(form, "annotation_positions", None)
                )
            else:
                self._current_annotation_offsets = {}

            form_fields = sorted(form.form_fields, key=lambda ff: (ff.order_index, ff.id))

            paper_orientation = getattr(form, "paper_orientation", "auto") or "auto"

            layout = self._classify_form_layout(form_fields, paper_orientation=paper_orientation)

            is_last_form = idx == total_forms



            if layout.mode == "mixed_landscape":

                self._switch_section(doc, WD_ORIENT.LANDSCAPE, project)

                self._add_toc_heading(
                    doc,
                    f"{idx}. {form.name}",
                    level=1,
                    form_domain=form.domain,
                    annotated=annotated,
                    annotation_delta_y_01cm=self._annotation_delta_y_for_key(ANNOTATION_FORM_KEY),
                )

                groups = self._group_form_fields(form_fields)

                if groups == [[]]:

                    groups = []

                for group in groups:

                    if not group:

                        continue

                    first_field = group[0]

                    if first_field.inline_mark == 1:

                        self._add_inline_table(
                            doc,
                            group,
                            True,
                            form_id=form.id,
                            available_cm=self.LANDSCAPE_CONTENT_WIDTH_CM,
                            annotated=annotated,
                        )

                    else:

                        self._build_form_table(
                            doc,
                            group,
                            form_id=form.id,
                            available_cm=self.LANDSCAPE_CONTENT_WIDTH_CM,
                            annotated=annotated,
                        )

                if not groups:

                    self._build_form_table(
                        doc,
                        [],
                        form_id=form.id,
                        available_cm=self.LANDSCAPE_CONTENT_WIDTH_CM,
                        annotated=annotated,
                    )

                self._add_applicable_visits_paragraph(doc, form_to_visits.get(form.id, []))

                if not is_last_form:

                    self._switch_section(doc, WD_ORIENT.PORTRAIT, project)

            elif layout.mode == "unified_landscape":

                # 统一横向布局路径

                self._switch_section(doc, WD_ORIENT.LANDSCAPE, project)

                self._add_toc_heading(
                    doc,
                    f"{idx}. {form.name}",
                    level=1,
                    form_domain=form.domain,
                    annotated=annotated,
                    annotation_delta_y_01cm=self._annotation_delta_y_for_key(ANNOTATION_FORM_KEY),
                )

                segments = self._build_unified_segments(form_fields)

                self._build_unified_table(
                    doc,
                    segments,
                    layout,
                    form_id=form.id,
                    available_cm=self.LANDSCAPE_CONTENT_WIDTH_CM,
                    annotated=annotated,
                )

                self._add_applicable_visits_paragraph(doc, form_to_visits.get(form.id, []))

                # 仅当后续还有表单时才切回 portrait，避免末尾空白页

                if not is_last_form:

                    self._switch_section(doc, WD_ORIENT.PORTRAIT, project)

            else:

                # legacy 路径（保持现有行为）

                if layout.force_landscape:

                    self._switch_section(doc, WD_ORIENT.LANDSCAPE, project)

                self._add_toc_heading(
                    doc,
                    f"{idx}. {form.name}",
                    level=1,
                    form_domain=form.domain,
                    annotated=annotated,
                    annotation_delta_y_01cm=self._annotation_delta_y_for_key(ANNOTATION_FORM_KEY),
                )



                groups = self._group_form_fields(form_fields)

                if groups == [[]]:

                    groups = []



                for group in groups:

                    if not group:

                        continue



                    first_field = group[0]

                    if first_field.inline_mark == 1:

                        needs_temporary_landscape = len(group) > 4 and not layout.force_portrait

                        if needs_temporary_landscape and not layout.force_landscape:

                            self._switch_section(doc, WD_ORIENT.LANDSCAPE, project)



                        inline_available_cm = (
                            self.LANDSCAPE_CONTENT_WIDTH_CM
                            if layout.force_landscape or needs_temporary_landscape
                            else self.PORTRAIT_CONTENT_WIDTH_CM
                        )
                        self._add_inline_table(
                            doc,
                            group,
                            needs_temporary_landscape,
                            form_id=form.id,
                            available_cm=inline_available_cm,
                            annotated=annotated,
                        )



                        if needs_temporary_landscape and not layout.force_landscape:

                            self._switch_section(doc, WD_ORIENT.PORTRAIT, project)

                        continue



                    self._build_form_table(
                        doc,
                        group,
                        form_id=form.id,
                        available_cm=(
                            self.LANDSCAPE_CONTENT_WIDTH_CM
                            if layout.force_landscape
                            else self.PORTRAIT_CONTENT_WIDTH_CM
                        ),
                        annotated=annotated,
                    )



                if not groups:
                    self._build_form_table(
                        doc,
                        [],
                        form_id=form.id,
                        available_cm=(
                            self.LANDSCAPE_CONTENT_WIDTH_CM
                            if layout.force_landscape
                            else self.PORTRAIT_CONTENT_WIDTH_CM
                        ),
                        annotated=annotated,
                    )

                self._add_applicable_visits_paragraph(doc, form_to_visits.get(form.id, []))

                # 仅当后续还有表单时才分页/切回 portrait，避免末尾空白页

                if not is_last_form:

                    if layout.force_landscape:

                        self._switch_section(doc, WD_ORIENT.PORTRAIT, project)

                    else:

                        self._switch_section(doc, WD_ORIENT.PORTRAIT, project)

            self._current_annotation_offsets = {}



    def _add_applicable_visits_paragraph(self, doc: Document, visits):
        """在表单末尾追加"适用访视：<name>、..."段落。"""
        if not visits:
            return
        para = doc.add_paragraph(style='ApplicableVisits')
        prefix_run = para.add_run('适用访视：')
        self._set_run_font(prefix_run, size=Pt(10.5), bold=True)
        names_run = para.add_run('、'.join(visit.name for visit in visits))
        self._set_run_font(names_run, size=Pt(10.5))

    def _get_column_width_override(self, form_id, table_kind: str, col_count: int):
        """获取指定表单和表格类型的列宽覆盖配置（旧格式兼容）。

        Args:
            form_id: 表单 ID（None 时返回 None）
            table_kind: 表格类型 ("normal", "inline", "unified")
            col_count: 列数，用于验证覆盖配置长度

        Returns:
            List[float] 或 None：列宽 fraction 数组，长度应等于 col_count
        """
        if form_id is None:
            return None
        form_key = str(form_id)
        if form_key not in self._column_width_overrides:
            return None
        form_overrides = self._column_width_overrides[form_key]
        if table_kind not in form_overrides:
            return None
        overrides = form_overrides[table_kind]
        if not isinstance(overrides, list) or len(overrides) != col_count:
            return None
        # 校验每个元素是否为数值且在 0.0~1.0 范围内
        if not all(isinstance(v, (int, float)) and 0.0 <= v <= 1.0 for v in overrides):
            return None
        return overrides

    def _get_column_width_override_by_instance_id(
        self, table_instance_id: str, col_count: int, form_id=None, table_kind: str = None
    ):
        """根据 table_instance_id 获取列宽覆盖配置。

        Args:
            table_instance_id: 表实例标识符，格式 "kind:fieldIds=..." 或 legacy "groupIndex-kind-colCount"
            col_count: 列数，用于验证覆盖配置长度
            form_id: 表单 ID（用于 legacy 格式兼容）
            table_kind: 表格类型（用于 legacy 格式兼容）

        Returns:
            List[float] 或 None：列宽 fraction 数组，长度应等于 col_count
        """
        if not self._column_width_overrides:
            return None

        # 新格式：直接用 table_instance_id 查询
        if table_instance_id in self._column_width_overrides:
            overrides = self._column_width_overrides[table_instance_id]
            if isinstance(overrides, list) and len(overrides) == col_count:
                if all(isinstance(v, (int, float)) and 0.0 <= v <= 1.0 for v in overrides):
                    return overrides

        # Legacy 格式兼容：fallback 到旧方法
        if form_id is not None and table_kind is not None:
            return self._get_column_width_override(form_id, table_kind, col_count)

        return None

    def _build_table_instance_id(self, table_kind: str, fields) -> str:
        """根据表格类型和字段列表构建 table_instance_id。

        Args:
            table_kind: 表格类型 ("normal", "inline", "unified")
            fields: 字段列表，每个字段需要有 id 属性

        Returns:
            table_instance_id: 格式 "kind:fieldIds=<ordered-field-ids>"
        """
        field_ids = [str(f.id) for f in (fields or []) if f and hasattr(f, 'id') and f.id is not None]
        return f"{table_kind}:fieldIds={','.join(field_ids)}"

    def _classify_form_layout(self, form_fields, paper_orientation: str = "auto") -> LayoutDecision:
        """判断表单是否需要走统一横向布局（unified landscape）。



        触发条件：同时存在普通字段和 inline 字段，且最大 inline block 宽度 > 4。

        paper_orientation 覆写：
            - 'landscape'：自动判定为 legacy 时附加 force_landscape 标记，强制切横向。
            - 'portrait' ：强制 legacy 并附加 force_portrait，抑制 inline 宽表的自动横向切换。
            - 'auto'    ：维持原有自动判定。
        """

        if not form_fields:

            decision = LayoutDecision("legacy", 0, 0, 0)

        else:

            sorted_fields = sorted(form_fields, key=lambda f: (f.order_index, f.id))

            has_regular = any(f.inline_mark == 0 for f in sorted_fields)



            # 计算连续 inline block 的最大宽度

            max_block_width = 0

            current_block_width = 0

            for f in sorted_fields:

                if f.inline_mark == 1:

                    current_block_width += 1

                else:

                    max_block_width = max(max_block_width, current_block_width)

                    current_block_width = 0

            max_block_width = max(max_block_width, current_block_width)



            has_inline = max_block_width > 0

            if has_regular and has_inline and max_block_width > 4:

                N = max_block_width

                decision = LayoutDecision("mixed_landscape", N, 0, 0)

            else:

                decision = LayoutDecision("legacy", 0, 0, 0)



        if paper_orientation == "portrait":

            return LayoutDecision("legacy", 0, 0, 0, force_portrait=True)

        if paper_orientation == "landscape" and decision.mode == "legacy":

            return LayoutDecision("legacy", 0, 0, 0, force_landscape=True)

        return decision



    @staticmethod

    def _compute_merge_spans(N: int, M: int) -> List[int]:

        """将 N 列均分为 M 个 span，前面的 span 优先分配余数列。



        前置条件：1 <= M <= N。若 M 超出范围，返回 [1] * N 作为安全回退。

        """

        if M <= 0 or M > N:

            return [1] * N

        base = N // M

        extra = N % M

        spans = []

        for i in range(M):

            spans.append(base + (1 if i < extra else 0))

        return spans



    def _build_unified_segments(self, form_fields) -> List[Segment]:

        """按字段顺序构建 unified landscape 所需的渲染片段。"""

        if not form_fields:

            return []



        sorted_fields = sorted(form_fields, key=lambda f: (f.order_index, f.id))

        segments: List[Segment] = []

        inline_buffer = []



        for form_field in sorted_fields:

            if form_field.inline_mark == 1:

                inline_buffer.append(form_field)

                continue



            if inline_buffer:

                segments.append(Segment("inline_block", list(inline_buffer)))

                inline_buffer = []



            field_def = form_field.field_definition

            if form_field.is_log_row or (field_def and field_def.field_type in ("日志行", "标签")):

                segments.append(Segment("full_row", [form_field]))

            else:

                segments.append(Segment("regular_field", [form_field]))



        if inline_buffer:

            segments.append(Segment("inline_block", list(inline_buffer)))



        return segments



    def _switch_section(self, doc: Document, orientation, project: Project):

        """新建分节并切换页面方向，同时重设页眉页脚。"""

        new_section = doc.add_section(WD_SECTION.NEW_PAGE)

        new_section.orientation = orientation



        if orientation == WD_ORIENT.LANDSCAPE:

            new_section.page_width = Cm(29.7)

            new_section.page_height = Cm(21)

        else:

            new_section.page_width = Cm(21)

            new_section.page_height = Cm(29.7)



        self._apply_header_to_section(new_section, project)

        self._apply_footer_to_section(new_section)

        return new_section



    def _build_unified_table(
        self,
        doc: Document,
        segments,
        layout: LayoutDecision,
        form_id=None,
        *,
        available_cm: float = LANDSCAPE_CONTENT_WIDTH_CM,
        annotated: bool = False,
    ):
        """创建 unified landscape 表格并按片段顺序渲染。

        Args:
            form_id: 表单 ID，用于获取列宽覆盖配置
            available_cm: 当前分节可用宽度，需与表单纸张方向保持一致。
        """
        N = layout.column_count
        table = doc.add_table(rows=1, cols=N)
        table.autofit = False

        # 收集 inline block 的内容用于宽度规划（使用语义需求）
        segment_data = []
        all_block_demands = []
        regular_field_demands = []
        # 收集所有字段用于构建 table_instance_id
        all_fields = []
        for segment in segments:
            if segment.type == "inline_block" and segment.fields:
                headers, row_values, _ = build_inline_table_model(segment.fields)
                segment_data.append(("inline_block", headers, row_values))
                # 使用包含 choice/fill-line/unit 语义的需求
                all_block_demands.append(build_inline_column_demands(segment.fields))
                all_fields.extend(segment.fields)
            elif segment.type == "regular_field" and segment.fields:
                form_field = segment.fields[0]
                field_def = getattr(form_field, "field_definition", None)
                label = getattr(form_field, "label_override", None) or (
                    getattr(field_def, "label", None) if field_def else None
                ) or ""
                regular_field_demands.append({
                    "label_weight": compute_text_weight(label),
                    "control_weight": build_field_control_weight(form_field),
                })
                all_fields.extend(segment.fields)

        # 检查是否有列宽覆盖配置 - 使用 table_instance_id
        table_instance_id = self._build_table_instance_id("unified", all_fields)
        overrides = self._get_column_width_override_by_instance_id(
            table_instance_id, N, form_id=form_id, table_kind="unified"
        )
        if overrides:
            # 直接使用覆盖的 fraction 转换为 cm
            col_widths = [overrides[i] * available_cm for i in range(N)]
        else:
            # 使用内容驱动的宽度规划（传入物理列数 N 确保 per-slot-max 聚合）
            col_widths = plan_unified_table_width(
                segment_data,
                available_cm,
                column_count=N,
                block_demands=all_block_demands,
                regular_field_demands=regular_field_demands,
            ) if segment_data or regular_field_demands else None

        if col_widths and len(col_widths) == N:
            # 应用规划的列宽
            final_col_widths = [Cm(col_widths[col_idx]) for col_idx in range(N)]
        else:
            # 回退到等宽分配
            avail = Cm(available_cm)
            col_w = int(avail / N)
            final_col_widths = [col_w] * N

        for col_idx, col in enumerate(table.columns):
            col.width = final_col_widths[col_idx]

        table._tbl.remove(table.rows[0]._tr)



        for segment in segments:

            if segment.type == "regular_field" and segment.fields:

                self._add_unified_regular_row(table, segment.fields[0], layout, annotated=annotated)

            elif segment.type == "full_row" and segment.fields:

                self._add_unified_full_row(table, segment.fields[0], N, annotated=annotated)

            elif segment.type == "inline_block" and segment.fields:

                self._add_unified_inline_band(table, segment.fields, N, annotated=annotated)



        # 行添加完成后，同步 cell 的 tcW，避免 python-docx 默认 1234 twips 覆盖 gridCol
        # 让 Word 渲染时 col 与 cell 宽度对齐（与 _add_inline_table 同等契约）。
        for col_idx, col in enumerate(table.columns):
            if col_idx >= len(final_col_widths):
                break
            for cell in col.cells:
                cell.width = final_col_widths[col_idx]

        self._apply_grid_table_style(table)

        return table



    def _add_unified_regular_row(
        self,
        table,
        form_field,
        layout: LayoutDecision,
        *,
        annotated: bool = False,
    ):

        """在 unified table 中添加普通字段行。"""

        field_def = form_field.field_definition

        if not field_def:

            return



        N = layout.column_count

        row = table.add_row()
        self._apply_exact_row_height(row)

        left_cell = row.cells[0]

        if layout.label_span > 1:

            left_cell = left_cell.merge(row.cells[layout.label_span - 1])



        right_start = layout.label_span

        right_cell = row.cells[right_start]

        if layout.value_span > 1:

            right_cell = right_cell.merge(row.cells[N - 1])



        label = form_field.label_override or field_def.label or ""

        left_para = left_cell.paragraphs[0]

        left_run = left_para.add_run(label)

        self._set_run_font(
            left_run,
            size=Pt(resolve_label_font_pt(form_field)),
            bold=resolve_label_bold(form_field),
        )

        self._apply_cell_paragraph_metrics(left_para)

        left_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        left_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER



        right_para = right_cell.paragraphs[0]

        default_lines = extract_default_lines(form_field)

        is_vertical_choice = (

            not default_lines

            and field_def.field_type in ["单选（纵向）", "多选（纵向）"]

        )

        if default_lines:

            for line_idx, line in enumerate(default_lines):

                if line_idx > 0:

                    right_para.add_run().add_break()

                right_run = right_para.add_run(line)

                self._set_run_font(right_run, size=Pt(10.5))

        else:

            if field_def.field_type in ["单选（纵向）", "多选（纵向）"]:

                self._render_vertical_choices(right_cell, field_def)

            elif field_def.field_type in ["单选", "多选"]:

                self._render_choice_field(right_para, field_def)

            else:

                right_run = right_para.add_run(self._render_field_control(field_def))

                self._set_run_font(right_run, size=Pt(10.5))



        self._apply_cell_paragraph_metrics(
            right_para, space_before=not is_vertical_choice, space_after=False
        )

        right_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        self._apply_cell_paragraph_metrics(
            right_cell.paragraphs[-1], space_before=False, space_after=not is_vertical_choice
        )

        right_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        annotation_text = self._field_annotation_text(field_def)
        if annotated and annotation_text:

            self._add_oid_annotation_box(
                right_cell.paragraphs[-1],
                annotation_text,
                delta_y_01cm=self._annotation_delta_y_for_key(annotation_text),
            )



        if form_field.bg_color:

            self._apply_cell_shading(left_cell, form_field.bg_color)

            self._apply_cell_shading(right_cell, form_field.bg_color)

        if form_field.text_color:

            text_color = RGBColor.from_string(form_field.text_color)

            self._set_run_font(left_run, color=text_color)

            for paragraph in right_cell.paragraphs:

                for run in paragraph.runs:

                    self._set_run_font(run, color=text_color)



    def _add_unified_full_row(self, table, form_field, N: int, *, annotated: bool = False):

        """在 unified table 中添加全宽行。"""

        row = table.add_row()
        self._apply_exact_row_height(row)

        merged_cell = row.cells[0]

        if N > 1:

            merged_cell = merged_cell.merge(row.cells[N - 1])



        para = merged_cell.paragraphs[0]

        field_def = form_field.field_definition

        is_log_row = form_field.is_log_row or (field_def and field_def.field_type == "日志行")



        if is_log_row:

            label = form_field.label_override or "以下为log行"

            run = para.add_run(label)

            self._set_run_font(
                run,
                size=Pt(resolve_label_font_pt(form_field)),
                bold=resolve_label_bold(form_field),
            )

            self._apply_cell_paragraph_metrics(para)

            para.alignment = WD_ALIGN_PARAGRAPH.LEFT

            merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            self._apply_cell_shading(merged_cell, form_field.bg_color or 'D9D9D9')

            if form_field.text_color:

                self._set_run_font(run, color=RGBColor.from_string(form_field.text_color))

            annotation_text = self._field_annotation_text(field_def) if field_def else ""
            if annotated and annotation_text:

                self._add_oid_annotation_box(
                    para,
                    annotation_text,
                    delta_y_01cm=self._annotation_delta_y_for_key(annotation_text),
                )

            return



        para.style = 'FormLabel'

        label = form_field.label_override or (field_def.label if field_def else "")

        run = para.add_run(label)

        self._set_run_font(
            run,
            size=Pt(resolve_label_font_pt(form_field)),
            bold=resolve_label_bold(form_field),
        )

        self._apply_cell_paragraph_metrics(para)

        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        if form_field.bg_color:

            self._apply_cell_shading(merged_cell, form_field.bg_color)

        if form_field.text_color:

            self._set_run_font(run, color=RGBColor.from_string(form_field.text_color))

        annotation_text = self._field_annotation_text(field_def) if field_def else ""
        if annotated and annotation_text:

            self._add_oid_annotation_box(
                para,
                annotation_text,
                delta_y_01cm=self._annotation_delta_y_for_key(annotation_text),
            )



    def _add_unified_inline_band(self, table, block_fields, N: int, *, annotated: bool = False):

        """在 unified table 中添加 inline block 的表头和数据行。"""

        if not block_fields:

            return



        headers, row_values, field_defs = build_inline_table_model(block_fields)

        M = len(block_fields)

        spans = self._compute_merge_spans(N, M) if M < N else [1] * M



        header_row = table.add_row()
        self._apply_exact_row_height(header_row)

        start_col = 0

        for col_idx, label in enumerate(headers):

            span = spans[col_idx]
            field_def = field_defs[col_idx]

            cell = header_row.cells[start_col]

            if span > 1:

                cell = cell.merge(header_row.cells[start_col + span - 1])



            para = cell.paragraphs[0]

            run = para.add_run(label)

            self._set_run_font(
                run,
                size=Pt(resolve_label_font_pt(block_fields[col_idx])),
                bold=resolve_label_bold(block_fields[col_idx]),
            )

            self._apply_cell_paragraph_metrics(para)

            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            self._apply_cell_shading(cell, 'D9D9D9')

            annotation_text = self._field_annotation_text(field_def) if field_def else ""
            if annotated and annotation_text:

                self._add_oid_annotation_box(
                    para,
                    annotation_text,
                    delta_y_01cm=self._annotation_delta_y_for_key(annotation_text),
                )

            start_col += span



        for row_values_item in row_values:

            data_row = table.add_row()
            self._apply_exact_row_height(data_row)

            start_col = 0

            for col_idx, cell_value in enumerate(row_values_item):

                span = spans[col_idx]

                cell = data_row.cells[start_col]

                if span > 1:

                    cell = cell.merge(data_row.cells[start_col + span - 1])



                para = cell.paragraphs[0]

                field_def = field_defs[col_idx]



                is_vertical_choice = (

                    cell_value is None

                    and field_def

                    and field_def.field_type in ["单选（纵向）", "多选（纵向）"]

                )

                if cell_value is not None:

                    run = para.add_run(cell_value)

                    self._set_run_font(run, size=Pt(10.5))

                elif field_def:

                    if is_vertical_choice:

                        self._render_vertical_choices(cell, field_def)

                    elif field_def.field_type in ["单选", "多选"]:

                        self._render_choice_field(para, field_def)

                    else:

                        run = para.add_run(self._render_field_control(field_def))

                        self._set_run_font(run, size=Pt(10.5))



                self._apply_cell_paragraph_metrics(
                    para, space_before=not is_vertical_choice, space_after=not is_vertical_choice
                )

                para.alignment = WD_ALIGN_PARAGRAPH.LEFT

                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER



                marked_field = block_fields[col_idx]

                if marked_field.bg_color:

                    self._apply_cell_shading(cell, marked_field.bg_color)

                if marked_field.text_color:

                    text_color = RGBColor.from_string(marked_field.text_color)

                    for paragraph in cell.paragraphs:

                        for run in paragraph.runs:

                            self._set_run_font(run, color=text_color)



                start_col += span



    def _group_form_fields(self, form_fields):

        """按连续普通字段组与 inline 组拆分，保持 order_index 渲染顺序。"""

        if not form_fields:

            return [[]]



        groups = []

        current_group = []

        current_inline = None



        for form_field in form_fields:

            is_inline = form_field.inline_mark == 1

            if current_inline is None or is_inline == current_inline:

                current_group.append(form_field)

            else:

                groups.append(current_group)

                current_group = [form_field]

            current_inline = is_inline



        if current_group:

            groups.append(current_group)



        return groups or [[]]



    def _apply_exact_row_height(self, row, height_cm: Optional[float] = None):

        """为导出表格行设置 1cm 最小行高，多行内容可自然增高。"""

        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Cm(height_cm if height_cm is not None else self.FORM_TABLE_ROW_HEIGHT_CM)

    def _build_form_table(
        self,
        doc: Document,
        fields,
        form_id=None,
        *,
        available_cm: float = PORTRAIT_CONTENT_WIDTH_CM,
        annotated: bool = False,
    ):
        """将一组普通字段渲染为单张表格。

        Args:
            form_id: 表单 ID，用于获取列宽覆盖配置
            available_cm: 当前分节可用宽度，需与表单纸张方向保持一致。
        """
        row_count = len(fields) if fields else 1
        table = doc.add_table(rows=row_count, cols=2)
        table.autofit = False

        # 内容驱动列宽：与前端 planNormalColumnFractions / 横向/统一表格语义一致。
        normal_widths = plan_normal_table_width(fields or [], available_cm=available_cm)

        # 应用列宽覆盖（如果有）- 使用 table_instance_id
        table_instance_id = self._build_table_instance_id("normal", fields)
        overrides = self._get_column_width_override_by_instance_id(
            table_instance_id, 2, form_id=form_id, table_kind="normal"
        )
        if overrides:
            normal_widths = [overrides[i] * available_cm for i in range(2)]

        table.columns[0].width = Cm(normal_widths[0])
        table.columns[1].width = Cm(normal_widths[1])
        self._apply_grid_table_style(table)



        if not fields:

            return table



        for row_idx, form_field in enumerate(fields):

            field_def = form_field.field_definition

            if form_field.is_log_row or (field_def and field_def.field_type == "日志行"):

                self._add_log_row(table, row_idx, form_field, annotated=annotated)

            elif field_def and field_def.field_type == "标签":

                self._add_label_row(table, row_idx, form_field, annotated=annotated)

            else:

                self._add_field_row(table, row_idx, form_field, normal_widths, annotated=annotated)



        return table



    def _add_log_row(self, table, row_idx: int, form_field, *, annotated: bool = False):

        """添加日志行。"""

        row = table.rows[row_idx]
        self._apply_exact_row_height(row)

        merged_cell = row.cells[0].merge(row.cells[1])

        para = merged_cell.paragraphs[0]

        label = form_field.label_override or "以下为log行"

        run = para.add_run(label)

        self._set_run_font(
            run,
            size=Pt(resolve_label_font_pt(form_field)),
            bold=resolve_label_bold(form_field),
        )

        self._apply_cell_paragraph_metrics(para)

        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        self._apply_cell_shading(merged_cell, form_field.bg_color or 'D9D9D9')

        if form_field.text_color:

            self._set_run_font(run, color=RGBColor.from_string(form_field.text_color))

        field_def = form_field.field_definition

        annotation_text = self._field_annotation_text(field_def) if field_def else ""
        if annotated and annotation_text:

            self._add_oid_annotation_box(
                para,
                annotation_text,
                delta_y_01cm=self._annotation_delta_y_for_key(annotation_text),
            )



    def _add_label_row(self, table, row_idx: int, form_field, *, annotated: bool = False):

        """添加标签字段行。"""

        row = table.rows[row_idx]
        self._apply_exact_row_height(row)

        merged_cell = row.cells[0].merge(row.cells[1])

        para = merged_cell.paragraphs[0]

        para.style = 'FormLabel'

        field_def = form_field.field_definition

        label = form_field.label_override or (field_def.label if field_def else "")

        run = para.add_run(label)

        self._set_run_font(
            run,
            size=Pt(resolve_label_font_pt(form_field)),
            bold=resolve_label_bold(form_field),
        )

        self._apply_cell_paragraph_metrics(para)

        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        annotation_text = self._field_annotation_text(field_def) if field_def else ""
        if annotated and annotation_text:

            self._add_oid_annotation_box(
                para,
                annotation_text,
                delta_y_01cm=self._annotation_delta_y_for_key(annotation_text),
            )



    def _add_field_row(self, table, row_idx: int, form_field, widths, *, annotated: bool = False):

        """添加普通字段行。"""

        row = table.rows[row_idx]
        self._apply_exact_row_height(row)

        field_def = form_field.field_definition

        if not field_def:

            return



        left_cell = row.cells[0]

        right_cell = row.cells[1]

        left_cell.width = Cm(widths[0])

        right_cell.width = Cm(widths[1])



        label = form_field.label_override or field_def.label or ""

        left_para = left_cell.paragraphs[0]

        left_run = left_para.add_run(label)

        self._set_run_font(
            left_run,
            size=Pt(resolve_label_font_pt(form_field)),
            bold=resolve_label_bold(form_field),
        )

        self._apply_cell_paragraph_metrics(left_para)

        left_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        left_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER



        right_para = right_cell.paragraphs[0]

        default_lines = extract_default_lines(form_field)

        is_vertical_choice = (

            not default_lines

            and field_def.field_type in ["单选（纵向）", "多选（纵向）"]

        )

        if default_lines:

            for line_idx, line in enumerate(default_lines):

                if line_idx > 0:

                    right_para.add_run().add_break()

                right_run = right_para.add_run(line)

                self._set_run_font(right_run, size=Pt(10.5))

        else:

            if field_def.field_type in ["单选（纵向）", "多选（纵向）"]:

                self._render_vertical_choices(right_cell, field_def, column_cm=widths[1])

            elif field_def.field_type in ["单选", "多选"]:

                self._render_choice_field(right_para, field_def, column_cm=widths[1])

            else:

                # 填写线下划线根数按 control 列实际宽度自适应（不换行），
                # 与前端预览共享同一估算公式以保证逐字一致。
                fill_chars = compute_fill_line_char_count(widths[1])
                right_run = right_para.add_run(
                    self._render_field_control(field_def, fill_line_chars=fill_chars)
                )

                self._set_run_font(right_run, size=Pt(10.5))



        self._apply_cell_paragraph_metrics(
            right_para, space_before=not is_vertical_choice, space_after=False
        )

        right_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        self._apply_cell_paragraph_metrics(
            right_cell.paragraphs[-1], space_before=False, space_after=not is_vertical_choice
        )

        # 文本/标签填写线贴单元格底部，使下划线呈“下划线”而非垂直居中时看起来像横线；
        # 方框占位（日期/数值等）与选项保持垂直居中。
        is_plain_fill_line = not default_lines and field_def.field_type in ("文本", "标签")
        right_cell.vertical_alignment = (
            WD_ALIGN_VERTICAL.BOTTOM if is_plain_fill_line else WD_ALIGN_VERTICAL.CENTER
        )

        annotation_text = self._field_annotation_text(field_def)
        if annotated and annotation_text:

            self._add_oid_annotation_box(
                right_cell.paragraphs[-1],
                annotation_text,
                delta_y_01cm=self._annotation_delta_y_for_key(annotation_text),
            )



        if form_field.bg_color:

            self._apply_cell_shading(left_cell, form_field.bg_color)

            self._apply_cell_shading(right_cell, form_field.bg_color)

        if form_field.text_color:

            text_color = RGBColor.from_string(form_field.text_color)

            self._set_run_font(left_run, color=text_color)

            for paragraph in right_cell.paragraphs:

                for run in paragraph.runs:

                    self._set_run_font(run, color=text_color)



    def _add_inline_table(
        self,
        doc: Document,
        marked_fields,
        is_wide=False,
        form_id=None,
        *,
        available_cm: float | None = None,
        annotated: bool = False,
    ):
        """添加横向表格（1表头行+N内容行），使用内容驱动的宽度规划

        Args:
            is_wide: 兼容旧调用的横向宽表标记；显式传入 available_cm 时仅保留语义。
            form_id: 表单 ID，用于获取列宽覆盖配置
            available_cm: 当前分节可用宽度，需与表单纸张方向保持一致。
        """
        if not marked_fields:
            return

        # 使用共享模块构建表格数据模型
        headers, row_values, field_defs = build_inline_table_model(marked_fields)

        # 创建表格：1+max_rows行，N列
        table = doc.add_table(rows=1 + len(row_values), cols=len(marked_fields))
        table.autofit = False
        self._apply_grid_table_style(table)

        # 使用内容驱动的宽度规划替代等宽分配
        if available_cm is None:
            avail_cm = self.LANDSCAPE_CONTENT_WIDTH_CM if is_wide else self.PORTRAIT_CONTENT_WIDTH_CM
        else:
            avail_cm = available_cm

        # 检查是否有列宽覆盖配置 - 使用 table_instance_id
        table_instance_id = self._build_table_instance_id("inline", marked_fields)
        overrides = self._get_column_width_override_by_instance_id(
            table_instance_id, len(marked_fields), form_id=form_id, table_kind="inline"
        )
        if overrides:
            # 直接使用覆盖的 fraction 转换为 cm
            col_widths = [overrides[i] * avail_cm for i in range(len(marked_fields))]
        else:
            # 构建包含 choice/fill-line/unit 等语义的列需求
            semantic_demands = build_inline_column_demands(marked_fields)
            col_widths = plan_inline_table_width(headers, row_values, avail_cm, semantic_demands=semantic_demands)

        # 应用列宽
        for col_idx, col in enumerate(table.columns):
            if col_idx < len(col_widths):
                width = Cm(col_widths[col_idx])
            else:
                # 回退到等宽分配
                width = int(Cm(avail_cm) / len(marked_fields))
            col.width = width
            for cell in col.cells:
                cell.width = width



        # 第一行：表头（字段名称）

        self._apply_exact_row_height(table.rows[0])

        for col_idx, label in enumerate(headers):

            field_def = field_defs[col_idx]

            if not field_def:

                continue



            cell = table.rows[0].cells[col_idx]

            para = cell.paragraphs[0]

            run = para.add_run(label)

            self._set_run_font(
                run,
                size=Pt(resolve_label_font_pt(marked_fields[col_idx])),
                bold=resolve_label_bold(marked_fields[col_idx]),
            )



            # 段落格式：单行 1cm 所需上下间距，固定 15.6pt 行距

            self._apply_cell_paragraph_metrics(para)

            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER



            # 表头底纹：#D9D9D9

            shading_elm = OxmlElement('w:shd')

            shading_elm.set(qn('w:fill'), 'D9D9D9')

            cell._tc.get_or_add_tcPr().append(shading_elm)

            annotation_text = self._field_annotation_text(field_def)
            if annotated and annotation_text:

                self._add_oid_annotation_box(
                    para,
                    annotation_text,
                    delta_y_01cm=self._annotation_delta_y_for_key(annotation_text),
                )



        # 内容行：根据row_values生成

        for row_idx, row in enumerate(row_values):

            self._apply_exact_row_height(table.rows[row_idx + 1])

            for col_idx, cell_value in enumerate(row):

                field_def = field_defs[col_idx]

                if not field_def:

                    continue



                cell = table.rows[row_idx + 1].cells[col_idx]

                para = cell.paragraphs[0]



                # 有默认值则显示默认值，否则显示控件占位符

                is_vertical_choice = (

                    cell_value is None

                    and field_def

                    and field_def.field_type in ["单选（纵向）", "多选（纵向）"]

                )

                if cell_value is not None:

                    run = para.add_run(cell_value)

                    self._set_run_font(run, size=Pt(10.5))

                else:

                    # 无默认值，显示控件占位符

                    if is_vertical_choice:

                        column_cm = col_widths[col_idx] if col_idx < len(col_widths) else None
                        self._render_vertical_choices(cell, field_def, column_cm=column_cm)

                    elif field_def.field_type in ["单选", "多选"]:

                        column_cm = col_widths[col_idx] if col_idx < len(col_widths) else None
                        self._render_choice_field(para, field_def, column_cm=column_cm)

                    else:

                        # inline 整格文本填写线：按该列实际宽度自适应（不换行），
                        # 与前端 getInlineRows 共享 compute_fill_line_char_count 公式以逐字一致。
                        inline_fill_chars = (
                            compute_fill_line_char_count(col_widths[col_idx])
                            if col_idx < len(col_widths) else None
                        )
                        run = para.add_run(
                            self._render_field_control(field_def, fill_line_chars=inline_fill_chars)
                        )

                        self._set_run_font(run, size=Pt(10.5))



                # 段落格式：单行 1cm 所需上下间距，固定 15.6pt 行距

                self._apply_cell_paragraph_metrics(
                    para, space_before=not is_vertical_choice, space_after=not is_vertical_choice
                )

                para.alignment = WD_ALIGN_PARAGRAPH.LEFT

                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER



                # 应用底纹颜色和文字颜色

                marked_field = marked_fields[col_idx]

                if marked_field:

                    if marked_field.bg_color:

                        self._apply_cell_shading(cell, marked_field.bg_color)

                    if marked_field.text_color:

                        text_color = RGBColor.from_string(marked_field.text_color)

                        for para in cell.paragraphs:

                            for run in para.runs:

                                self._set_run_font(run, color=text_color)



    def _render_field_control(self, field_def, fill_line_chars: int | None = None) -> str:

        """渲染字段控件文本。

        Args:
            fill_line_chars: 文本/标签及选项尾部填写线场景的下划线根数。None 时回退到旧的
                固定 16 个（整格填写线）或 6 个（选项尾线），保持未接入列宽调用方兼容。
        """

        field_type = field_def.field_type or ""
        if field_type == "复选":
            return f"□{resolve_checkbox_label(field_def)}"

        fill_line = "_" * fill_line_chars if fill_line_chars else "________________"



        # 获取单位符号

        unit_text = ""

        if hasattr(field_def, "unit") and field_def.unit:

            unit_text = f" {field_def.unit.symbol}"



        if field_type == "单选":

            return self._render_single_choice(field_def, fill_line_chars=fill_line_chars)

        elif field_type == "多选":

            return self._render_multi_choice(field_def, fill_line_chars=fill_line_chars)

        elif field_type in ["单选（纵向）", "下拉框"]:

            return self._render_single_choice_vertical(field_def, fill_line_chars=fill_line_chars)

        elif field_type == "多选（纵向）":

            return self._render_multi_choice_vertical(field_def, fill_line_chars=fill_line_chars)

        elif field_type == "日期":

            return "|__|__|__|__|年|__|__|月|__|__|日"

        elif field_type == "日期时间":

            fmt = (getattr(field_def, "date_format", "") or "").lower()

            if "ss" in fmt:

                return "|__|__|__|__|年|__|__|月|__|__|日  |__|__|时|__|__|分|__|__|秒"

            return "|__|__|__|__|年|__|__|月|__|__|日  |__|__|时|__|__|分"

        elif field_type == "时间":

            fmt = (getattr(field_def, "date_format", "") or "").lower()

            if "ss" in fmt:

                return "|__|__|时|__|__|分|__|__|秒"

            return "|__|__|时|__|__|分"

        elif field_type == "数值":

            integers = field_def.integer_digits or 10

            decimals = field_def.decimal_digits if field_def.decimal_digits is not None else 2

            # 生成 |__|__| 格式

            parts = []

            for _ in range(integers):

                parts.append("|__|")

            line = "".join(parts)

            if decimals > 0:

                line += "."

                for _ in range(decimals):

                    line += "|__|"

            return line + unit_text

        elif field_type in ["文本", "标签"]:

            return fill_line + unit_text

        else:

            return fill_line



    def _render_single_choice(self, field_def, fill_line_chars: int | None = None) -> str:

        """渲染单选控件"""

        options = self._get_option_labels(field_def, fill_line_chars=fill_line_chars)

        if not options:

            return "________________"

        return "  ".join([f"○{opt}" for opt in options])



    def _render_single_choice_vertical(self, field_def, fill_line_chars: int | None = None) -> str:

        """渲染纵向单选控件"""

        options = self._get_option_labels(field_def, fill_line_chars=fill_line_chars)

        if not options:

            return "________________"

        return "\n".join([f"○{opt}" for opt in options])



    def _render_multi_choice(self, field_def, fill_line_chars: int | None = None) -> str:

        """渲染多选控件"""

        options = self._get_option_labels(field_def, fill_line_chars=fill_line_chars)

        if not options:

            return "________________"

        return "  ".join([f"□{opt}" for opt in options])



    def _render_multi_choice_vertical(self, field_def, fill_line_chars: int | None = None) -> str:

        """渲染纵向多选控件"""

        options = self._get_option_labels(field_def, fill_line_chars=fill_line_chars)

        if not options:

            return "________________"

        return "\n".join([f"□{opt}" for opt in options])



    def _choice_trailing_fill_chars(
        self,
        label: str,
        fill_line_chars: int | None = None,
        column_cm: float | None = None,
    ) -> int:
        if column_cm is not None:
            return compute_choice_trailing_fill_char_count(column_cm, label)
        if fill_line_chars is None:
            return 6
        marker_label_count = math.ceil(compute_choice_atom_weight(label or "", False))
        return max(0, fill_line_chars - marker_label_count)



    def _render_vertical_choices(
        self,
        cell,
        field_def,
        fill_line_chars: int | None = None,
        column_cm: float | None = None,
    ):

        """纵向排列选项：每个选项独占单元格内一个独立段落。



        choice atom：选项文本 + 尾部填写线作为原子 token，不可拆行。

        """

        field_type = field_def.field_type

        option_data = self._get_option_data(field_def)

        if not option_data:

            run = cell.paragraphs[0].add_run("________________")

            self._set_run_font(run, size=Pt(10.5))

            return



        # 单元格上下加内边距，避免纵向选项紧贴上/下框线（与段落间距正交，行仍随内容自然增高）
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = OxmlElement("w:tcMar")
        for side in ("top", "bottom"):
            node = OxmlElement(f"w:{side}")
            node.set(qn("w:w"), str(int(self.CELL_VPAD_PT * 20)))
            node.set(qn("w:type"), "dxa")
            tcMar.append(node)
        tcPr.append(tcMar)

        symbol = "○" if "单选" in field_type else "□"

        for idx, (label, has_trailing) in enumerate(option_data):

            if idx == 0:

                para = cell.paragraphs[0]

                para.paragraph_format.space_before = Pt(0)

                para.paragraph_format.space_after = Pt(0)

                self._apply_exact_line_spacing(para)

                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            else:

                para = cell.add_paragraph()

                # 非首项加段前间距，使纵向选项之间留出与预览一致的间隔
                para.paragraph_format.space_before = Pt(self.VERTICAL_OPTION_GAP_PT)

                para.paragraph_format.space_after = Pt(0)

                self._apply_exact_line_spacing(para)

                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY



            # 关闭网格吸附，使各纵向选项之间的间距在 Word 中均匀呈现
            self._disable_snap_to_grid(para)

            symbol_run = para.add_run(symbol)

            self._set_run_font(symbol_run, size=Pt(10.5))

            symbol_run.font.name = self.FONT_EAST_ASIA

            rPr = symbol_run._element.get_or_add_rPr()

            rFonts = rPr.get_or_add_rFonts()

            rFonts.set(qn("w:ascii"), self.FONT_EAST_ASIA)

            rFonts.set(qn("w:hAnsi"), self.FONT_EAST_ASIA)

            rFonts.set(qn("w:eastAsia"), self.FONT_EAST_ASIA)



            # choice atom：选项文本 + 尾部填写线作为原子 token

            if has_trailing:

                trailing_fill = "_" * self._choice_trailing_fill_chars(
                    label,
                    fill_line_chars=fill_line_chars,
                    column_cm=column_cm,
                )
                atom_text = label + trailing_fill

                atom_run = para.add_run(atom_text)

                self._set_run_font(atom_run, size=Pt(10.5))

            else:

                # 普通选项文本

                opt_run = para.add_run(label)

                self._set_run_font(opt_run, size=Pt(10.5))



    def _render_choice_field(
        self,
        paragraph,
        field_def,
        fill_line_chars: int | None = None,
        column_cm: float | None = None,
    ):

        """渲染单选或多选字段，确保○□符号使用宋体



        choice atom：选项文本 + 尾部填写线作为原子 token，不可拆行。

        """

        field_type = field_def.field_type

        option_data = self._get_option_data(field_def)

        # 没有选项时显示下划线占位符

        if not option_data:

            run = paragraph.add_run("________________")

            self._set_run_font(run, size=Pt(10.5))

            return



        symbol = "○" if "单选" in field_type else "□"

        # 横向所有选项共享一行：尾线按扣除全部选项 marker+label+分隔符后的剩余宽计算，
        # 平均分给带尾线的选项，避免单个尾线按整列计算导致整行换行。
        horizontal_trailing = (
            compute_horizontal_choice_trailing_fill_chars(column_cm, option_data)
            if column_cm is not None
            else None
        )

        for idx, (label, has_trailing) in enumerate(option_data):

            if idx > 0:

                # 横向排列，添加空格分隔

                space_run = paragraph.add_run("  ")

                self._set_run_font(space_run, size=Pt(10.5))



            # 添加符号run，使用宋体

            symbol_run = paragraph.add_run(symbol)

            self._set_run_font(symbol_run, size=Pt(10.5))

            # 强制符号使用宋体

            symbol_run.font.name = self.FONT_EAST_ASIA

            rPr = symbol_run._element.get_or_add_rPr()

            rFonts = rPr.get_or_add_rFonts()

            rFonts.set(qn("w:ascii"), self.FONT_EAST_ASIA)

            rFonts.set(qn("w:hAnsi"), self.FONT_EAST_ASIA)

            rFonts.set(qn("w:eastAsia"), self.FONT_EAST_ASIA)



            # choice atom：选项文本 + 尾部填写线作为原子 token

            if has_trailing:

                fill_count = (
                    horizontal_trailing
                    if horizontal_trailing is not None
                    else self._choice_trailing_fill_chars(
                        label,
                        fill_line_chars=fill_line_chars,
                        column_cm=column_cm,
                    )
                )
                trailing_fill = "_" * fill_count
                atom_text = label + trailing_fill

                atom_run = paragraph.add_run(atom_text)

                self._set_run_font(atom_run, size=Pt(10.5))

            else:

                # 普通选项文本

                opt_run = paragraph.add_run(label)

                self._set_run_font(opt_run, size=Pt(10.5))



    def _get_option_labels(
        self,
        field_def,
        fill_line_chars: int | None = None,
        column_cm: float | None = None,
    ) -> list:

        """获取选项标签列表，trailing_underscore=1 时在标签末尾拼接下划线



        若标签文本本身已以下划线结尾，则不重复追加（兼容 docx 导入的字面下划线场景）

        """

        labels = []
        for label, has_trailing in self._get_option_data(field_def):
            if has_trailing and not label.endswith("_"):
                trailing_fill = "_" * self._choice_trailing_fill_chars(
                    label,
                    fill_line_chars=fill_line_chars,
                    column_cm=column_cm,
                )
                labels.append(f"{label}{trailing_fill}")
            else:
                labels.append(label)
        return labels



    def _get_option_data(self, field_def) -> List[Tuple[str, bool]]:

        """获取选项数据列表：(原始标签文本, 是否有后加下划线)



        排序规则：order_index 为主，id 为稳定回退键。

        """

        if not hasattr(field_def, "codelist") or not field_def.codelist:

            return []

        if not hasattr(field_def.codelist, "options") or not field_def.codelist.options:

            return []

        # 按 order_index 排序，缺失时回退到 id

        options = sorted(

            field_def.codelist.options,

            key=lambda o: (o.order_index if o.order_index is not None else float('inf'), o.id or 0)

        )

        result: List[Tuple[str, bool]] = []

        for opt in options:

            if not opt.decode:

                continue

            result.append((opt.decode, bool(getattr(opt, "trailing_underscore", 0))))

        return result



    def _add_fill_line_run(self, paragraph, length: int = 6):

        """添加填写线 run（纯下划线字符，与文本字段填写线风格一致）"""

        fill_text = "_" * length

        run = paragraph.add_run(fill_text)

        self._set_run_font(run, size=Pt(10.5))

        return run



    def _apply_document_style(self, doc: Document):

        """统一文档字体：中文宋体，英文Times New Roman，并设置标题样式"""

        # 更新基础样式

        for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:

            if style_name not in doc.styles:

                continue

            style = doc.styles[style_name]

            style.font.name = self.FONT_ASCII

            rPr = style.element.get_or_add_rPr()

            rFonts = rPr.get_or_add_rFonts()

            rFonts.set(qn("w:ascii"), self.FONT_ASCII)

            rFonts.set(qn("w:hAnsi"), self.FONT_ASCII)

            rFonts.set(qn("w:eastAsia"), self.FONT_EAST_ASIA)



        # 设置 Heading 1 为 14pt 粗体

        if "Heading 1" in doc.styles:

            h1_style = doc.styles["Heading 1"]

            h1_style.font.size = Pt(14)

            h1_style.font.bold = True

            h1_style.font.color.rgb = RGBColor(0, 0, 0)

            h1_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT



        # 新增自定义样式：封面信息表格

        if "CoverInfo" not in doc.styles:

            cover_style = doc.styles.add_style("CoverInfo", WD_STYLE_TYPE.PARAGRAPH)

            cover_style.font.size = Pt(10)

            cover_style.font.name = self.FONT_ASCII

            rPr = cover_style.element.get_or_add_rPr()

            rFonts = rPr.get_or_add_rFonts()

            rFonts.set(qn("w:ascii"), self.FONT_ASCII)

            rFonts.set(qn("w:hAnsi"), self.FONT_ASCII)

            rFonts.set(qn("w:eastAsia"), self.FONT_EAST_ASIA)

            cover_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

            self._apply_cell_paragraph_metrics(cover_style)



        # 新增自定义样式：访视分布图表格

        if "VisitFlow" not in doc.styles:

            visit_style = doc.styles.add_style("VisitFlow", WD_STYLE_TYPE.PARAGRAPH)

            visit_style.font.size = Pt(10.5)

            visit_style.font.name = self.FONT_ASCII

            rPr = visit_style.element.get_or_add_rPr()

            rFonts = rPr.get_or_add_rFonts()

            rFonts.set(qn("w:ascii"), self.FONT_ASCII)

            rFonts.set(qn("w:hAnsi"), self.FONT_ASCII)

            rFonts.set(qn("w:eastAsia"), self.FONT_EAST_ASIA)

            self._apply_cell_paragraph_metrics(visit_style)



        # 新增自定义样式：表单标签字段

        if "FormLabel" not in doc.styles:

            label_style = doc.styles.add_style("FormLabel", WD_STYLE_TYPE.PARAGRAPH)

            label_style.font.size = Pt(10.5)

            label_style.font.name = self.FONT_ASCII

            rPr = label_style.element.get_or_add_rPr()

            rFonts = rPr.get_or_add_rFonts()

            rFonts.set(qn("w:ascii"), self.FONT_ASCII)

            rFonts.set(qn("w:hAnsi"), self.FONT_ASCII)

            rFonts.set(qn("w:eastAsia"), self.FONT_EAST_ASIA)

            self._apply_cell_paragraph_metrics(label_style)



        # 新增自定义样式：适用访视 footer

        if "ApplicableVisits" not in doc.styles:

            applicable_style = doc.styles.add_style("ApplicableVisits", WD_STYLE_TYPE.PARAGRAPH)

            applicable_style.font.size = Pt(10.5)

            applicable_style.font.name = self.FONT_ASCII

            rPr = applicable_style.element.get_or_add_rPr()

            rFonts = rPr.get_or_add_rFonts()

            rFonts.set(qn("w:ascii"), self.FONT_ASCII)

            rFonts.set(qn("w:hAnsi"), self.FONT_ASCII)

            rFonts.set(qn("w:eastAsia"), self.FONT_EAST_ASIA)

            self._apply_cell_paragraph_metrics(applicable_style)

            applicable_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT



    def _apply_grid_table_style(self, table):

        """为表格应用网格边框样式。



        添加表级 tblBorders（含 insideH/insideV）确保 Word 能稳定渲染网格边框，

        同时为每个单元格设置 tcBorders 作为冗余保障。

        """

        table.style = None



        # 添加表级边框（关键：insideH/insideV 确保内部网格线可见）

        tblPr = table._tbl.tblPr

        if tblPr is None:

            tblPr = OxmlElement('w:tblPr')

            table._tbl.insert(0, tblPr)



        tblBorders = tblPr.find(qn('w:tblBorders'))

        if tblBorders is None:

            tblBorders = OxmlElement('w:tblBorders')

            tblPr.append(tblBorders)



        for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:

            existing = tblBorders.find(qn(f'w:{border_name}'))

            if existing is not None:

                tblBorders.remove(existing)

            border_elm = OxmlElement(f'w:{border_name}')

            border_elm.set(qn('w:val'), 'single')

            border_elm.set(qn('w:sz'), '4')

            border_elm.set(qn('w:space'), '0')

            border_elm.set(qn('w:color'), 'auto')

            tblBorders.append(border_elm)



        # 单元格级边框作为冗余保障

        for row in table.rows:

            for cell in row.cells:

                self._apply_cell_borders(cell)



    def _apply_cell_borders(self, cell):

        """为单元格应用边框"""

        tcPr = cell._tc.get_or_add_tcPr()

        tcBorders = tcPr.find(qn('w:tcBorders'))

        if tcBorders is None:

            tcBorders = OxmlElement('w:tcBorders')

            tcPr.append(tcBorders)



        for border_name in ["top", "left", "bottom", "right"]:

            border_elm = OxmlElement(f'w:{border_name}')

            border_elm.set(qn('w:val'), 'single')

            border_elm.set(qn('w:sz'), '4')

            border_elm.set(qn('w:space'), '0')

            border_elm.set(qn('w:color'), 'auto')



            existing_border = tcBorders.find(qn(f'w:{border_name}'))

            if existing_border is not None:

                tcBorders.remove(existing_border)

            tcBorders.append(border_elm)



    def _apply_cell_shading(self, cell, color_hex: str):

        """为单元格应用颜色底纹"""

        shading_elm = OxmlElement('w:shd')

        shading_elm.set(qn('w:fill'), color_hex)

        cell._tc.get_or_add_tcPr().append(shading_elm)



    def _apply_cover_page_table_style(self, table):

        """封面信息表格专用样式。"""

        table.style = None

        table.autofit = False



        tblPr = table._tbl.tblPr

        if tblPr is None:

            tblPr = OxmlElement('w:tblPr')

            table._tbl.insert(0, tblPr)



        tblBorders = tblPr.find(qn('w:tblBorders'))

        if tblBorders is None:

            tblBorders = OxmlElement('w:tblBorders')

            tblPr.append(tblBorders)



        for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:

            border_elm = OxmlElement(f'w:{border_name}')

            border_elm.set(qn('w:val'), 'nil')

            border_elm.set(qn('w:sz'), '0')

            tblBorders.append(border_elm)



        if table.columns:

            for col in table.columns:

                col.width = Cm(5)



        for row in table.rows:

            for cell in row.cells:

                for para in cell.paragraphs:

                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                self._remove_cell_borders(cell)



    def _remove_cell_borders(self, cell):

        """移除单元格边框"""

        tcPr = cell._tc.get_or_add_tcPr()

        tcBorders = tcPr.find(qn('w:tcBorders'))

        if tcBorders is None:

            tcBorders = OxmlElement('w:tcBorders')

            tcPr.append(tcBorders)



        for border_name in ["top", "left", "bottom", "right"]:

            border_elm = OxmlElement(f'w:{border_name}')

            border_elm.set(qn('w:val'), 'nil')

            border_elm.set(qn('w:sz'), '0')

            border_elm.set(qn('w:space'), '0')

            border_elm.set(qn('w:color'), 'auto')



            existing_border = tcBorders.find(qn(f'w:{border_name}'))

            if existing_border is not None:

                tcBorders.remove(existing_border)

            tcBorders.append(border_elm)



    def _make_picture_float(self, picture):

        """设置图片浮于文字上方"""

        from docx.oxml import parse_xml



        # 获取inline元素

        inline = picture._inline



        # 创建完整的anchor XML结构

        anchor_xml = '''

        <wp:anchor xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"

                   xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"

                   distT="0" distB="0" distL="114300" distR="114300" simplePos="0"

                   relativeHeight="251658240" behindDoc="0" locked="0" layoutInCell="1" allowOverlap="1">

            <wp:simplePos x="0" y="0"/>

            <wp:positionH relativeFrom="column">

                <wp:posOffset>0</wp:posOffset>

            </wp:positionH>

            <wp:positionV relativeFrom="paragraph">

                <wp:posOffset>-288000</wp:posOffset>

            </wp:positionV>

            <wp:extent cx="{cx}" cy="{cy}"/>

            <wp:effectExtent l="0" t="0" r="0" b="0"/>

            <wp:wrapNone/>

            <wp:docPr id="{id}" name="图片 {id}"/>

            <wp:cNvGraphicFramePr/>

            {graphic}

        </wp:anchor>

        '''



        # 获取extent和graphic元素

        extent = inline.find(qn('wp:extent'))

        graphic = inline.find(qn('a:graphic'))

        docPr = inline.find(qn('wp:docPr'))



        if extent is not None and graphic is not None:

            cx = extent.get('cx')

            cy = extent.get('cy')

            pic_id = str(self._next_annotation_docpr_id())



            # 构建完整的anchor XML

            from xml.etree.ElementTree import tostring

            graphic_str = tostring(graphic, encoding='unicode')



            anchor_xml = anchor_xml.format(cx=cx, cy=cy, id=pic_id, graphic=graphic_str)

            anchor = parse_xml(anchor_xml)



            # 替换inline为anchor

            parent = inline.getparent()

            parent.replace(inline, anchor)



    def _set_run_font(self, run, size: Optional[Pt] = None, bold: Optional[bool] = None, color: Optional[RGBColor] = None):

        """設置Run的中英文字體與字號"""

        if size is not None:

            run.font.size = size

        if bold is not None:

            run.font.bold = bold

        if color is not None:

            run.font.color.rgb = color



        # 仅重新着色（未指定 size/bold）时直接返回，不重置字体，
        # 避免覆盖 _render_choice_field/_render_vertical_choices 已强制为宋体的 ○/□ 标记字体。
        if size is None and bold is None and color is not None:
            return

        run.font.name = self.FONT_ASCII

        rPr = run._element.get_or_add_rPr()

        rFonts = rPr.get_or_add_rFonts()

        rFonts.set(qn("w:ascii"), self.FONT_ASCII)

        rFonts.set(qn("w:hAnsi"), self.FONT_ASCII)

        rFonts.set(qn("w:eastAsia"), self.FONT_EAST_ASIA)

        # ○/□ 选项标记按 ascii 字体解析，统一强制宋体，覆盖默认值/字符串路径下的 Times New Roman
        if "○" in run.text or "□" in run.text:
            rFonts.set(qn("w:ascii"), self.FONT_EAST_ASIA)
            rFonts.set(qn("w:hAnsi"), self.FONT_EAST_ASIA)





def export_full_database(db_path: str) -> str:

    """使用 sqlite3.backup() 安全复制运行中数据库到临时文件，返回临时文件路径。"""

    tmp = tempfile.NamedTemporaryFile(suffix=".db", delete=False)

    tmp_path = tmp.name

    tmp.close()



    src_conn = sqlite3.connect(db_path)

    dst_conn = sqlite3.connect(tmp_path)

    try:

        src_conn.backup(dst_conn)

    finally:

        dst_conn.close()

        src_conn.close()



    return tmp_path





def _vacuum_sqlite_file(db_path: str) -> None:
    """对导出后的 SQLite 文件执行 VACUUM。"""
    conn = sqlite3.connect(db_path, isolation_level=None)
    try:
        conn.execute("VACUUM")
    finally:
        conn.close()


def export_project_database(db_path: str, project_id: int, project_name: str) -> str:
    """导出单项目数据库：先验证兼容性，再 backup 完整快照，最后裁剪非目标数据。

    Task 4.5: 在导出前验证 form_field 结构，不兼容则抛出 ExportError。
    """
    # 验证 form_field 结构兼容性
    _validate_form_field_schema(db_path)

    tmp_path = export_full_database(db_path)

    conn = sqlite3.connect(tmp_path)
    try:
        conn.execute("PRAGMA foreign_keys = ON")
        # 解除所有项目与 user 的外键关联
        conn.execute("UPDATE project SET owner_id = NULL")
        # 清除用户敏感数据
        conn.execute("DELETE FROM user")
        # 删除其他项目（级联删除关联数据）
        conn.execute("DELETE FROM project WHERE id != ?", (project_id,))
        conn.commit()
    finally:
        conn.close()

    _vacuum_sqlite_file(tmp_path)
    return tmp_path


def export_user_projects_database(db_path: str, owner_id: int, export_name: str = "user_projects") -> str:
    """导出当前用户全部项目数据库：先验证兼容性，再备份，最后保留该用户拥有的项目集合。

    Task 4.5: 在导出前验证 form_field 结构，不兼容则抛出 ExportError。
    """
    # 验证 form_field 结构兼容性
    _validate_form_field_schema(db_path)

    tmp_path = export_full_database(db_path)

    conn = sqlite3.connect(tmp_path)
    try:
        conn.execute("PRAGMA foreign_keys = ON")
        project_ids = [
            row[0]
            for row in conn.execute(
                "SELECT id FROM project WHERE owner_id = ? ORDER BY id",
                (owner_id,),
            ).fetchall()
        ]
        if not project_ids:
            raise ValueError("当前用户没有可导出的项目")

        placeholders = ",".join("?" for _ in project_ids)
        conn.execute("UPDATE project SET owner_id = NULL")
        conn.execute("DELETE FROM user")
        conn.execute(
            f"DELETE FROM project WHERE id NOT IN ({placeholders})",
            tuple(project_ids),
        )
        conn.commit()
    finally:
        conn.close()

    _vacuum_sqlite_file(tmp_path)
    return tmp_path
