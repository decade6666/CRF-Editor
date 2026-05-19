from __future__ import annotations

import json
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

FORM_HEADING_RE = re.compile(r"^\s*\d+\.\s*(.+?)\s*$")


@dataclass(frozen=True)
class TableFieldForm:
    name: str
    tables: list[list[list[str]]]


@dataclass(frozen=True)
class CellMismatch:
    form_index: int
    form_name: str
    table_index: int
    row_index: int
    cell_index: int
    kind: str
    preview: str | None
    export: str | None

    def to_dict(self) -> dict[str, Any]:
        return {
            "form_index": self.form_index,
            "form_name": self.form_name,
            "table_index": self.table_index,
            "row_index": self.row_index,
            "cell_index": self.cell_index,
            "kind": self.kind,
            "preview": self.preview,
            "export": self.export,
        }


@dataclass(frozen=True)
class TableFieldParityReport:
    preview_form_count: int
    export_form_count: int
    form_order_matches: bool
    preview_row_count: int
    export_row_count: int
    preview_cell_count: int
    export_cell_count: int
    exact_cell_count: int
    exact_row_count: int
    mismatches: list[CellMismatch]

    @property
    def exact_cell_denominator(self) -> int:
        return max(self.preview_cell_count, self.export_cell_count)

    @property
    def exact_row_denominator(self) -> int:
        return max(self.preview_row_count, self.export_row_count)

    @property
    def exact_cell_ratio(self) -> float:
        if self.exact_cell_denominator == 0:
            return 1.0
        return self.exact_cell_count / self.exact_cell_denominator

    @property
    def exact_row_ratio(self) -> float:
        if self.exact_row_denominator == 0:
            return 1.0
        return self.exact_row_count / self.exact_row_denominator

    def to_dict(self) -> dict[str, Any]:
        return {
            "preview_form_count": self.preview_form_count,
            "export_form_count": self.export_form_count,
            "form_order_matches": self.form_order_matches,
            "preview_row_count": self.preview_row_count,
            "export_row_count": self.export_row_count,
            "preview_cell_count": self.preview_cell_count,
            "export_cell_count": self.export_cell_count,
            "exact_cell_count": self.exact_cell_count,
            "exact_cell_denominator": self.exact_cell_denominator,
            "exact_cell_ratio": self.exact_cell_ratio,
            "exact_row_count": self.exact_row_count,
            "exact_row_denominator": self.exact_row_denominator,
            "exact_row_ratio": self.exact_row_ratio,
            "mismatches": [mismatch.to_dict() for mismatch in self.mismatches],
        }


def load_preview_table_fields(source: str | Path | list[Any] | dict[str, Any]) -> list[TableFieldForm]:
    if isinstance(source, str | Path):
        raw = json.loads(Path(source).read_text(encoding="utf-8"))
    else:
        raw = source

    raw_forms = raw.get("forms") if isinstance(raw, dict) else raw
    if not isinstance(raw_forms, list):
        raise ValueError("preview table fields must be a list or an object with a forms list")

    return [_parse_form(raw_form, index) for index, raw_form in enumerate(raw_forms)]


def extract_docx_form_table_fields(docx_path: str | Path) -> list[TableFieldForm]:
    doc = Document(str(docx_path))
    forms: list[TableFieldForm] = []
    current_form: TableFieldForm | None = None

    for block in _iter_block_items(doc):
        if isinstance(block, Paragraph):
            match = FORM_HEADING_RE.match(block.text.strip())
            if match:
                current_form = TableFieldForm(name=match.group(1), tables=[])
                forms.append(current_form)
        elif isinstance(block, Table) and current_form is not None:
            current_form.tables.append(_extract_table(block))

    return forms


def compare_table_field_forms(
    preview_forms: list[TableFieldForm],
    export_forms: list[TableFieldForm],
    *,
    max_mismatches: int = 100,
) -> TableFieldParityReport:
    exact_cells = 0
    exact_rows = 0
    mismatches: list[CellMismatch] = []

    for form_index in range(max(len(preview_forms), len(export_forms))):
        preview_form = _get_item(preview_forms, form_index)
        export_form = _get_item(export_forms, form_index)
        form_name = _form_name(preview_form, export_form)
        for table_index in range(max(_table_count(preview_form), _table_count(export_form))):
            preview_table = _get_table(preview_form, table_index)
            export_table = _get_table(export_form, table_index)
            for row_index in range(max(len(preview_table), len(export_table))):
                preview_row = _get_item(preview_table, row_index)
                export_row = _get_item(export_table, row_index)
                if preview_row is not None and export_row is not None and preview_row == export_row:
                    exact_rows += 1
                exact_cells += _compare_row_cells(
                    mismatches,
                    form_index,
                    form_name,
                    table_index,
                    row_index,
                    preview_row,
                    export_row,
                    max_mismatches,
                )

    return TableFieldParityReport(
        preview_form_count=len(preview_forms),
        export_form_count=len(export_forms),
        form_order_matches=[form.name for form in preview_forms] == [form.name for form in export_forms],
        preview_row_count=_count_rows(preview_forms),
        export_row_count=_count_rows(export_forms),
        preview_cell_count=_count_cells(preview_forms),
        export_cell_count=_count_cells(export_forms),
        exact_cell_count=exact_cells,
        exact_row_count=exact_rows,
        mismatches=mismatches,
    )


def _parse_form(raw_form: Any, index: int) -> TableFieldForm:
    if not isinstance(raw_form, dict):
        raise ValueError(f"form #{index} must be an object")
    name = str(raw_form.get("name") or raw_form.get("title") or "")
    raw_tables = raw_form.get("tables")
    if not isinstance(raw_tables, list):
        raise ValueError(f"form #{index} tables must be a list")
    return TableFieldForm(name=name, tables=[_parse_table(table) for table in raw_tables])


def _parse_table(raw_table: Any) -> list[list[str]]:
    if not isinstance(raw_table, list):
        raise ValueError("table must be a list of rows")
    return [[_normalize_cell_text(cell) for cell in row] for row in raw_table]


def _iter_block_items(doc: DocxDocument) -> Iterable[Paragraph | Table]:
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield Table(child, doc)


def _extract_table(table: Table) -> list[list[str]]:
    extracted: list[list[str]] = []
    for row in table.rows:
        seen_cells = []
        values: list[str] = []
        for cell in row.cells:
            if any(cell._tc is seen for seen in seen_cells):
                continue
            seen_cells.append(cell._tc)
            values.append(_normalize_cell_text(cell.text))
        extracted.append(values)
    return extracted


def _normalize_cell_text(value: Any) -> str:
    return str(value).replace("\r\n", "\n").replace("\r", "\n")


def _compare_row_cells(
    mismatches: list[CellMismatch],
    form_index: int,
    form_name: str,
    table_index: int,
    row_index: int,
    preview_row: list[str] | None,
    export_row: list[str] | None,
    max_mismatches: int,
) -> int:
    exact_cells = 0
    for cell_index in range(max(len(preview_row or []), len(export_row or []))):
        preview_cell = _get_item(preview_row or [], cell_index)
        export_cell = _get_item(export_row or [], cell_index)
        if preview_cell == export_cell and preview_cell is not None:
            exact_cells += 1
        elif len(mismatches) < max_mismatches:
            mismatches.append(_cell_mismatch(
                form_index, form_name, table_index, row_index, cell_index, preview_cell, export_cell
            ))
    return exact_cells


def _cell_mismatch(
    form_index: int,
    form_name: str,
    table_index: int,
    row_index: int,
    cell_index: int,
    preview: str | None,
    export: str | None,
) -> CellMismatch:
    if preview is None:
        kind = "missing_preview_cell"
    elif export is None:
        kind = "missing_export_cell"
    else:
        kind = "text"
    return CellMismatch(form_index, form_name, table_index, row_index, cell_index, kind, preview, export)


def _count_rows(forms: list[TableFieldForm]) -> int:
    return sum(len(table) for form in forms for table in form.tables)


def _count_cells(forms: list[TableFieldForm]) -> int:
    return sum(len(row) for form in forms for table in form.tables for row in table)


def _table_count(form: TableFieldForm | None) -> int:
    return len(form.tables) if form is not None else 0


def _get_table(form: TableFieldForm | None, index: int) -> list[list[str]]:
    if form is None or index >= len(form.tables):
        return []
    return form.tables[index]


def _get_item(items: list[Any] | None, index: int) -> Any | None:
    if items is None or index >= len(items):
        return None
    return items[index]


def _form_name(preview_form: TableFieldForm | None, export_form: TableFieldForm | None) -> str:
    if preview_form is not None:
        return preview_form.name
    if export_form is not None:
        return export_form.name
    return ""
