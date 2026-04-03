"""Word文档导入服务 - 解析eCRF Word文档，提取表单和字段结构"""
from __future__ import annotations

import html as _html
import re
import os
import uuid
import logging
from pathlib import Path
from typing import Dict, List, Optional, Tuple

from docx import Document
from docx.oxml.ns import qn
from sqlalchemy import select
from sqlalchemy.orm import Session

from src.models.form import Form
from src.models.form_field import FormField
from src.models.field_definition import FieldDefinition
from src.models.codelist import CodeList, CodeListOption
from src.models.unit import Unit
from src.utils import generate_code
from src.services.order_service import OrderService

logger = logging.getLogger(__name__)


# ── Word表格 → HTML 渲染 ──

def _get_grid_span(tc) -> int:
    """读取单元格的 gridSpan（横向合并列数），默认1"""
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        return 1
    gs = tcPr.find(qn("w:gridSpan"))
    if gs is None:
        return 1
    return max(int(gs.get(qn("w:val"), "1")), 1)


def _get_vmerge(tc) -> Optional[str]:
    """读取单元格的 vMerge 属性

    Returns:
        "restart" = 纵向合并起点
        "continue" = 被上方单元格合并（无 val 属性时也视为 continue）
        None = 无纵向合并
    """
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        return None
    vm = tcPr.find(qn("w:vMerge"))
    if vm is None:
        return None
    val = vm.get(qn("w:val"), "")
    return "restart" if val == "restart" else "continue"


def _extract_cell_text(tc) -> str:
    """从 tc XML 元素中提取纯文本，段落间用换行分隔"""
    paragraphs = []
    for p in tc.findall(qn("w:p")):
        texts = []
        for r in p.findall(qn("w:r")):
            for t in r.findall(qn("w:t")):
                if t.text:
                    texts.append(t.text)
        paragraphs.append("".join(texts))
    return "\n".join(paragraphs).strip()


def _table_to_html(table) -> str:
    """将 python-docx Table 对象转换为 HTML 字符串

    基于 OXML 级别精确处理 gridSpan（colspan）和 vMerge（rowspan），
    保留基本表格结构和内联样式，用于前端预览展示。
    """
    tbl = table._tbl
    # 确定列数：优先从 tblGrid 读取
    grid_node = tbl.find(qn("w:tblGrid"))
    if grid_node is not None:
        grid_cols = grid_node.findall(qn("w:gridCol"))
        col_count = len(grid_cols) if grid_cols else 0
    else:
        col_count = 0

    row_elements = tbl.findall(qn("w:tr"))
    row_count = len(row_elements)

    if row_count == 0:
        return ""

    # 回退：取每行 gridSpan 之和的最大值
    if col_count == 0:
        for tr in row_elements:
            tcs = tr.findall(qn("w:tc"))
            row_width = sum(_get_grid_span(tc) for tc in tcs)
            col_count = max(col_count, row_width)

    if col_count == 0:
        return ""

    # 占用矩阵：跟踪被 rowspan 占据的格子
    occupied = [[False] * col_count for _ in range(row_count)]

    # 第一遍扫描：收集所有单元格的位置、colspan、rowspan 信息
    cell_map = []  # [(row, col, colspan, rowspan, text)]
    for ri, tr in enumerate(row_elements):
        tcs = tr.findall(qn("w:tc"))
        col = 0
        for tc in tcs:
            # 跳过已被占用的列
            while col < col_count and occupied[ri][col]:
                col += 1
            if col >= col_count:
                break

            colspan = min(_get_grid_span(tc), col_count - col)
            vm = _get_vmerge(tc)

            if vm == "continue":
                # 被上方合并，标记占用但不输出
                for c in range(col, min(col + colspan, col_count)):
                    occupied[ri][c] = True
                col += colspan
                continue

            # 计算 rowspan（vm == "restart" 时向下扫描）
            rowspan = 1
            if vm == "restart":
                for scan_ri in range(ri + 1, row_count):
                    scan_col = 0
                    found = False
                    for scan_tc in row_elements[scan_ri].findall(qn("w:tc")):
                        while scan_col < col_count and occupied[scan_ri][scan_col]:
                            scan_col += 1
                        if scan_col == col:
                            scan_vm = _get_vmerge(scan_tc)
                            if scan_vm == "continue":
                                rowspan += 1
                                found = True
                            break
                        scan_colspan = _get_grid_span(scan_tc)
                        scan_col += scan_colspan
                    if not found:
                        break

            # 标记占用区域
            for dr in range(rowspan):
                for dc in range(colspan):
                    r, c = ri + dr, col + dc
                    if r < row_count and c < col_count:
                        occupied[r][c] = True

            text = _html.escape(_extract_cell_text(tc))
            cell_map.append((ri, col, colspan, rowspan, text))
            col += colspan

    # 第二遍：生成 HTML
    # 按行分组
    rows_dict: Dict[int, list] = {}
    for ri, col, colspan, rowspan, text in cell_map:
        rows_dict.setdefault(ri, []).append((col, colspan, rowspan, text))

    html_parts = [
        '<table style="border-collapse:collapse;width:100%;font-size:13px;">'
    ]
    for ri in range(row_count):
        cells = rows_dict.get(ri, [])
        # 按列顺序排列
        cells.sort(key=lambda x: x[0])
        html_parts.append("<tr>")
        for _, colspan, rowspan, text in cells:
            attrs = [
                'style="border:1px solid #bbb;padding:6px 8px;'
                'vertical-align:top;text-align:left;"'
            ]
            if colspan > 1:
                attrs.append(f'colspan="{colspan}"')
            if rowspan > 1:
                attrs.append(f'rowspan="{rowspan}"')
            # 换行符转 <br>
            display_text = text.replace("\n", "<br>") if text else ""
            html_parts.append(f'<td {" ".join(attrs)}>{display_text}</td>')
        html_parts.append("</tr>")
    html_parts.append("</table>")

    return "".join(html_parts)


# ── 字段类型检测 ──

def _detect_field_type(value_text: str) -> Tuple[str, dict]:
    """根据单元格内容格式推断字段类型和配置参数

    检测顺序：日期 → 时间 → 单选(有选项) → 多选(有选项) → 小数数值 → 整数数值 → 标签(纯文本) → 文本(空)
    """
    text = value_text.strip()
    if not text:
        return "文本", {}

    # 日期: |__|__|__|__|年|__|__|月|__|__|日
    if "年" in text and "月" in text and "日" in text and "|" in text:
        return "日期", {"date_format": "YYYY-MM-DD"}

    # 时间: |__|__|:|__|__| 或含"时""分"的格式（必须在数值检测之前）
    if re.search(r"\|__\|.*[:：].*\|__\|", text):
        # 判断是否包含秒
        colon_count = text.count(":") + text.count("：")
        if colon_count >= 2:
            return "时间", {"date_format": "HH:mm:ss"}
        return "时间", {"date_format": "HH:mm"}
    if "时" in text and "分" in text and "|" in text:
        return "时间", {"date_format": "HH:mm"}

    # 单选: ○选项1  ○选项2（选项为空时不判定为单选）
    if "○" in text:
        options = [o.strip() for o in re.split(r"○", text) if o.strip()]
        if options:
            return "单选", {"options": options}

    # 多选: □选项1  □选项2（选项为空时不判定为多选）
    if "□" in text:
        options = [o.strip() for o in re.split(r"□", text) if o.strip()]
        if options:
            return "多选", {"options": options}

    # 带小数的数值: |__|__|.|__|
    if "." in text and "|" in text:
        parts = text.split(".")
        int_digits = parts[0].count("__")
        dec_digits = parts[1].count("__") if len(parts) > 1 else 0
        return "数值", {
            "integer_digits": min(max(int_digits, 1), 20),
            "decimal_digits": min(max(dec_digits, 1), 15),
        }

    # 整数数值: |__|__|__|
    if "|__" in text:
        digits = text.count("__")
        return "数值", {"integer_digits": min(max(digits, 1), 20), "decimal_digits": 0}

    # 标签: 纯文本无输入控件标记（无|__|、○、□）
    # 只有○/□但无选项文本的也会落到这里
    return "标签", {}


def _extract_unit_from_text(text: str) -> Optional[str]:
    """从文本中提取单位（如 |__|__|.|__|℃ 中的 ℃）"""
    cleaned = re.sub(r"[|_\d.]+", "", text).strip()
    if cleaned and len(cleaned) <= 10:
        return cleaned
    return None


# ── 文档结构解析 ──

def _get_paragraph_text(element) -> str:
    """从 XML 元素中提取段落文本"""
    texts = []
    for r in element.findall(qn("w:r")):
        for t in r.findall(qn("w:t")):
            if t.text:
                texts.append(t.text)
    return "".join(texts).strip()


def _is_form_title(element) -> Optional[str]:
    """判断段落是否为表单标题

    识别策略：带有列表/编号样式的段落（ListParagraph 或自定义样式如 af6 等）。
    排除明显不是标题的段落（注释说明、长文本等）。
    """
    pPr = element.find(qn("w:pPr"))
    if pPr is None:
        return None
    text = _get_paragraph_text(element)
    if not text:
        return None

    # 排除明显不是表单标题的段落
    if len(text) > 30:
        return None
    skip_prefixes = ("注：", "注:", "周岁", "赛美斯", "太美", "方案版本号")
    if any(text.startswith(p) for p in skip_prefixes):
        return None

    # 检查是否有编号/列表样式（numPr 表示有编号）
    numPr = pPr.find(qn("w:numPr"))
    if numPr is not None:
        return text

    # 兼容 ListParagraph 及其变体样式
    pStyle = pPr.find(qn("w:pStyle"))
    if pStyle is not None:
        style_val = pStyle.get(qn("w:val"), "")
        if style_val in ("ListParagraph", "af6", "a0"):
            return text

    return None


def _is_log_row(row) -> bool:
    """检测是否为 log 行标记"""
    cells_text = [cell.text.strip() for cell in row.cells]
    return any("以下为log行" in t for t in cells_text)


def _is_horizontal_table(table) -> bool:
    """检测是否为横向多列表格（如生命体征、实验室检查）"""
    if len(table.columns) <= 2:
        return False
    for row in table.rows:
        row_text = " ".join(cell.text.strip() for cell in row.cells)
        if any(kw in row_text for kw in ["项目", "结果", "单位", "检测项名称", "测定值"]):
            return True
    return False


def _parse_simple_table(table) -> List[dict]:
    """解析简单2列表格，返回字段列表"""
    fields = []
    for row in table.rows:
        cells = row.cells
        if len(cells) < 2:
            continue
        label = cells[0].text.strip()
        value = cells[1].text.strip()
        if not label:
            continue

        # log 行标记
        if _is_log_row(row):
            fields.append({"type": "log_row"})
            continue

        field_type, config = _detect_field_type(value)
        field_info = {"label": label, "field_type": field_type, **config}

        # 标签类型：保留右侧单元格文本作为默认值
        if field_type == "标签" and value:
            field_info["default_value"] = value

        # 从值格式中提取单位（仅数值类型）
        if field_type == "数值":
            unit = _extract_unit_from_text(value)
            if unit:
                field_info["unit_symbol"] = unit

        fields.append(field_info)
    return fields


def _find_header_row_index(table) -> int:
    """找到横向表格的列头行索引"""
    for i, row in enumerate(table.rows):
        row_text = " ".join(cell.text.strip() for cell in row.cells)
        if any(kw in row_text for kw in ["项目", "结果", "单位", "检测项名称", "测定值"]):
            return i
    return -1


def _parse_horizontal_table(table) -> List[dict]:
    """解析横向多列表格（生命体征、实验室检查等）"""
    fields = []
    header_idx = _find_header_row_index(table)
    if header_idx < 0:
        return _parse_simple_table(table)

    # 列头行之前的行作为表单级字段（取前2列当作label/value）
    for i in range(header_idx):
        row = table.rows[i]
        cells = row.cells
        # 合并单元格可能导致重复，去重取前两个不同的文本
        seen = []
        for c in cells:
            t = c.text.strip()
            if t and t not in seen:
                seen.append(t)
        if len(seen) >= 2:
            label, value = seen[0], seen[1]
            ft, cfg = _detect_field_type(value)
            fields.append({"label": label, "field_type": ft, **cfg})
        elif len(seen) == 1:
            fields.append({"label": seen[0], "field_type": "文本"})

    # 解析列头，确定各列含义
    header_cells = [c.text.strip() for c in table.rows[header_idx].cells]
    # 去重（合并单元格会重复）
    col_roles = []
    for i, h in enumerate(header_cells):
        if i > 0 and h == header_cells[i - 1]:
            col_roles.append(col_roles[-1] if col_roles else "unknown")
        else:
            col_roles.append(_classify_column(h))

    # ── 辅助：从数据行取样推断列类型 ──

    def _first_non_empty_in_col(col_idx: int) -> str:
        """取指定列首个非空数据行值"""
        for ri in range(header_idx + 1, len(table.rows)):
            row = table.rows[ri]
            if col_idx >= len(row.cells):
                continue
            val = row.cells[col_idx].text.strip()
            if val:
                return val
        return ""

    def _collect_select_options(col_idx: int) -> List[str]:
        """从数据行中提取○选项（取首个含○的单元格）"""
        for ri in range(header_idx + 1, len(table.rows)):
            row = table.rows[ri]
            if col_idx >= len(row.cells):
                continue
            val = row.cells[col_idx].text.strip()
            if "○" in val:
                return [o.strip() for o in re.split(r"○", val) if o.strip()]
        return []

    def _select_field_type(header_label: str) -> str:
        """根据列头关键词决定单选子类型"""
        if any(kw in header_label for kw in ("临床意义", "异常")):
            return "单选（纵向）"
        return "单选"

    # ── 辅助：收集指定列所有数据行的值，用于 name/unit 列的默认值 ──

    def _collect_column_values(col_idx: int) -> Optional[str]:
        """收集指定列所有数据行非空值，换行拼接作为 default_value"""
        vals = []
        for ri in range(header_idx + 1, len(table.rows)):
            row = table.rows[ri]
            if col_idx >= len(row.cells):
                continue
            vals.append(row.cells[col_idx].text.strip())
        # 至少有一个非空值才返回
        if any(vals):
            return "\n".join(vals)
        return None

    # ── 列头字段：基于列头生成字段模板，不再展开数据行 ──
    seen_headers: set = set()
    for ci, role in enumerate(col_roles):
        if ci >= len(header_cells):
            continue
        header_label = header_cells[ci].strip()
        if not header_label or header_label in seen_headers:
            continue
        seen_headers.add(header_label)

        # 未知角色回退为 value
        if role not in ("name", "value", "unit", "select"):
            role = "value"

        if role == "name":
            field = {"label": header_label, "field_type": "文本", "inline_mark": True}
            dv = _collect_column_values(ci)
            if dv:
                field["default_value"] = dv
            fields.append(field)
            continue

        if role == "unit":
            field = {"label": header_label, "field_type": "文本", "inline_mark": True}
            dv = _collect_column_values(ci)
            if dv:
                field["default_value"] = dv
            fields.append(field)
            continue

        if role == "select":
            options = _collect_select_options(ci)
            ft = _select_field_type(header_label)
            field_info: dict = {"label": header_label, "field_type": ft, "inline_mark": True}
            if options:
                field_info["options"] = options
            fields.append(field_info)
            continue

        # role == "value"：从数据行抽样推断类型
        sample = _first_non_empty_in_col(ci)
        ft, cfg = _detect_field_type(sample)
        field_info = {"label": header_label, "field_type": ft or "文本", "inline_mark": True}
        field_info.update(cfg)
        # 标签类型也收集默认值（保留列数据）
        if ft == "标签":
            dv = _collect_column_values(ci)
            if dv:
                field_info["default_value"] = dv
        fields.append(field_info)

    return fields


def _classify_column(header: str) -> str:
    """根据列头文本分类列的角色"""
    h = header.strip()
    # 名称/项目列
    if h in ("项目", "检测项名称", "检查项目", "检验项目", "名称", "检测项目"):
        return "name"
    if any(kw in h for kw in ("项目名", "检查项", "检验项", "指标")):
        return "name"
    # 结果/值列
    if h in ("结果", "测定值", "检测结果", "检验结果", "数值"):
        return "value"
    # 单位列
    if h in ("单位",):
        return "unit"
    # "正常值范围下限/上限"是数值输入，不是选择列
    if any(kw in h for kw in ("下限", "上限")):
        return "value"
    # 已知选择列
    if any(kw in h for kw in [
        "临床意义", "未查", "异常", "正常值范围",
    ]):
        return "select"
    return "value"


# ── AI覆盖后的字段配置清理 ──

def _cleanup_field_config(field_info: dict, new_type: str) -> None:
    """覆盖字段类型后，清理不一致的配置属性

    例如：类型从"数值"改为"文本"时，需移除 integer_digits 等数值属性。
    """
    # 非数值类型：清除数值相关属性
    if new_type not in ("数值",):
        field_info.pop("integer_digits", None)
        field_info.pop("decimal_digits", None)
        field_info.pop("unit_symbol", None)
    # 非选择类型：清除选项
    if new_type not in ("单选", "多选", "单选（纵向）", "多选（纵向）"):
        field_info.pop("options", None)
    # 非日期/时间类型：清除日期格式
    if new_type not in ("日期", "时间"):
        field_info.pop("date_format", None)


# ── 主服务类 ──

class DocxImportService:
    """Word文档导入服务"""

    TEMP_DIR = "uploads/docx_temp"

    def __init__(self, session: Session):
        self.session = session

    # 上传文件大小上限（10MB）
    MAX_FILE_SIZE = 10 * 1024 * 1024

    @staticmethod
    def save_temp_file(content: bytes, filename: str) -> Tuple[str, str]:
        """保存上传的文件到临时目录，返回 (temp_id, file_path)

        对文件名做安全处理，防止路径遍历攻击。
        """
        if len(content) > DocxImportService.MAX_FILE_SIZE:
            raise ValueError(
                f"文件大小超过限制（最大 {DocxImportService.MAX_FILE_SIZE // 1024 // 1024}MB）"
            )
        temp_dir = Path(DocxImportService.TEMP_DIR)
        temp_dir.mkdir(parents=True, exist_ok=True)
        temp_id = uuid.uuid4().hex[:12]
        # 安全处理：只取文件名部分，过滤路径分隔符
        basename = os.path.basename(filename).replace("..", "")
        safe_name = f"{temp_id}_{basename}"
        file_path = str(temp_dir / safe_name)
        with open(file_path, "wb") as f:
            f.write(content)
        return temp_id, file_path

    @staticmethod
    def get_temp_path(temp_id: str) -> Optional[str]:
        """根据 temp_id 查找临时文件路径（只返回.docx文件，不返回目录）"""
        temp_dir = Path(DocxImportService.TEMP_DIR)
        if not temp_dir.exists():
            return None
        for f in temp_dir.iterdir():
            if f.name.startswith(temp_id) and f.is_file() and f.suffix == ".docx":
                return str(f)
        return None

    @staticmethod
    def cleanup_temp(temp_id: str) -> None:
        """清理临时文件（安全清理，避免异常覆盖原始错误）"""
        try:
            path = DocxImportService.get_temp_path(temp_id)
            if path and os.path.exists(path):
                path_obj = Path(path)
                if path_obj.is_file():
                    os.remove(path)
                elif path_obj.is_dir():
                    import shutil
                    shutil.rmtree(path)
        except Exception as e:
            logger.warning("清理临时文件失败: temp_id=%s, 错误: %s", temp_id, str(e))

    # ── 解析预览 ──

    @staticmethod
    def parse_preview(file_path: str) -> List[dict]:
        """解析Word文档，返回表单预览列表

        Returns:
            [{"index": 0, "name": "知情同意", "field_count": 4}, ...]
        """
        doc = Document(file_path)
        forms = DocxImportService._extract_forms(doc)
        return [
            {
                "index": i,
                "name": f["name"],
                "field_count": len([
                    fd for fd in f["fields"] if fd.get("type") != "log_row"
                ]),
            }
            for i, f in enumerate(forms)
        ]

    @staticmethod
    def parse_full(file_path: str) -> List[dict]:
        """解析Word文档，返回完整的表单+字段结构"""
        doc = Document(file_path)
        return DocxImportService._extract_forms(doc)

    # ── 核心解析：按顺序遍历文档元素，匹配标题与表格 ──

    @staticmethod
    def _extract_forms(doc: Document) -> List[dict]:
        """遍历文档body元素，将表单标题与后续表格配对"""
        body = doc.element.body
        forms = []
        current_title = None
        table_idx = 0
        tables = doc.tables
        skip_tables = 2  # 跳过封面信息表和访视分布图

        for child in body:
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

            if tag == "p":
                title = _is_form_title(child)
                if title:
                    current_title = title

            elif tag == "tbl":
                if table_idx < skip_tables:
                    table_idx += 1
                    continue

                real_idx = table_idx
                table_idx += 1

                if real_idx >= len(tables):
                    continue

                table = tables[real_idx]
                form_name = current_title or f"未命名表单_{real_idx}"

                # 根据表格类型选择解析策略
                if _is_horizontal_table(table):
                    fields = _parse_horizontal_table(table)
                else:
                    fields = _parse_simple_table(table)

                if fields:
                    forms.append({
                        "name": form_name,
                        "fields": fields,
                        "raw_html": _table_to_html(table),
                    })
                current_title = None

        return forms

    # ── 导入执行 ──

    def import_forms(
        self,
        target_project_id: int,
        file_path: str,
        form_indices: List[int],
        ai_overrides: Optional[list] = None,
    ) -> dict:
        """将解析出的表单写入数据库

        Args:
            ai_overrides: AI建议覆盖列表，每项含 form_index 和 overrides
        """
        all_forms = self.parse_full(file_path)
        # 过滤非法索引：去重 + 排除负数 + 排除越界
        valid_indices = list(dict.fromkeys(
            i for i in form_indices if 0 <= i < len(all_forms)
        ))
        selected = [all_forms[i] for i in valid_indices]
        if not selected:
            return {"imported_form_count": 0, "detail": []}

        # 构建 AI 覆盖映射：{form_index: {field_index: field_type}}
        override_map: Dict[int, Dict[int, str]] = {}
        if ai_overrides:
            for fo in ai_overrides:
                fi = fo.form_index if hasattr(fo, "form_index") else fo.get("form_index")
                overrides = fo.overrides if hasattr(fo, "overrides") else fo.get("overrides", [])
                field_map = {}
                for o in overrides:
                    idx = o.index if hasattr(o, "index") else o.get("index")
                    ft = o.field_type if hasattr(o, "field_type") else o.get("field_type")
                    field_map[idx] = ft
                override_map[fi] = field_map

        s = self.session
        summary = {"imported_form_count": 0, "detail": []}

        # 缓存已有数据，避免重复创建
        existing_forms = {
            f.name for f in s.scalars(
                select(Form).where(Form.project_id == target_project_id)
            ).all()
        }
        existing_units: Dict[str, int] = {
            u.symbol: u.id for u in s.scalars(
                select(Unit).where(Unit.project_id == target_project_id)
            ).all()
        }
        existing_codelists: Dict[str, int] = {
            c.name: c.id for c in s.scalars(
                select(CodeList).where(CodeList.project_id == target_project_id)
            ).all()
        }
        existing_vars: set = {
            fd.variable_name for fd in s.scalars(
                select(FieldDefinition).where(
                    FieldDefinition.project_id == target_project_id
                )
            ).all()
        }

        for form_index, form_data in zip(valid_indices, selected):
            field_overrides = override_map.get(form_index, {})
            result = self._create_form(
                s, target_project_id, form_data,
                existing_forms, existing_units,
                existing_codelists, existing_vars,
                field_overrides=field_overrides,
            )
            summary["imported_form_count"] += 1
            summary["detail"].append(result)

        # 显式flush，让数据库约束错误在此处抛出，而不是延迟到事务提交
        s.flush()
        return summary

    def _create_form(
        self,
        s: Session,
        project_id: int,
        form_data: dict,
        existing_forms: set,
        existing_units: Dict[str, int],
        existing_codelists: Dict[str, int],
        existing_vars: set,
        field_overrides: Optional[Dict[int, str]] = None,
    ) -> dict:
        """创建单个表单及其字段"""
        form_name = form_data["name"]
        # 表单名冲突处理
        if form_name in existing_forms:
            base = form_name
            suffix = "_导入"
            candidate = f"{base}{suffix}"
            idx = 2
            while candidate in existing_forms:
                candidate = f"{base}{suffix}{idx}"
                idx += 1
            form_name = candidate
        existing_forms.add(form_name)

        new_form = Form(
            project_id=project_id,
            name=form_name,
            code=generate_code("FORM"),
            order_index=OrderService.get_next_order(s, Form, Form.project_id == project_id),
        )
        s.add(new_form)
        s.flush()

        field_count = 0
        order_index = 0

        for fi, field_info in enumerate(form_data["fields"]):
            order_index += 1

            # log 行
            if field_info.get("type") == "log_row":
                s.add(FormField(
                    form_id=new_form.id,
                    field_definition_id=None,
                    is_log_row=1,
                    order_index=order_index,
                ))
                continue

            # 应用 AI 建议覆盖：替换字段类型并清理不一致的配置
            if field_overrides and fi in field_overrides:
                override_type = field_overrides[fi]
                logger.info(
                    "AI覆盖: 表单=%s 字段#%d '%s' -> %s",
                    form_name, fi, field_info.get("field_type"), override_type,
                )
                field_info = dict(field_info)  # 浅拷贝，避免污染原数据
                field_info["field_type"] = override_type
                _cleanup_field_config(field_info, override_type)

            # 创建字段定义
            try:
                fd = self._create_field_definition(
                    s, project_id, field_info,
                    existing_units, existing_codelists, existing_vars,
                )
                if fd is None:
                    continue
            except Exception as e:
                logger.error(
                    "创建字段定义失败: 表单=%s 字段#%d label='%s' type=%s, 错误: %s",
                    form_name, fi, field_info.get("label"), field_info.get("field_type"), str(e)
                )
                raise

            # 创建表单字段关联
            s.add(FormField(
                form_id=new_form.id,
                field_definition_id=fd.id,
                order_index=order_index,
                inline_mark=1 if field_info.get("inline_mark") else 0,
                default_value=field_info.get("default_value"),
            ))
            field_count += 1

        return {"name": form_name, "field_count": field_count}

    def _create_field_definition(
        self,
        s: Session,
        project_id: int,
        field_info: dict,
        existing_units: Dict[str, int],
        existing_codelists: Dict[str, int],
        existing_vars: set,
    ) -> Optional[FieldDefinition]:
        """创建单个字段定义，自动处理 CodeList 和 Unit"""
        label = field_info.get("label", "")
        field_type = field_info.get("field_type", "文本")
        if not label:
            return None

        # 生成唯一变量名
        var_base = generate_code("FIELD")
        var_name = var_base
        while var_name in existing_vars:
            var_name = generate_code("FIELD")
        existing_vars.add(var_name)

        # 处理单位
        unit_id = None
        unit_symbol = field_info.get("unit_symbol")
        if unit_symbol:
            if unit_symbol in existing_units:
                unit_id = existing_units[unit_symbol]
            else:
                new_unit = Unit(
                    project_id=project_id,
                    symbol=unit_symbol,
                    code=generate_code("UNIT"),
                    order_index=OrderService.get_next_order(s, Unit, Unit.project_id == project_id),
                )
                s.add(new_unit)
                s.flush()
                unit_id = new_unit.id
                existing_units[unit_symbol] = unit_id

        # 处理选项列表
        codelist_id = None
        options = field_info.get("options")
        if field_type in ("单选", "多选", "单选（纵向）", "多选（纵向）") and options:
            normalized_options = [str(opt_text or "").strip() for opt_text in options]
            normalized_options = [opt_text for opt_text in normalized_options if opt_text]
            if normalized_options:
                cl_name = f"{label}_选项"
                # 限制CodeList名称长度不超过255字符
                if len(cl_name) > 255:
                    cl_name = cl_name[:252] + "..."
                if cl_name in existing_codelists:
                    codelist_id = existing_codelists[cl_name]
                else:
                    new_cl = CodeList(
                        project_id=project_id,
                        name=cl_name,
                        code=generate_code("CL"),
                    )
                    s.add(new_cl)
                    s.flush()
                    codelist_id = new_cl.id
                    existing_codelists[cl_name] = codelist_id
                    for i, opt_text in enumerate(normalized_options, start=1):
                        s.add(CodeListOption(
                            codelist_id=codelist_id,
                            code=f"C.{i}",
                            decode=opt_text,
                            trailing_underscore=0,
                            order_index=i,
                        ))

        # 创建字段定义
        fd = FieldDefinition(
            project_id=project_id,
            variable_name=var_name,
            label=label,
            field_type=field_type,
            integer_digits=field_info.get("integer_digits"),
            decimal_digits=field_info.get("decimal_digits"),
            date_format=field_info.get("date_format"),
            codelist_id=codelist_id,
            unit_id=unit_id,
            order_index=OrderService.get_next_order(s, FieldDefinition, FieldDefinition.project_id == project_id),
        )
        s.add(fd)
        s.flush()
        return fd
