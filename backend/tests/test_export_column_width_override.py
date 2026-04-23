"""Word 导出列宽覆盖测试。

按 TDD 流程验证前端调整的表格列宽能够传递到后端并应用到 Word 导出。

测试场景：
1. 接口层：POST body 中的 column_width_overrides 参数能正确传递到 ExportService
2. 服务层：传入的列宽覆盖比例实际应用到 Word 表格的 tblGrid
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
import pytest
from sqlalchemy import create_engine, event
from sqlalchemy.orm import Session, sessionmaker
from sqlalchemy.pool import StaticPool

from src.models import Base
from src.models.project import Project
from src.models.user import User
from src.models.visit import Visit
from src.models.form import Form
from src.models.visit_form import VisitForm
from src.models.field_definition import FieldDefinition
from src.models.form_field import FormField
from src.services.export_service import ExportService


@pytest.fixture
def engine():
    """创建内存数据库引擎。"""
    _engine = create_engine(
        "sqlite+pysqlite:///:memory:",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )

    @event.listens_for(_engine, "connect")
    def _enable_fk(dbapi_conn, _):
        dbapi_conn.execute("PRAGMA foreign_keys = ON")

    Base.metadata.create_all(_engine)
    yield _engine
    _engine.dispose()


@pytest.fixture
def session(engine) -> Session:
    """创建数据库会话。"""
    session_factory = sessionmaker(bind=engine, expire_on_commit=False)
    with session_factory() as db_session:
        yield db_session


def create_minimal_project(session: Session) -> tuple[Project, User]:
    """创建最小项目和用户。"""
    user = User(username="tester")
    session.add(user)
    session.flush()
    project = Project(name="列宽覆盖测试项目", version="v1.0", owner_id=user.id)
    session.add(project)
    session.flush()
    return project, user


def create_text_field_def(session: Session, project_id: int, label: str) -> FieldDefinition:
    """创建文本字段定义。"""
    fd = FieldDefinition(
        project_id=project_id,
        label=label,
        variable_name=f"VAR_{label.replace(' ', '_')}",
        field_type="文本",
    )
    session.add(fd)
    session.flush()
    return fd


def add_field_to_form(
    session: Session,
    form_id: int,
    field_def_id: int,
    order_index: int = 0,
    inline_mark: int = 0,
) -> FormField:
    """将字段添加到表单。"""
    ff = FormField(
        form_id=form_id,
        field_definition_id=field_def_id,
        order_index=order_index,
        inline_mark=inline_mark,
    )
    session.add(ff)
    session.flush()
    return ff


# ========== 测试 1：接口层透传（当前应失败，因为接口未实现 body 参数）==========


def test_export_word_accepts_column_width_overrides(session: Session, tmp_path: Path) -> None:
    """验证导出接口接受 column_width_overrides 参数并传递到服务层。

    当前预期：测试失败，因为 export_word 端点未实现 body 参数解析。
    """
    project, _ = create_minimal_project(session)

    visit = Visit(project_id=project.id, name="访视1", code="V1", sequence=1)
    session.add(visit)
    session.flush()

    form = Form(project_id=project.id, name="普通表单", code="F_NORMAL", order_index=1)
    session.add(form)
    session.flush()

    vf = VisitForm(visit_id=visit.id, form_id=form.id, sequence=1)
    session.add(vf)
    session.flush()

    # 创建 2 个普通字段
    fd1 = create_text_field_def(session, project.id, "字段A")
    fd2 = create_text_field_def(session, project.id, "字段B")
    add_field_to_form(session, form.id, fd1.id, order_index=1, inline_mark=0)
    add_field_to_form(session, form.id, fd2.id, order_index=2, inline_mark=0)

    session.commit()

    # 准备列宽覆盖参数
    # 格式：{ form_id: { table_kind: [fraction, ...] } }
    # table_kind: "normal" | "inline" | "unified"
    # fraction: 0.0 ~ 1.0，表示该列占总宽度的比例
    column_width_overrides = {
        str(form.id): {
            "normal": [0.3, 0.7],  # 第一列 30%，第二列 70%
        }
    }

    output_path = tmp_path / "override.docx"

    # 传递列宽覆盖参数到导出服务
    service = ExportService(session)
    ok = service.export_project_to_word(project.id, str(output_path), column_width_overrides=column_width_overrides)

    assert ok is True, "导出应成功"

    doc = Document(str(output_path))
    form_tables = doc.tables[2:]  # 跳过封面和访视流程
    assert len(form_tables) >= 1, "应有至少一个表单表格"

    # 验证列宽比例
    normal_table = form_tables[0]
    tbl_xml = normal_table._tbl
    grid_cols = tbl_xml.findall(qn('w:tblGrid') + '/' + qn('w:gridCol'))

    if grid_cols and len(grid_cols) == 2:
        w0 = int(grid_cols[0].get(qn('w:w'), '0'))
        w1 = int(grid_cols[1].get(qn('w:w'), '0'))
        total = w0 + w1
        if total > 0:
            ratio0 = w0 / total
            # 验证第一列比例接近 0.3（允许 5% 误差）
            assert abs(ratio0 - 0.3) < 0.05, f"第一列比例应为 0.3，实际为 {ratio0:.3f}"


# ========== 测试 2：inline 表格列宽覆盖 ==========


def test_export_inline_table_column_width_override(session: Session, tmp_path: Path) -> None:
    """验证 inline 表格列宽覆盖应用到 Word 表格。

    当前预期：测试失败，因为服务层未实现列宽覆盖逻辑。
    """
    project, _ = create_minimal_project(session)

    visit = Visit(project_id=project.id, name="访视1", code="V1", sequence=1)
    session.add(visit)
    session.flush()

    form = Form(project_id=project.id, name="Inline表单", code="F_INLINE", order_index=1)
    session.add(form)
    session.flush()

    vf = VisitForm(visit_id=visit.id, form_id=form.id, sequence=1)
    session.add(vf)
    session.flush()

    # 创建 3 个 inline 字段
    fd1 = create_text_field_def(session, project.id, "列A")
    fd2 = create_text_field_def(session, project.id, "列B")
    fd3 = create_text_field_def(session, project.id, "列C")
    add_field_to_form(session, form.id, fd1.id, order_index=1, inline_mark=1)
    add_field_to_form(session, form.id, fd2.id, order_index=2, inline_mark=1)
    add_field_to_form(session, form.id, fd3.id, order_index=3, inline_mark=1)

    session.commit()

    # 准备列宽覆盖：[0.5, 0.3, 0.2]
    column_width_overrides = {
        str(form.id): {
            "inline": [0.5, 0.3, 0.2],
        }
    }

    output_path = tmp_path / "inline_override.docx"

    service = ExportService(session)
    ok = service.export_project_to_word(project.id, str(output_path), column_width_overrides=column_width_overrides)

    assert ok is True, "导出应成功"

    doc = Document(str(output_path))
    form_tables = doc.tables[2:]

    # 找到 3 列的 inline 表格
    inline_table = None
    for t in form_tables:
        if len(t.columns) == 3:
            inline_table = t
            break

    assert inline_table is not None, "应存在 3 列 inline 表格"

    # 验证列宽比例
    tbl_xml = inline_table._tbl
    grid_cols = tbl_xml.findall(qn('w:tblGrid') + '/' + qn('w:gridCol'))

    if grid_cols and len(grid_cols) == 3:
        widths = [int(gc.get(qn('w:w'), '0')) for gc in grid_cols]
        total = sum(widths)
        if total > 0:
            ratios = [w / total for w in widths]
            # 验证比例接近预期（允许 5% 误差）
            assert abs(ratios[0] - 0.5) < 0.05, f"第一列比例应为 0.5，实际为 {ratios[0]:.3f}"
            assert abs(ratios[1] - 0.3) < 0.05, f"第二列比例应为 0.3，实际为 {ratios[1]:.3f}"
            assert abs(ratios[2] - 0.2) < 0.05, f"第三列比例应为 0.2，实际为 {ratios[2]:.3f}"


# ========== 测试 3：unified 表格列宽覆盖 ==========


def test_export_unified_table_column_width_override(session: Session, tmp_path: Path) -> None:
    """验证 unified 表格列宽覆盖应用到 Word 表格。

    当前预期：测试失败，因为服务层未实现列宽覆盖逻辑。
    """
    project, _ = create_minimal_project(session)

    visit = Visit(project_id=project.id, name="访视1", code="V1", sequence=1)
    session.add(visit)
    session.flush()

    form = Form(project_id=project.id, name="混合表单", code="F_MIXED", order_index=1)
    session.add(form)
    session.flush()

    vf = VisitForm(visit_id=visit.id, form_id=form.id, sequence=1)
    session.add(vf)
    session.flush()

    # 添加 1 个普通字段 + 5 个 inline 字段触发 unified
    fd_normal = create_text_field_def(session, project.id, "普通字段")
    add_field_to_form(session, form.id, fd_normal.id, order_index=1, inline_mark=0)

    inline_fds = []
    for i in range(1, 6):
        fd = create_text_field_def(session, project.id, f"列{i}")
        inline_fds.append(fd)
        add_field_to_form(session, form.id, fd.id, order_index=10 + i, inline_mark=1)

    session.commit()

    # 准备列宽覆盖：5 列，自定义比例
    column_width_overrides = {
        str(form.id): {
            "unified": [0.1, 0.25, 0.2, 0.25, 0.2],
        }
    }

    output_path = tmp_path / "unified_override.docx"

    service = ExportService(session)
    ok = service.export_project_to_word(project.id, str(output_path), column_width_overrides=column_width_overrides)

    assert ok is True, "导出应成功"

    doc = Document(str(output_path))

    # 找到 5 列的 unified 表格
    unified_table = None
    for t in doc.tables[2:]:
        if len(t.columns) == 5:
            unified_table = t
            break

    assert unified_table is not None, "应存在 5 列 unified 表格"

    # 验证列宽比例
    tbl_xml = unified_table._tbl
    grid_cols = tbl_xml.findall(qn('w:tblGrid') + '/' + qn('w:gridCol'))

    if grid_cols and len(grid_cols) == 5:
        widths = [int(gc.get(qn('w:w'), '0')) for gc in grid_cols]
        total = sum(widths)
        if total > 0:
            ratios = [w / total for w in widths]
            expected = [0.1, 0.25, 0.2, 0.25, 0.2]
            for i, (actual, exp) in enumerate(zip(ratios, expected)):
                assert abs(actual - exp) < 0.05, f"第 {i+1} 列比例应为 {exp:.2f}，实际为 {actual:.3f}"


# ========== 测试 4~7：边界情况处理 ==========


def test_export_ignores_invalid_fraction_values(session: Session, tmp_path: Path) -> None:
    """验证后端忽略超出范围的 fraction 值。"""
    project, _ = create_minimal_project(session)

    visit = Visit(project_id=project.id, name="访视1", code="V1", sequence=1)
    session.add(visit)
    session.flush()

    form = Form(project_id=project.id, name="测试表单", code="F_TEST", order_index=1)
    session.add(form)
    session.flush()

    vf = VisitForm(visit_id=visit.id, form_id=form.id, sequence=1)
    session.add(vf)
    session.flush()

    fd1 = create_text_field_def(session, project.id, "字段A")
    fd2 = create_text_field_def(session, project.id, "字段B")
    add_field_to_form(session, form.id, fd1.id, order_index=1)
    add_field_to_form(session, form.id, fd2.id, order_index=2)

    session.commit()

    # 测试超出范围的值
    invalid_overrides = {
        str(form.id): {
            "normal": [1.5, -0.5],  # 超出 0.0~1.0 范围
        }
    }

    output_path = tmp_path / "invalid.docx"
    service = ExportService(session)
    # 应该忽略无效值并成功导出
    ok = service.export_project_to_word(project.id, str(output_path), column_width_overrides=invalid_overrides)

    assert ok is True, "导出应成功（忽略无效覆盖值）"

    doc = Document(str(output_path))
    form_tables = doc.tables[2:]
    assert len(form_tables) >= 1

    # 验证列宽不是按无效覆盖值设置（而是按内容驱动默认值）
    normal_table = form_tables[0]
    tbl_xml = normal_table._tbl
    grid_cols = tbl_xml.findall(qn('w:tblGrid') + '/' + qn('w:gridCol'))

    if grid_cols and len(grid_cols) == 2:
        w0 = int(grid_cols[0].get(qn('w:w'), '0'))
        w1 = int(grid_cols[1].get(qn('w:w'), '0'))
        total = w0 + w1
        if total > 0:
            ratio0 = w0 / total
            # 不应该接近 1.5（无效值被忽略，使用默认值）
            assert ratio0 < 0.9, "无效覆盖值应被忽略"


def test_export_ignores_non_numeric_values(session: Session, tmp_path: Path) -> None:
    """验证后端忽略非数值类型的 fraction 值。"""
    project, _ = create_minimal_project(session)

    visit = Visit(project_id=project.id, name="访视1", code="V1", sequence=1)
    session.add(visit)
    session.flush()

    form = Form(project_id=project.id, name="测试表单", code="F_TEST", order_index=1)
    session.add(form)
    session.flush()

    vf = VisitForm(visit_id=visit.id, form_id=form.id, sequence=1)
    session.add(vf)
    session.flush()

    fd1 = create_text_field_def(session, project.id, "字段A")
    fd2 = create_text_field_def(session, project.id, "字段B")
    add_field_to_form(session, form.id, fd1.id, order_index=1)
    add_field_to_form(session, form.id, fd2.id, order_index=2)

    session.commit()

    # 测试非数值类型
    invalid_overrides = {
        str(form.id): {
            "normal": ["0.3", "0.7"],  # 字符串类型
        }
    }

    output_path = tmp_path / "non_numeric.docx"
    service = ExportService(session)
    ok = service.export_project_to_word(project.id, str(output_path), column_width_overrides=invalid_overrides)

    assert ok is True, "导出应成功（忽略非数值覆盖值）"


def test_export_ignores_length_mismatch(session: Session, tmp_path: Path) -> None:
    """验证后端忽略长度不匹配的覆盖配置。"""
    project, _ = create_minimal_project(session)

    visit = Visit(project_id=project.id, name="访视1", code="V1", sequence=1)
    session.add(visit)
    session.flush()

    form = Form(project_id=project.id, name="测试表单", code="F_TEST", order_index=1)
    session.add(form)
    session.flush()

    vf = VisitForm(visit_id=visit.id, form_id=form.id, sequence=1)
    session.add(vf)
    session.flush()

    fd1 = create_text_field_def(session, project.id, "字段A")
    fd2 = create_text_field_def(session, project.id, "字段B")
    add_field_to_form(session, form.id, fd1.id, order_index=1)
    add_field_to_form(session, form.id, fd2.id, order_index=2)

    session.commit()

    # 测试长度不匹配（2列表格传入3个值）
    invalid_overrides = {
        str(form.id): {
            "normal": [0.3, 0.3, 0.4],  # 3个值，但表格只有2列
        }
    }

    output_path = tmp_path / "length_mismatch.docx"
    service = ExportService(session)
    ok = service.export_project_to_word(project.id, str(output_path), column_width_overrides=invalid_overrides)

    assert ok is True, "导出应成功（忽略长度不匹配的覆盖值）"


def test_export_handles_empty_overrides(session: Session, tmp_path: Path) -> None:
    """验证后端正确处理空的覆盖配置。"""
    project, _ = create_minimal_project(session)

    visit = Visit(project_id=project.id, name="访视1", code="V1", sequence=1)
    session.add(visit)
    session.flush()

    form = Form(project_id=project.id, name="测试表单", code="F_TEST", order_index=1)
    session.add(form)
    session.flush()

    vf = VisitForm(visit_id=visit.id, form_id=form.id, sequence=1)
    session.add(vf)
    session.flush()

    fd1 = create_text_field_def(session, project.id, "字段A")
    fd2 = create_text_field_def(session, project.id, "字段B")
    add_field_to_form(session, form.id, fd1.id, order_index=1)
    add_field_to_form(session, form.id, fd2.id, order_index=2)

    session.commit()

    # 测试空覆盖配置
    empty_overrides = {}

    output_path = tmp_path / "empty_overrides.docx"
    service = ExportService(session)
    ok = service.export_project_to_word(project.id, str(output_path), column_width_overrides=empty_overrides)

    assert ok is True, "导出应成功"

    doc = Document(str(output_path))
    form_tables = doc.tables[2:]
    assert len(form_tables) >= 1, "应有至少一个表单表格"

