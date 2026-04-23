import sqlite3
from pathlib import Path
from unittest.mock import patch

from docx import Document
import pytest
from fastapi.testclient import TestClient
from sqlalchemy import create_engine, event
from sqlalchemy.orm import Session, sessionmaker
from sqlalchemy.pool import StaticPool

from main import app
from src.config import AdminConfig, AppConfig, AuthConfig
from src.database import get_read_session, get_session
from src.models import Base
from src.models.project import Project
from src.models.user import User
from src.services.export_service import (
    ExportService,
    export_full_database,
    export_project_database,
    export_user_projects_database,
)
from tests.helpers import auth_headers, login_as


_TEST_CONFIG = AppConfig(
    auth=AuthConfig(secret_key="test-secret-key-for-testing"),
    admin=AdminConfig(username="admin"),
)


@pytest.fixture
def session() -> Session:
    engine = create_engine("sqlite+pysqlite:///:memory:")
    Base.metadata.create_all(engine)
    session_factory = sessionmaker(bind=engine, expire_on_commit=False)

    with session_factory() as db_session:
        yield db_session

    engine.dispose()


@pytest.fixture
def engine():
    _engine = create_engine(
        "sqlite+pysqlite:///:memory:",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )

    @event.listens_for(_engine, "connect")
    def _enable_fk(dbapi_conn, _):
        dbapi_conn.execute("PRAGMA foreign_keys = ON")

    Base.metadata.create_all(_engine)
    yield _engine
    _engine.dispose()


@pytest.fixture
def client(engine) -> TestClient:
    def _override():
        with Session(engine) as db_session:
            with db_session.begin():
                yield db_session

    app.dependency_overrides[get_session] = _override
    app.dependency_overrides[get_read_session] = _override
    with patch("main.get_config", return_value=_TEST_CONFIG), \
         patch("src.services.auth_service.get_config", return_value=_TEST_CONFIG), \
         patch("src.services.user_admin_service.get_config", return_value=_TEST_CONFIG):
        with TestClient(app, raise_server_exceptions=False) as test_client:
            yield test_client
    app.dependency_overrides.clear()


def create_project(session: Session, name: str = "项目") -> Project:
    project = Project(name=name, version="v1.0")
    session.add(project)
    session.flush()
    return project


# ── 验证相关 ──


def test_validate_output_accepts_valid_docx(tmp_path: Path) -> None:
    output_path = tmp_path / "valid.docx"
    doc = Document()
    doc.add_table(rows=2, cols=3)  # 封面
    doc.add_table(rows=1, cols=1)  # 访视流程
    doc.add_table(rows=1, cols=2)  # 表单
    doc.save(output_path)

    ok, reason = ExportService._validate_output(str(output_path))

    assert ok is True
    assert reason == ""


def test_validate_output_rejects_insufficient_tables(tmp_path: Path) -> None:
    output_path = tmp_path / "insufficient.docx"
    doc = Document()
    doc.add_table(rows=2, cols=3)  # 封面
    doc.add_table(rows=1, cols=1)  # 访视流程
    doc.save(output_path)

    ok, reason = ExportService._validate_output(str(output_path))

    assert ok is False
    assert "結構不完整" in reason


def test_validate_output_rejects_zero_byte_file(tmp_path: Path) -> None:
    output_path = tmp_path / "empty.docx"
    output_path.write_bytes(b"")

    ok, reason = ExportService._validate_output(str(output_path))

    assert ok is False
    assert "0 字节" in reason


# ── POST /export/word 直接下载 ──


def test_export_word_returns_docx_file(
    client: TestClient,
    engine,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    with Session(engine) as session:
        user = User(username="alice")
        session.add(user)
        session.flush()
        project = Project(name="导出项目", version="v1.0", owner_id=user.id)
        session.add(project)
        session.commit()
        project_id = project.id

    token = login_as(client, "alice")
    monkeypatch.setattr(
        ExportService,
        "export_project_to_word",
        lambda self, pid, output_path, column_width_overrides=None: Document().save(output_path) or True,
    )
    monkeypatch.setattr(
        ExportService,
        "_validate_output",
        staticmethod(lambda path: (True, "")),
    )

    response = client.post(
        f"/api/projects/{project_id}/export/word",
        headers=auth_headers(token),
    )

    assert response.status_code == 200
    assert response.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


def test_export_word_rejects_invalid_docx(
    client: TestClient,
    engine,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    with Session(engine) as session:
        user = User(username="alice")
        session.add(user)
        session.flush()
        project = Project(name="导出项目", version="v1.0", owner_id=user.id)
        session.add(project)
        session.commit()
        project_id = project.id

    token = login_as(client, "alice")
    monkeypatch.setattr(
        ExportService,
        "export_project_to_word",
        lambda self, pid, output_path, column_width_overrides=None: Path(output_path).write_bytes(b"PK") or True,
    )
    monkeypatch.setattr(
        ExportService,
        "_validate_output",
        staticmethod(lambda path: (False, "无效 docx")),
    )

    response = client.post(
        f"/api/projects/{project_id}/export/word",
        headers=auth_headers(token),
    )

    assert response.status_code == 500
    assert "无效 docx" in response.text


# ── 数据库导出 ──


def test_export_full_database_returns_valid_sqlite(tmp_path: Path) -> None:
    """整库导出返回可用的 .db 文件"""
    src_path = str(tmp_path / "source.db")
    conn = sqlite3.connect(src_path)
    conn.execute("CREATE TABLE test (id INTEGER PRIMARY KEY, name TEXT)")
    conn.execute("INSERT INTO test VALUES (1, 'hello')")
    conn.commit()
    conn.close()

    result_path = export_full_database(src_path)
    try:
        result_conn = sqlite3.connect(result_path)
        rows = result_conn.execute("SELECT * FROM test").fetchall()
        result_conn.close()
        assert rows == [(1, "hello")]
    finally:
        Path(result_path).unlink(missing_ok=True)


def test_export_project_database_prunes_correctly(tmp_path: Path) -> None:
    """单项目导出裁剪完整性：仅含目标项目、user 表为空、owner_id 为 NULL"""
    src_path = str(tmp_path / "source.db")
    conn = sqlite3.connect(src_path)
    conn.execute("PRAGMA foreign_keys = ON")
    # 模拟最小 schema
    conn.execute("CREATE TABLE user (id INTEGER PRIMARY KEY, username TEXT)")
    conn.execute("CREATE TABLE project (id INTEGER PRIMARY KEY, name TEXT, owner_id INTEGER REFERENCES user(id))")
    conn.execute("INSERT INTO user VALUES (1, 'alice')")
    conn.execute("INSERT INTO project VALUES (1, 'KeepMe', 1)")
    conn.execute("INSERT INTO project VALUES (2, 'RemoveMe', 1)")
    conn.commit()
    conn.close()

    result_path = export_project_database(src_path, 1, "KeepMe")
    try:
        result_conn = sqlite3.connect(result_path)
        projects = result_conn.execute("SELECT id, name, owner_id FROM project").fetchall()
        users = result_conn.execute("SELECT * FROM user").fetchall()
        result_conn.close()

        assert len(projects) == 1
        assert projects[0][0] == 1
        assert projects[0][1] == "KeepMe"
        assert projects[0][2] is None  # owner_id 已置 NULL
        assert users == []  # user 表为空
    finally:
        Path(result_path).unlink(missing_ok=True)



def test_export_user_projects_database_prunes_to_owner_scope(tmp_path: Path) -> None:
    """用户聚合导出仅保留当前用户项目，并清空 user 表与 owner_id。"""
    src_path = str(tmp_path / "user_scope.db")
    conn = sqlite3.connect(src_path)
    conn.execute("PRAGMA foreign_keys = ON")
    conn.execute("CREATE TABLE user (id INTEGER PRIMARY KEY, username TEXT)")
    conn.execute("CREATE TABLE project (id INTEGER PRIMARY KEY, name TEXT, owner_id INTEGER REFERENCES user(id))")
    conn.executemany(
        "INSERT INTO user VALUES (?, ?)",
        [(1, 'alice'), (2, 'bob')],
    )
    conn.executemany(
        "INSERT INTO project VALUES (?, ?, ?)",
        [
            (1, 'Alice-A', 1),
            (2, 'Alice-B', 1),
            (3, 'Bob-A', 2),
        ],
    )
    conn.commit()
    conn.close()

    result_path = export_user_projects_database(src_path, 1, 'alice')
    try:
        result_conn = sqlite3.connect(result_path)
        projects = result_conn.execute("SELECT id, name, owner_id FROM project ORDER BY id").fetchall()
        users = result_conn.execute("SELECT * FROM user").fetchall()
        result_conn.close()

        assert projects == [
            (1, 'Alice-A', None),
            (2, 'Alice-B', None),
        ]
        assert users == []
    finally:
        Path(result_path).unlink(missing_ok=True)
