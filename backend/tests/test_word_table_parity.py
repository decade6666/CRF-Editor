from __future__ import annotations

import json
from pathlib import Path

from docx import Document

from src.services.word_table_parity import (
    compare_table_field_forms,
    extract_docx_form_table_fields,
    load_preview_table_fields,
)


def _write_docx_with_forms(path: Path) -> None:
    doc = Document()
    doc.add_table(rows=1, cols=1).cell(0, 0).text = "cover"
    doc.add_heading("表单访视分布图", level=1)
    doc.add_table(rows=1, cols=1).cell(0, 0).text = "visit flow"

    doc.add_heading("1. 生命体征", level=1)
    first_table = doc.add_table(rows=2, cols=2)
    first_table.cell(0, 0).text = "收缩压"
    first_table.cell(0, 1).text = "|__||__|"
    first_table.cell(1, 0).text = "是否吸烟"
    first_table.cell(1, 1).text = "○是  ○否"

    doc.add_heading("2. 实验室", level=1)
    second_table = doc.add_table(rows=1, cols=2)
    second_table.cell(0, 0).text = "其他"
    second_table.cell(0, 1).text = "其他，请描述______"

    doc.save(path)


def test_extract_docx_form_table_fields_ignores_scaffolding_tables(tmp_path: Path) -> None:
    docx_path = tmp_path / "forms.docx"
    _write_docx_with_forms(docx_path)

    forms = extract_docx_form_table_fields(docx_path)

    assert [form.name for form in forms] == ["生命体征", "实验室"]
    assert forms[0].tables == [[
        ["收缩压", "|__||__|"],
        ["是否吸烟", "○是  ○否"],
    ]]
    assert forms[1].tables == [[["其他", "其他，请描述______"]]]


def test_extract_docx_form_table_fields_collapses_merged_cells(tmp_path: Path) -> None:
    docx_path = tmp_path / "merged.docx"
    doc = Document()
    doc.add_heading("1. 合并行", level=1)
    table = doc.add_table(rows=1, cols=2)
    merged = table.cell(0, 0).merge(table.cell(0, 1))
    merged.text = "以下为log行"
    doc.save(docx_path)

    forms = extract_docx_form_table_fields(docx_path)

    assert forms[0].tables == [[['以下为log行']]]



def test_compare_table_field_forms_reports_exact_counts_and_mismatches(tmp_path: Path) -> None:
    docx_path = tmp_path / "forms.docx"
    preview_path = tmp_path / "preview.json"
    _write_docx_with_forms(docx_path)
    preview_path.write_text(
        json.dumps(
            {
                "forms": [
                    {
                        "name": "生命体征",
                        "tables": [[
                            ["收缩压", "|__||__|"],
                            ["是否吸烟", "○是 ○否"],
                        ]],
                    },
                    {
                        "name": "实验室",
                        "tables": [[["其他", "其他，请描述______"]]],
                    },
                ]
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    preview_forms = load_preview_table_fields(preview_path)
    export_forms = extract_docx_form_table_fields(docx_path)
    report = compare_table_field_forms(preview_forms, export_forms)

    assert report.preview_form_count == 2
    assert report.export_form_count == 2
    assert report.form_order_matches is True
    assert report.preview_row_count == 3
    assert report.export_row_count == 3
    assert report.preview_cell_count == 6
    assert report.export_cell_count == 6
    assert report.exact_cell_count == 5
    assert report.exact_cell_ratio == 5 / 6
    assert report.exact_row_count == 2
    assert report.exact_row_ratio == 2 / 3
    assert report.mismatches[0].to_dict() == {
        "form_index": 0,
        "form_name": "生命体征",
        "table_index": 0,
        "row_index": 1,
        "cell_index": 1,
        "kind": "text",
        "preview": "○是 ○否",
        "export": "○是  ○否",
    }


def test_compare_table_field_forms_counts_extra_cells_in_denominator() -> None:
    preview_forms = load_preview_table_fields([
        {"name": "A", "tables": [[['x']]]},
    ])
    export_forms = load_preview_table_fields([
        {"name": "A", "tables": [[['x', 'extra']]]},
    ])

    report = compare_table_field_forms(preview_forms, export_forms)

    assert report.preview_cell_count == 1
    assert report.export_cell_count == 2
    assert report.exact_cell_count == 1
    assert report.exact_cell_denominator == 2
    assert report.exact_cell_ratio == 0.5
    assert report.mismatches[0].kind == "missing_preview_cell"
