from __future__ import annotations

from pathlib import Path
import logging
import math
import os
import re
import tempfile
from unittest.mock import patch

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.oxml.ns import qn
from docx.shared import Cm
import pytest
from hypothesis import given, settings
from hypothesis import strategies as st
from sqlalchemy import create_engine
from sqlalchemy.orm import Session, sessionmaker

from src.models import Base
from src.models.field_definition import FieldDefinition
from src.models.form import Form
from src.models.form_field import FormField
from src.models.project import Project
from src.models.visit import Visit
from src.models.visit_form import VisitForm
from src.services import toc_pagination
from src.services.export_service import ExportService, LayoutDecision, Segment


@pytest.fixture
def session() -> Session:
    engine = create_engine("sqlite+pysqlite:///:memory:")
    Base.metadata.create_all(engine)
    session_factory = sessionmaker(bind=engine, expire_on_commit=False)

    with session_factory() as db_session:
        yield db_session

    engine.dispose()



def create_project(session: Session, name: str = "项目") -> Project:
    project = Project(name=name, version="v1.0")
    session.add(project)
    session.flush()
    return project



def create_form(
    session: Session,
    project_id: int,
    *,
    name: str,
    order_index: int | None = None,
) -> Form:
    form = Form(
        project_id=project_id,
        name=name,
        code=f"{name}_CODE",
        order_index=order_index,
    )
    session.add(form)
    session.flush()
    return form



def create_visit(
    session: Session,
    project_id: int,
    *,
    name: str,
    sequence: int,
) -> Visit:
    visit = Visit(
        project_id=project_id,
        name=name,
        code=f"{name}_CODE",
        sequence=sequence,
    )
    session.add(visit)
    session.flush()
    return visit



def create_visit_form(
    session: Session,
    visit_id: int,
    form_id: int,
    *,
    sequence: int,
) -> VisitForm:
    visit_form = VisitForm(visit_id=visit_id, form_id=form_id, sequence=sequence)
    session.add(visit_form)
    session.flush()
    return visit_form



def create_field_definition(
    session: Session,
    project_id: int,
    *,
    variable_name: str,
    label: str,
    field_type: str = "文本",
) -> FieldDefinition:
    field_definition = FieldDefinition(
        project_id=project_id,
        variable_name=variable_name,
        label=label,
        field_type=field_type,
    )
    session.add(field_definition)
    session.flush()
    return field_definition



def create_form_field(
    session: Session,
    form_id: int,
    field_definition_id: int,
    *,
    order_index: int,
    default_value: str | None = None,
) -> FormField:
    form_field = FormField(
        form_id=form_id,
        field_definition_id=field_definition_id,
        order_index=order_index,
        default_value=default_value,
    )
    session.add(form_field)
    session.flush()
    return form_field



def export_document(session: Session, project_id: int, tmp_path: Path) -> Document:
    output_path = tmp_path / f"project-{project_id}.docx"
    # 关闭 LibreOffice 页码预计算：保持单测快速、确定（不依赖外部渲染进程）
    ok = ExportService(session).export_project_to_word(
        project_id, str(output_path), bake_toc_page_numbers=False
    )
    assert ok is True
    return Document(output_path)



def assert_table_rows_at_least_one_centimeter(table) -> None:
    for row in table.rows:
        tr_pr = row._tr.trPr
        assert tr_pr is not None
        tr_height = tr_pr.find(qn('w:trHeight'))
        assert tr_height is not None
        assert tr_height.get(qn('w:hRule')) == 'atLeast'
        assert tr_height.get(qn('w:val')) == str(Cm(1).twips)



def extract_form_headings(doc: Document) -> list[str]:
    headings: list[str] = []
    for paragraph in doc.paragraphs:
        # 跳过含域代码的段落（如 TOC 预渲染条目），避免误判为表单标题
        if any(True for _ in paragraph._p.iter(qn("w:fldChar"))):
            continue
        text = paragraph.text.strip()
        if re.match(r"^\d+\.\s+", text):
            headings.append(text)
    return headings


def _find_toc_paragraph_tokens(doc: Document) -> list[tuple[str, str]]:
    for paragraph in doc.paragraphs:
        tokens: list[tuple[str, str]] = []
        for element in paragraph._p.iter():
            if element.tag == qn('w:fldChar'):
                tokens.append(("fldChar", element.get(qn('w:fldCharType')) or ""))
            elif element.tag == qn('w:instrText'):
                tokens.append(("instrText", element.text or ""))
            elif element.tag == qn('w:t') and element.text:
                tokens.append(("text", element.text))
        if any(kind == "instrText" and "TOC" in value for kind, value in tokens):
            return tokens
    return []


def _first_token_index(tokens: list[tuple[str, str]], kind: str, value: str) -> int:
    for index, token in enumerate(tokens):
        if token == (kind, value):
            return index
    raise AssertionError(f"missing token: {(kind, value)}")



def test_export_project_renders_one_table_per_form_for_standard_fields(
    session: Session,
    tmp_path: Path,
) -> None:
    project = create_project(session)
    form = create_form(session, project.id, name="生命体征", order_index=1)
    visit = create_visit(session, project.id, name="筛选期", sequence=1)
    create_visit_form(session, visit.id, form.id, sequence=7)

    systolic = create_field_definition(
        session,
        project.id,
        variable_name="SYSBP",
        label="收缩压",
    )
    diastolic = create_field_definition(
        session,
        project.id,
        variable_name="DIABP",
        label="舒张压",
    )
    create_form_field(session, form.id, systolic.id, order_index=1, default_value="120")
    create_form_field(session, form.id, diastolic.id, order_index=2, default_value="80")

    doc = export_document(session, project.id, tmp_path)

    assert len(doc.tables) == 3
    form_table = doc.tables[2]
    assert len(form_table.rows) == 2
    assert len(form_table.columns) == 2
    assert form_table.cell(0, 0).text.strip() == "收缩压"
    assert form_table.cell(1, 0).text.strip() == "舒张压"



def test_export_text_field_fill_line_scales_with_column_width(
    session: Session,
    tmp_path: Path,
) -> None:
    from src.services.width_planning import (
        compute_fill_line_char_count,
        plan_normal_table_width,
    )

    project = create_project(session)
    form = create_form(session, project.id, name="备注表", order_index=1)
    visit = create_visit(session, project.id, name="筛选期", sequence=1)
    create_visit_form(session, visit.id, form.id, sequence=1)
    note = create_field_definition(
        session,
        project.id,
        variable_name="NOTE",
        label="备注",
        field_type="文本",
    )
    note_field = create_form_field(session, form.id, note.id, order_index=1)

    doc = export_document(session, project.id, tmp_path)
    form_table = doc.tables[2]
    fill_cell_text = form_table.cell(0, 1).text

    widths = plan_normal_table_width([note_field], available_cm=14.66)
    expected = compute_fill_line_char_count(widths[1])

    # 填写线根数随 control 列宽自适应，比旧固定 16 更长且不换行
    assert fill_cell_text == "_" * expected
    assert expected > 16
    assert "\n" not in fill_cell_text



def test_export_choice_trailing_fill_line_scales_with_column_width(
    session: Session,
    tmp_path: Path,
) -> None:
    from src.models.codelist import CodeList, CodeListOption
    from src.services.width_planning import (
        CELL_HPAD_CM,
        FILL_LINE_SAFETY_CM,
        UNDERSCORE_CHAR_CM,
        compute_choice_atom_weight,
        compute_choice_trailing_fill_char_count,
        compute_horizontal_choice_trailing_fill_chars,
        compute_fill_line_char_count,
        plan_normal_table_width,
    )

    project = create_project(session)
    form = create_form(session, project.id, name="选择表", order_index=1)
    visit = create_visit(session, project.id, name="筛选期", sequence=1)
    create_visit_form(session, visit.id, form.id, sequence=1)

    codelist = CodeList(project_id=project.id, name="诊断结果", code="CL_DIAG")
    session.add(codelist)
    session.flush()
    session.add_all([
        CodeListOption(codelist_id=codelist.id, code="1", decode="有尾线", trailing_underscore=1, order_index=1),
        CodeListOption(codelist_id=codelist.id, code="2", decode="无尾线", trailing_underscore=0, order_index=2),
    ])
    session.flush()

    choice = create_field_definition(
        session,
        project.id,
        variable_name="DIAG",
        label="诊断",
        field_type="单选",
    )
    choice.codelist_id = codelist.id
    session.flush()
    choice_field = create_form_field(session, form.id, choice.id, order_index=1)

    doc = export_document(session, project.id, tmp_path)
    fill_cell_text = doc.tables[2].cell(0, 1).text

    widths = plan_normal_table_width([choice_field], available_cm=14.66)
    full_line_count = compute_fill_line_char_count(widths[1])
    # 横向单选：尾线按扣除所有选项 marker+label+分隔符后的剩余宽度计算
    expected = compute_horizontal_choice_trailing_fill_chars(
        widths[1], [("有尾线", True), ("无尾线", False)]
    )

    usable_cm = widths[1] - CELL_HPAD_CM - FILL_LINE_SAFETY_CM
    # 整行（两个选项 marker+label + 分隔符 + 尾线）估算宽度不超过列宽预算 → 不换行
    line_chars = (
        math.ceil(compute_choice_atom_weight("有尾线", False))
        + math.ceil(compute_choice_atom_weight("无尾线", False))
        + 2
        + expected
    )

    assert fill_cell_text == f"○有尾线{'_' * expected}  ○无尾线"
    assert 6 < expected < full_line_count
    assert line_chars * UNDERSCORE_CHAR_CM <= usable_cm
    assert "\n" not in fill_cell_text


def test_render_field_control_defaults_to_legacy_sixteen_underscores(
    session: Session,
) -> None:
    project = create_project(session)
    text_field = create_field_definition(
        session,
        project.id,
        variable_name="MEMO",
        label="说明",
        field_type="文本",
    )

    # 未传 fill_line_chars 的调用方（inline / unified / 空占位）保持旧行为
    rendered = ExportService(session)._render_field_control(text_field)

    assert rendered == "________________"


def test_export_project_sorts_form_headings_by_order_index_then_id(
    session: Session,
    tmp_path: Path,
) -> None:
    project = create_project(session)
    create_form(session, project.id, name="Alpha", order_index=2)
    create_form(session, project.id, name="Zeta", order_index=1)
    create_form(session, project.id, name="LaterById", order_index=None)
    create_form(session, project.id, name="LastById", order_index=None)

    doc = export_document(session, project.id, tmp_path)

    assert extract_form_headings(doc) == [
        "1. Zeta",
        "2. Alpha",
        "3. LaterById",
        "4. LastById",
    ]



def test_export_sets_update_fields_on_open(
    session: Session,
    tmp_path: Path,
) -> None:
    project = create_project(session)
    create_form(session, project.id, name="生命体征", order_index=1)

    doc = export_document(session, project.id, tmp_path)

    settings = doc.settings.element
    update_fields = settings.find(qn('w:updateFields'))
    assert update_fields is not None
    assert update_fields.get(qn('w:val')) == 'true'
    # CT_Settings 顺序：updateFields 必须排在 compat 之前
    children = list(settings)
    compat = settings.find(qn('w:compat'))
    if compat is not None:
        assert children.index(update_fields) < children.index(compat)


def test_export_toc_field_is_well_formed_and_keeps_heading_extraction_clean(
    session: Session,
    tmp_path: Path,
) -> None:
    project = create_project(session)
    create_form(session, project.id, name="生命体征", order_index=1)
    create_form(session, project.id, name="实验室", order_index=2)

    doc = export_document(session, project.id, tmp_path)

    tokens = _find_toc_paragraph_tokens(doc)
    assert tokens
    assert any(
        kind == "instrText"
        and 'TOC \\o "1-3" \\h \\z \\u' in value
        for kind, value in tokens
    )
    begin_index = _first_token_index(tokens, "fldChar", "begin")
    instr_index = next(
        index
        for index, (kind, value) in enumerate(tokens)
        if kind == "instrText" and "TOC" in value and "\\h" in value
    )
    separate_index = _first_token_index(tokens, "fldChar", "separate")

    # 首条条目以 TOC 域起始开头：begin → instrText(TOC) → separate；
    # 其后是该条目自身的超链接与 PAGEREF 域（含 end），故本段允许出现 end
    assert begin_index == 0
    assert begin_index < instr_index < separate_index
    assert extract_form_headings(doc) == ["1. 生命体征", "2. 实验室"]


def test_export_headings_have_unique_toc_bookmarks(
    session: Session,
    tmp_path: Path,
) -> None:
    """每个标题段落包含唯一 _Toc 书签，无重复 name。"""
    project = create_project(session)
    create_form(session, project.id, name="生命体征", order_index=1)
    create_form(session, project.id, name="实验室", order_index=2)

    doc = export_document(session, project.id, tmp_path)

    bookmark_names: list[str] = []
    for paragraph in doc.paragraphs:
        for bm_start in paragraph._p.findall(qn("w:bookmarkStart")):
            name = bm_start.get(qn("w:name"))
            if name and name.startswith("_Toc"):
                bookmark_names.append(name)

    assert len(bookmark_names) >= 3  # 访视分布图 + 2 个表单
    assert len(bookmark_names) == len(set(bookmark_names)), "书签 name 有重复"


def test_export_toc_prerendered_entries_match_headings(
    session: Session,
    tmp_path: Path,
) -> None:
    """预渲染 TOC 条目数 = 标题段落数（含访视分布图）。"""
    project = create_project(session)
    create_form(session, project.id, name="生命体征", order_index=1)
    create_form(session, project.id, name="实验室", order_index=2)

    doc = export_document(session, project.id, tmp_path)

    heading_count = 0
    for paragraph in doc.paragraphs:
        for bm_start in paragraph._p.findall(qn("w:bookmarkStart")):
            name = bm_start.get(qn("w:name"))
            if name and name.startswith("_Toc"):
                heading_count += 1
                break

    toc_hyperlink_count = 0
    for paragraph in doc.paragraphs:
        for hl in paragraph._p.findall(qn("w:hyperlink")):
            anchor = hl.get(qn("w:anchor"))
            if anchor and anchor.startswith("_Toc"):
                toc_hyperlink_count += 1

    assert heading_count >= 3
    assert toc_hyperlink_count == heading_count


def test_export_toc_entries_have_pageref(
    session: Session,
    tmp_path: Path,
) -> None:
    """预渲染条目的 PAGEREF 域引用的书签与 w:hyperlink@anchor 一致。"""
    project = create_project(session)
    create_form(session, project.id, name="生命体征", order_index=1)

    doc = export_document(session, project.id, tmp_path)

    entry_anchors: list[str] = []
    pageref_targets: list[str] = []
    for paragraph in doc.paragraphs:
        for hl in paragraph._p.findall(qn("w:hyperlink")):
            anchor = hl.get(qn("w:anchor"))
            if anchor and anchor.startswith("_Toc"):
                entry_anchors.append(anchor)
        for instr in paragraph._p.iter(qn("w:instrText")):
            text = (instr.text or "").strip()
            if text.startswith("PAGEREF"):
                target = text.split()[1] if len(text.split()) > 1 else ""
                pageref_targets.append(target)

    assert entry_anchors, "未找到预渲染条目"
    assert entry_anchors == pageref_targets


def test_export_toc_entries_use_song_font_defined_styles_and_fallback_page_no(
    session: Session,
    tmp_path: Path,
) -> None:
    """TOC 条目使用宋体、引用已定义的 TOC1 样式，页码有非空回退值。"""
    project = create_project(session)
    create_form(session, project.id, name="生命体征", order_index=1)

    doc = export_document(session, project.id, tmp_path)

    # 1) TOC1/2/3 段落样式已注入 styles.xml，pStyle 不再悬空
    style_ids = {
        s.get(qn("w:styleId"))
        for s in doc.styles.element.findall(qn("w:style"))
    }
    assert {"TOC1", "TOC2", "TOC3"} <= style_ids

    # 2) 定位首条预渲染条目
    entry = None
    for paragraph in doc.paragraphs:
        for hl in paragraph._p.findall(qn("w:hyperlink")):
            if (hl.get(qn("w:anchor")) or "").startswith("_Toc"):
                entry = paragraph
                break
        if entry is not None:
            break
    assert entry is not None, "未找到预渲染条目"

    # 3) 超链接文本 run 写入宋体（eastAsia=SimSun），与正文字体一致
    hyperlink = entry._p.find(qn("w:hyperlink"))
    rfonts = hyperlink.find(qn("w:r")).find(qn("w:rPr")).find(qn("w:rFonts"))
    assert rfonts is not None
    assert rfonts.get(qn("w:eastAsia")) == "SimSun"

    # 4) 页码有非空数字回退值，整段不含 '?'
    all_text = "".join(t.text or "" for t in entry._p.iter(qn("w:t")))
    numbers = _toc_entry_page_numbers(doc)
    assert "?" not in all_text
    assert numbers
    assert all(number.strip() for number in numbers)
    assert all(number.isdigit() for number in numbers)


def test_export_toc_field_end_wraps_prerendered_entries(
    session: Session,
    tmp_path: Path,
) -> None:
    """外层 TOC 域起始合入首条条目、end 合入末条条目，整段条目即域结果。

    首条条目以 TOC begin 开启（含 separate），末条条目以外层 TOC end 收尾。
    这样 Word 更新域时整体替换预渲染条目，不会在其上方再生成一份导致目录重复。
    """
    project = create_project(session)
    create_form(session, project.id, name="生命体征", order_index=1)
    create_form(session, project.id, name="实验室", order_index=2)

    doc = export_document(session, project.id, tmp_path)

    paragraphs = doc.paragraphs

    # 首条条目（含 TOC instrText）：以 TOC begin 开启、含 separate；
    # 该段同时含自身 PAGEREF（begin/separate/end），故 begin 出现两次
    toc_idx = next(
        i
        for i, para in enumerate(paragraphs)
        if any("TOC" in (instr.text or "") for instr in para._p.iter(qn("w:instrText")))
    )
    toc_fld_types = [
        fc.get(qn("w:fldCharType"))
        for fc in paragraphs[toc_idx]._p.iter(qn("w:fldChar"))
    ]
    assert toc_fld_types[0] == "begin"
    assert "separate" in toc_fld_types
    assert toc_fld_types.count("begin") == 2, "TOC begin + PAGEREF begin"

    # 末条预渲染条目段落（含 _Toc 书签超链接）：最后一个 fldChar 是外层 TOC end
    entry_indices = [
        i
        for i in range(toc_idx, len(paragraphs))
        if any(
            (hl.get(qn("w:anchor")) or "").startswith("_Toc")
            for hl in paragraphs[i]._p.findall(qn("w:hyperlink"))
        )
    ]
    assert entry_indices, "未找到预渲染条目"
    last_entry_fld_types = [
        fc.get(qn("w:fldCharType"))
        for fc in paragraphs[entry_indices[-1]]._p.iter(qn("w:fldChar"))
    ]
    # 末条条目内：PAGEREF(begin/separate/end) + 外层 TOC end
    assert last_entry_fld_types[-1] == "end"
    assert last_entry_fld_types.count("begin") == 1, "仅应有 PAGEREF 的 begin"
    assert last_entry_fld_types.count("end") == 2, "PAGEREF end + 外层 TOC end"


def test_export_toc_entries_immediately_follow_title_without_blank_line(
    session: Session,
    tmp_path: Path,
) -> None:
    """"目录"标题段之后紧跟首条目录条目，中间无空行；条目紧邻标题不在文档末尾。"""
    project = create_project(session)
    create_form(session, project.id, name="生命体征", order_index=1)
    create_form(session, project.id, name="实验室", order_index=2)

    doc = export_document(session, project.id, tmp_path)

    paragraphs = doc.paragraphs

    # 找"目录"标题段
    title_idx = next(
        i for i, para in enumerate(paragraphs) if (para.text or "").strip() == "目录"
    )

    # 紧邻的下一段即首条目录条目：含 _Toc 超链接（若中间有空行，此段将无超链接）
    nxt = paragraphs[title_idx + 1]
    anchors = [hl.get(qn("w:anchor")) for hl in nxt._p.findall(qn("w:hyperlink"))]
    assert any(a and a.startswith("_Toc") for a in anchors), (
        "目录标题段后应紧跟首条目录条目（中间不得有空行/空域壳段）"
    )
    # 首条条目同时合入 TOC 域起始（begin/separate）
    fld_types = [fc.get(qn("w:fldCharType")) for fc in nxt._p.iter(qn("w:fldChar"))]
    assert "begin" in fld_types and "separate" in fld_types

    # 条目紧邻标题，不应远在文档末尾
    last_entry_idx = max(
        i
        for i in range(title_idx + 1, len(paragraphs))
        if any(
            (hl.get(qn("w:anchor")) or "").startswith("_Toc")
            for hl in paragraphs[i]._p.findall(qn("w:hyperlink"))
        )
    )
    assert last_entry_idx - title_idx <= 10


def _toc_entry_page_numbers(doc: Document) -> list[str]:
    """提取每条目录条目 PAGEREF 域的页码文本（页码是条目段落最后一个 w:t）。"""
    numbers: list[str] = []
    for paragraph in doc.paragraphs:
        has_toc_anchor = any(
            (hl.get(qn("w:anchor")) or "").startswith("_Toc")
            for hl in paragraph._p.findall(qn("w:hyperlink"))
        )
        if not has_toc_anchor:
            continue
        texts = [t.text or "" for t in paragraph._p.iter(qn("w:t"))]
        numbers.append(texts[-1] if texts else "")
    return numbers


def test_export_toc_keeps_non_empty_fallback_when_baking_returns_empty(
    session: Session,
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
    caplog: pytest.LogCaptureFixture,
) -> None:
    """真实页码计算失败时，目录页码保留非空回退值并保留 Word 域。"""
    project = create_project(session)
    create_form(session, project.id, name="生命体征", order_index=1)
    create_form(session, project.id, name="实验室", order_index=2)
    monkeypatch.setattr(toc_pagination, "compute_heading_pages", lambda _path: {})

    output_path = tmp_path / "fallback-empty.docx"
    with caplog.at_level(logging.WARNING, logger="src.services.export_service"):
        ok = ExportService(session).export_project_to_word(
            project.id,
            str(output_path),
            bake_toc_page_numbers=True,
        )

    assert ok is True
    doc = Document(output_path)
    numbers = _toc_entry_page_numbers(doc)
    assert numbers
    assert all(number.strip() for number in numbers)
    assert all(number.isdigit() for number in numbers)
    assert "非空回退页码" in caplog.text
    assert doc.settings.element.find(qn("w:updateFields")) is not None
    assert any(
        "PAGEREF" in (instr.text or "")
        for paragraph in doc.paragraphs
        for instr in paragraph._p.iter(qn("w:instrText"))
    )


def test_export_toc_bakes_known_pages_and_keeps_fallback_for_missing(
    session: Session,
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
    caplog: pytest.LogCaptureFixture,
) -> None:
    """真实页码只覆盖命中的目录项，未命中项仍保留非空回退值。"""
    project = create_project(session)
    create_form(session, project.id, name="生命体征", order_index=1)
    create_form(session, project.id, name="实验室", order_index=2)
    monkeypatch.setattr(
        toc_pagination,
        "compute_heading_pages",
        lambda _path: {"1. 生命体征": 7},
    )

    output_path = tmp_path / "fallback-partial.docx"
    with caplog.at_level(logging.WARNING, logger="src.services.export_service"):
        ok = ExportService(session).export_project_to_word(
            project.id,
            str(output_path),
            bake_toc_page_numbers=True,
        )

    assert ok is True
    numbers = _toc_entry_page_numbers(Document(output_path))
    assert numbers
    assert "7" in numbers
    assert any(number == "1" for number in numbers)
    assert all(number.strip() for number in numbers)
    assert "部分目录页码未取得" in caplog.text


@pytest.mark.skipif(
    toc_pagination.find_libreoffice() is None,
    reason="需要 LibreOffice 渲染真实页码",
)
def test_export_toc_bakes_real_page_numbers_with_libreoffice(
    session: Session,
    tmp_path: Path,
) -> None:
    """LibreOffice 可用时，目录每条页码被写死为真实页码（正整数、非空、递增）。"""
    project = create_project(session)
    for i in range(1, 6):
        create_form(session, project.id, name=f"表单{i}", order_index=i)

    output_path = tmp_path / "baked.docx"
    ok = ExportService(session).export_project_to_word(
        project.id, str(output_path), bake_toc_page_numbers=True
    )
    assert ok is True

    doc = Document(output_path)
    numbers = _toc_entry_page_numbers(doc)
    assert numbers, "未找到目录条目"
    # 每条目录都写入了真实页码（数字），且随条目递增（表单依次在后续页）
    assert all(n.isdigit() for n in numbers), f"存在非数字页码: {numbers}"
    pages = [int(n) for n in numbers]
    assert all(p >= 1 for p in pages)
    assert pages == sorted(pages)


def test_export_project_uses_next_page_section_break_between_portrait_forms(
    session: Session,
    tmp_path: Path,
) -> None:
    project = create_project(session)
    first = create_form(session, project.id, name="生命体征", order_index=1)
    second = create_form(session, project.id, name="实验室", order_index=2)
    visit = create_visit(session, project.id, name="筛选期", sequence=1)
    create_visit_form(session, visit.id, first.id, sequence=1)
    create_visit_form(session, visit.id, second.id, sequence=2)

    first_field = create_field_definition(session, project.id, variable_name="SYSBP", label="收缩压")
    second_field = create_field_definition(session, project.id, variable_name="ALT", label="谷丙转氨酶")
    create_form_field(session, first.id, first_field.id, order_index=1)
    create_form_field(session, second.id, second_field.id, order_index=1)

    output_path = tmp_path / "sections.docx"
    switch_calls: list[WD_ORIENT] = []
    original_switch = ExportService._switch_section

    def _spy(self, doc, orientation, project_arg):
        switch_calls.append(orientation)
        return original_switch(self, doc, orientation, project_arg)

    with patch.object(ExportService, "_switch_section", _spy):
        ok = ExportService(session).export_project_to_word(project.id, str(output_path))

    assert ok is True
    doc = Document(output_path)
    assert switch_calls.count(WD_ORIENT.PORTRAIT) >= 2
    assert len(doc.sections) >= 4
    assert all(section.start_type == WD_SECTION.NEW_PAGE for section in doc.sections[1:])



def test_export_project_sets_form_table_rows_to_at_least_one_centimeter(
    session: Session,
    tmp_path: Path,
) -> None:
    project = create_project(session)
    form = create_form(session, project.id, name="生命体征", order_index=1)
    visit = create_visit(session, project.id, name="筛选期", sequence=1)
    create_visit_form(session, visit.id, form.id, sequence=1)

    systolic = create_field_definition(
        session,
        project.id,
        variable_name="SYSBP",
        label="收缩压",
    )
    diastolic = create_field_definition(
        session,
        project.id,
        variable_name="DIABP",
        label="舒张压",
    )
    section_label = create_field_definition(
        session,
        project.id,
        variable_name="SECTION_LABEL",
        label="给药后记录",
        field_type="标签",
    )
    create_form_field(session, form.id, systolic.id, order_index=1, default_value="120")
    session.add(
        FormField(
            form_id=form.id,
            field_definition_id=None,
            is_log_row=1,
            order_index=2,
            label_override="日志记录",
        )
    )
    session.flush()
    create_form_field(session, form.id, section_label.id, order_index=3)
    create_form_field(session, form.id, diastolic.id, order_index=4, default_value="80")

    doc = export_document(session, project.id, tmp_path)

    cover_table = doc.tables[0]
    assert_table_rows_at_least_one_centimeter(cover_table)

    form_table = doc.tables[2]
    assert len(form_table.rows) == 4
    assert_table_rows_at_least_one_centimeter(form_table)


def test_export_project_preserves_mixed_normal_inline_group_order(
    session: Session,
    tmp_path: Path,
) -> None:
    project = create_project(session)
    form = create_form(session, project.id, name="12导联心电图", order_index=1)
    visit = create_visit(session, project.id, name="筛选期", sequence=1)
    create_visit_form(session, visit.id, form.id, sequence=1)

    date_field = create_field_definition(session, project.id, variable_name="ECG_DATE", label="检查日期")
    item_field = create_field_definition(session, project.id, variable_name="ECG_ITEM", label="项目")
    result_field = create_field_definition(session, project.id, variable_name="ECG_RESULT", label="结果")
    judgement_field = create_field_definition(session, project.id, variable_name="ECG_JUDGE", label="综合判定结果")
    create_form_field(session, form.id, date_field.id, order_index=1)
    inline_item = create_form_field(session, form.id, item_field.id, order_index=2)
    inline_result = create_form_field(session, form.id, result_field.id, order_index=3)
    inline_item.inline_mark = 1
    inline_result.inline_mark = 1
    create_form_field(session, form.id, judgement_field.id, order_index=4)

    doc = export_document(session, project.id, tmp_path)

    form_tables = doc.tables[2:]
    assert len(form_tables) == 3
    assert form_tables[0].cell(0, 0).text.strip() == "检查日期"
    assert form_tables[1].cell(0, 0).text.strip() == "项目"
    assert form_tables[1].cell(0, 1).text.strip() == "结果"
    assert form_tables[2].cell(0, 0).text.strip() == "综合判定结果"



def test_export_project_visit_flow_uses_cross_marks_and_order_index_sorting(
    session: Session,
    tmp_path: Path,
) -> None:
    project = create_project(session)
    alpha = create_form(session, project.id, name="Alpha", order_index=2)
    zeta = create_form(session, project.id, name="Zeta", order_index=1)
    later_by_id = create_form(session, project.id, name="LaterById", order_index=None)
    last_by_id = create_form(session, project.id, name="LastById", order_index=None)
    visit = create_visit(session, project.id, name="筛选期", sequence=1)
    create_visit_form(session, visit.id, alpha.id, sequence=9)
    create_visit_form(session, visit.id, zeta.id, sequence=5)
    create_visit_form(session, visit.id, later_by_id.id, sequence=4)
    create_visit_form(session, visit.id, last_by_id.id, sequence=3)

    doc = export_document(session, project.id, tmp_path)

    visit_flow_table = doc.tables[1]
    assert visit_flow_table.cell(1, 0).text.strip() == "Zeta"
    assert visit_flow_table.cell(2, 0).text.strip() == "Alpha"
    assert visit_flow_table.cell(3, 0).text.strip() == "LaterById"
    assert visit_flow_table.cell(4, 0).text.strip() == "LastById"
    assert visit_flow_table.cell(1, 1).text.strip() == "×"
    assert visit_flow_table.cell(2, 1).text.strip() == "×"
    assert visit_flow_table.cell(3, 1).text.strip() == "×"
    assert visit_flow_table.cell(4, 1).text.strip() == "×"
    assert_table_rows_at_least_one_centimeter(visit_flow_table)


def _find_tbl_headers(tr) -> list:
    from docx.oxml.ns import qn
    trPr = tr.find(qn('w:trPr'))
    if trPr is None:
        return []
    return [el for el in trPr.findall(qn('w:tblHeader')) if el.get(qn('w:val')) == 'true']


def _find_applicable_visits_paragraphs(doc) -> list:
    return [p for p in doc.paragraphs if p.style and p.style.name == 'ApplicableVisits']


def test_export_project_visit_flow_header_row_sets_tblHeader(
    session: Session,
    tmp_path: Path,
) -> None:
    project = create_project(session)
    form = create_form(session, project.id, name="生命体征", order_index=1)
    visit = create_visit(session, project.id, name="筛选期", sequence=1)
    create_visit_form(session, visit.id, form.id, sequence=1)

    doc = export_document(session, project.id, tmp_path)

    visit_flow_table = doc.tables[1]
    assert len(_find_tbl_headers(visit_flow_table.rows[0]._tr)) == 1


def test_export_project_no_visits_skeleton_header_row_still_sets_tblHeader(
    session: Session,
    tmp_path: Path,
) -> None:
    project = create_project(session)
    create_form(session, project.id, name="孤立表单", order_index=1)

    doc = export_document(session, project.id, tmp_path)

    visit_flow_table = doc.tables[1]
    assert len(_find_tbl_headers(visit_flow_table.rows[0]._tr)) == 1


def test_export_project_applicable_visits_footer_uses_sequence_order_and_matches_header_order(
    session: Session,
    tmp_path: Path,
) -> None:
    project = create_project(session)
    form = create_form(session, project.id, name="实验室", order_index=1)
    v3 = create_visit(session, project.id, name="V3", sequence=3)
    v1 = create_visit(session, project.id, name="V1", sequence=1)
    v2 = create_visit(session, project.id, name="V2", sequence=2)
    create_visit_form(session, v3.id, form.id, sequence=10)
    create_visit_form(session, v1.id, form.id, sequence=7)
    create_visit_form(session, v2.id, form.id, sequence=5)

    doc = export_document(session, project.id, tmp_path)

    visit_flow_table = doc.tables[1]
    assert visit_flow_table.cell(0, 1).text.strip() == "V1"
    assert visit_flow_table.cell(0, 2).text.strip() == "V2"
    assert visit_flow_table.cell(0, 3).text.strip() == "V3"

    paragraphs = _find_applicable_visits_paragraphs(doc)
    assert len(paragraphs) == 1
    assert paragraphs[0].text == "适用访视：V1、V2、V3"


def test_export_project_skips_applicable_visits_footer_for_orphan_form(
    session: Session,
    tmp_path: Path,
) -> None:
    project = create_project(session)
    create_form(session, project.id, name="孤立表单", order_index=1)
    create_visit(session, project.id, name="筛选期", sequence=1)

    doc = export_document(session, project.id, tmp_path)

    assert _find_applicable_visits_paragraphs(doc) == []


def test_export_project_applicable_visits_footer_prefix_bold_and_names_run_use_expected_fonts(
    session: Session,
    tmp_path: Path,
) -> None:
    from docx.oxml.ns import qn

    project = create_project(session)
    form = create_form(session, project.id, name="实验室", order_index=1)
    visit = create_visit(session, project.id, name="筛选期", sequence=1)
    create_visit_form(session, visit.id, form.id, sequence=1)

    doc = export_document(session, project.id, tmp_path)

    paragraphs = _find_applicable_visits_paragraphs(doc)
    assert len(paragraphs) == 1
    paragraph = paragraphs[0]
    assert paragraph.style.name == "ApplicableVisits"
    assert len(paragraph.runs) == 2
    prefix_run, names_run = paragraph.runs
    assert prefix_run.text == "适用访视："
    assert prefix_run.bold is True
    assert names_run.text == "筛选期"
    assert names_run.bold is not True

    for run in paragraph.runs:
        rFonts = run._element.find(qn('w:rPr')).find(qn('w:rFonts'))
        assert rFonts.get(qn('w:eastAsia')) == "SimSun"
        assert rFonts.get(qn('w:ascii')) == "Times New Roman"


def test_export_project_groups_adjacent_inline_fields_into_one_table(
    session: Session,
    tmp_path: Path,
) -> None:
    project = create_project(session)
    form = create_form(session, project.id, name="实验室检查", order_index=1)

    field_a = create_field_definition(
        session,
        project.id,
        variable_name="LAB_A",
        label="字段A",
    )
    field_b = create_field_definition(
        session,
        project.id,
        variable_name="LAB_B",
        label="字段B",
    )
    create_form_field(session, form.id, field_a.id, order_index=1, default_value="第一行\n第二行")
    create_form_field(session, form.id, field_b.id, order_index=2, default_value="仅一行")

    form.form_fields[0].inline_mark = 1
    form.form_fields[0].label_override = "覆盖标签A"
    form.form_fields[1].inline_mark = 1
    session.flush()

    doc = export_document(session, project.id, tmp_path)

    assert len(doc.tables) == 3
    inline_table = doc.tables[2]
    assert len(inline_table.columns) == 2
    assert inline_table.cell(0, 0).text.strip() == "覆盖标签A"
    assert inline_table.cell(0, 1).text.strip() == "字段B"
    assert inline_table.cell(1, 0).text.strip() == "第一行"
    assert inline_table.cell(1, 1).text.strip() == "仅一行"
    assert inline_table.cell(2, 0).text.strip() == "第二行"
    assert_table_rows_at_least_one_centimeter(inline_table)


def test_build_unified_table_sets_all_rows_to_at_least_one_centimeter(
    session: Session,
) -> None:
    project = create_project(session)
    form = create_form(session, project.id, name="统一横向表", order_index=1)
    regular_def = create_field_definition(
        session,
        project.id,
        variable_name="REGULAR",
        label="普通字段",
    )
    label_def = create_field_definition(
        session,
        project.id,
        variable_name="SECTION",
        label="分区标签",
        field_type="标签",
    )
    inline_a_def = create_field_definition(
        session,
        project.id,
        variable_name="INLINE_A",
        label="内联A",
    )
    inline_b_def = create_field_definition(
        session,
        project.id,
        variable_name="INLINE_B",
        label="内联B",
    )
    regular = create_form_field(session, form.id, regular_def.id, order_index=1)
    full_row = create_form_field(session, form.id, label_def.id, order_index=2)
    inline_a = create_form_field(
        session,
        form.id,
        inline_a_def.id,
        order_index=3,
        default_value="第一行\n第二行",
    )
    inline_b = create_form_field(
        session,
        form.id,
        inline_b_def.id,
        order_index=4,
        default_value="仅一行",
    )
    inline_a.inline_mark = 1
    inline_b.inline_mark = 1
    session.flush()

    service = ExportService(session)
    service._column_width_overrides = {}
    doc = Document()
    service._apply_document_style(doc)
    table = service._build_unified_table(
        doc,
        [
            Segment("regular_field", [regular]),
            Segment("full_row", [full_row]),
            Segment("inline_block", [inline_a, inline_b]),
        ],
        LayoutDecision("unified_landscape", 4, 1, 3),
        form_id=form.id,
    )

    assert len(table.rows) == 5
    assert_table_rows_at_least_one_centimeter(table)


def test_export_project_renders_cover_table_with_three_rows_two_cols(
    session: Session,
    tmp_path: Path,
) -> None:
    project = create_project(session)
    project.trial_name = "试验名称"
    project.protocol_number = "P-001"
    project.crf_version = "V1.0"
    doc = export_document(session, project.id, tmp_path)

    assert any("试验名称" in p.text for p in doc.paragraphs)
    assert any("V1.0" in p.text for p in doc.paragraphs)
    cover_table = doc.tables[0]
    assert len(cover_table.rows) == 3
    assert len(cover_table.columns) == 2
    assert cover_table.cell(0, 0).text.strip() == "方案编号"
    assert cover_table.cell(0, 1).text.strip() == "P-001"
    assert cover_table.cell(1, 0).text.strip() == "中心编号"
    assert cover_table.cell(2, 0).text.strip() == "筛选号"


def test_export_project_uses_visit_flow_skeleton_when_no_visits(
    session: Session,
    tmp_path: Path,
) -> None:
    project = create_project(session)

    doc = export_document(session, project.id, tmp_path)

    visit_flow_table = doc.tables[1]
    assert len(visit_flow_table.rows) == 1
    assert len(visit_flow_table.columns) == 1
    assert visit_flow_table.cell(0, 0).text.strip() == "访视名称"
    assert "暂无访视数据" not in [paragraph.text for paragraph in doc.paragraphs]


def test_export_project_renders_empty_form_as_single_empty_skeleton_row(
    session: Session,
    tmp_path: Path,
) -> None:
    project = create_project(session)
    create_form(session, project.id, name="空白表", order_index=1)

    doc = export_document(session, project.id, tmp_path)

    form_table = doc.tables[2]
    assert len(form_table.rows) == 1
    assert len(form_table.columns) == 2
    assert form_table.cell(0, 0).text.strip() == ""
    assert form_table.cell(0, 1).text.strip() == ""


def test_export_project_table_count_matches_cover_visit_and_forms(
    session: Session,
    tmp_path: Path,
) -> None:
    project = create_project(session)
    form_a = create_form(session, project.id, name="实验室检查", order_index=1)
    form_b = create_form(session, project.id, name="生命体征", order_index=2)
    visit_a = create_visit(session, project.id, name="筛选期", sequence=1)
    visit_b = create_visit(session, project.id, name="基线期", sequence=2)
    create_visit_form(session, visit_a.id, form_a.id, sequence=1)
    create_visit_form(session, visit_b.id, form_b.id, sequence=1)

    fd_a = create_field_definition(session, project.id, variable_name="LAB", label="实验室")
    fd_b = create_field_definition(session, project.id, variable_name="VITAL", label="体征")
    create_form_field(session, form_a.id, fd_a.id, order_index=1, default_value="A")
    create_form_field(session, form_b.id, fd_b.id, order_index=1, default_value="B")

    doc = export_document(session, project.id, tmp_path)

    assert len(doc.tables) == 2 + len(project.forms)
    visit_flow_table = doc.tables[1]
    assert len(visit_flow_table.rows) == len(project.forms) + 1
    assert len(visit_flow_table.columns) == len(project.visits) + 1


def test_export_no_forms_produces_3_tables(
    session: Session,
    tmp_path: Path,
) -> None:
    project = create_project(session)

    doc = export_document(session, project.id, tmp_path)

    assert len(doc.tables) >= 3
    empty_form_table = doc.tables[2]
    assert len(empty_form_table.rows) == 1
    assert len(empty_form_table.columns) == 2


def test_export_project_preserves_skip_first_two_tables_import_assumption(
    session: Session,
    tmp_path: Path,
) -> None:
    project = create_project(session)
    form = create_form(session, project.id, name="实验室检查", order_index=1)
    visit = create_visit(session, project.id, name="筛选期", sequence=1)
    create_visit_form(session, visit.id, form.id, sequence=1)
    field_definition = create_field_definition(session, project.id, variable_name="LAB", label="实验室")
    create_form_field(session, form.id, field_definition.id, order_index=1, default_value="值")

    doc = export_document(session, project.id, tmp_path)

    assert doc.tables[0].cell(0, 0).text.strip() != "访视名称"
    assert doc.tables[1].cell(0, 0).text.strip() == "访视名称"
    assert doc.tables[2].cell(0, 0).text.strip() == "实验室"


def test_export_cover_uses_project_screening_number_format_when_set(
    session: Session,
    tmp_path: Path,
) -> None:
    project = create_project(session)
    project.screening_number_format = "X|__|__|-|__|__|"

    doc = export_document(session, project.id, tmp_path)

    cover_table = doc.tables[0]
    assert cover_table.cell(2, 0).text.strip() == "筛选号"
    assert cover_table.cell(2, 1).text.strip() == "X|__|__|-|__|__|"


def test_export_cover_falls_back_to_default_screening_number_format(
    session: Session,
    tmp_path: Path,
) -> None:
    project = create_project(session)
    project.screening_number_format = None

    doc = export_document(session, project.id, tmp_path)

    cover_table = doc.tables[0]
    assert cover_table.cell(2, 0).text.strip() == "筛选号"
    assert cover_table.cell(2, 1).text.strip() == "S|__|__||__|__|__|"


def test_export_cover_falls_back_to_default_for_legacy_whitespace_value(
    session: Session,
    tmp_path: Path,
) -> None:
    """legacy 存量 whitespace 值也必须走默认回退，与 UI/API 的归一化语义保持一致。"""
    project = create_project(session)
    project.screening_number_format = "   "

    doc = export_document(session, project.id, tmp_path)

    cover_table = doc.tables[0]
    assert cover_table.cell(2, 1).text.strip() == "S|__|__||__|__|__|"


def test_export_visit_flow_has_landscape_section_restored_to_portrait(
    session: Session,
    tmp_path: Path,
) -> None:
    project = create_project(session)
    form = create_form(session, project.id, name="实验室检查", order_index=1)
    visit = create_visit(session, project.id, name="筛选期", sequence=1)
    create_visit_form(session, visit.id, form.id, sequence=1)
    field_definition = create_field_definition(session, project.id, variable_name="LAB", label="实验室")
    create_form_field(session, form.id, field_definition.id, order_index=1, default_value="值")

    doc = export_document(session, project.id, tmp_path)

    orientations = [section.orientation for section in doc.sections]
    assert len(orientations) >= 3
    assert orientations[0] == WD_ORIENT.PORTRAIT
    assert orientations[1] == WD_ORIENT.LANDSCAPE
    assert orientations[2] == WD_ORIENT.PORTRAIT
    assert doc.tables[0].cell(0, 0).text.strip() != "访视名称"
    assert doc.tables[1].cell(0, 0).text.strip() == "访视名称"


# ---------------------------------------------------------------------------
# Phase 4.4: PBT 属性测试（hypothesis）
# ---------------------------------------------------------------------------

def _make_in_memory_session():
    """创建内存 SQLite session（供 hypothesis 测试使用，独立于 pytest fixture）"""
    engine = create_engine("sqlite+pysqlite:///:memory:")
    Base.metadata.create_all(engine)
    return sessionmaker(bind=engine, expire_on_commit=False)(), engine


@given(form_count=st.integers(min_value=0, max_value=5))
@settings(max_examples=15, deadline=None)
def test_pbt_p1_export_always_produces_nonempty_file(form_count: int) -> None:
    """P1: 任意表单数量的项目导出后文件大小 > 0"""
    db, engine = _make_in_memory_session()
    try:
        from src.models.project import Project
        from src.models.form import Form

        project = Project(name="PBT项目", version="v1.0")
        db.add(project)
        db.flush()
        for i in range(form_count):
            db.add(Form(project_id=project.id, name=f"表单{i}", code=f"F{i}", order_index=i))
        db.flush()

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            output_path = f.name
        try:
            ok = ExportService(db).export_project_to_word(project.id, output_path)
            assert ok is True, "export_project_to_word 应返回 True"
            assert os.path.getsize(output_path) > 0, "导出文件不应为 0 字节"
        finally:
            os.unlink(output_path)
    finally:
        db.close()
        engine.dispose()


@given(form_count=st.integers(min_value=0, max_value=5))
@settings(max_examples=15, deadline=None)
def test_pbt_p2_export_always_produces_valid_docx(form_count: int) -> None:
    """P2: 任意表单数量的项目导出后 Document() 可正常打开"""
    db, engine = _make_in_memory_session()
    try:
        from src.models.project import Project
        from src.models.form import Form

        project = Project(name="PBT项目", version="v1.0")
        db.add(project)
        db.flush()
        for i in range(form_count):
            db.add(Form(project_id=project.id, name=f"表单{i}", code=f"F{i}", order_index=i))
        db.flush()

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            output_path = f.name
        try:
            ExportService(db).export_project_to_word(project.id, output_path)
            doc = Document(output_path)  # 不应抛出异常
            assert len(doc.tables) >= 3, "导出文档应至少包含3张表格"
        finally:
            os.unlink(output_path)
    finally:
        db.close()
        engine.dispose()


@given(form_count=st.integers(min_value=1, max_value=4))
@settings(max_examples=10, deadline=None)
def test_pbt_p8_export_is_idempotent_in_structure(form_count: int) -> None:
    """P8: 同一项目导出两次，表格数量结构一致（幂等性）"""
    db, engine = _make_in_memory_session()
    try:
        from src.models.project import Project
        from src.models.form import Form

        project = Project(name="PBT项目", version="v1.0")
        db.add(project)
        db.flush()
        for i in range(form_count):
            db.add(Form(project_id=project.id, name=f"表单{i}", code=f"F{i}", order_index=i))
        db.flush()

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f1, \
             tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f2:
            path1, path2 = f1.name, f2.name
        try:
            ExportService(db).export_project_to_word(project.id, path1)
            ExportService(db).export_project_to_word(project.id, path2)
            doc1, doc2 = Document(path1), Document(path2)
            assert len(doc1.tables) == len(doc2.tables), "两次导出表格数量应相同"
        finally:
            os.unlink(path1)
            os.unlink(path2)
    finally:
        db.close()
        engine.dispose()


@given(trial_name=st.one_of(
    st.none(),
    st.text(
        alphabet=st.characters(min_codepoint=32, max_codepoint=0x9FFF, blacklist_categories=("Cc", "Cs")),
        min_size=1,
        max_size=50,
    ),
))
@settings(max_examples=10, deadline=None)
def test_pbt_p9_empty_form_always_has_skeleton_table(trial_name) -> None:
    """P9: 空表单（无字段）导出后仍有骨架表格"""
    db, engine = _make_in_memory_session()
    try:
        from src.models.project import Project
        from src.models.form import Form

        project = Project(name="PBT空表单", version="v1.0", trial_name=trial_name)
        db.add(project)
        db.flush()
        db.add(Form(project_id=project.id, name="空白表单", code="EMPTY", order_index=1))
        db.flush()

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            output_path = f.name
        try:
            ExportService(db).export_project_to_word(project.id, output_path)
            doc = Document(output_path)
            assert len(doc.tables) >= 3, "空表单项目也应有 >= 3 张表格（封面+访视图+骨架）"
            form_table = doc.tables[2]
            assert len(form_table.rows) >= 1
            assert len(form_table.columns) == 2
        finally:
            os.unlink(output_path)
    finally:
        db.close()
        engine.dispose()
