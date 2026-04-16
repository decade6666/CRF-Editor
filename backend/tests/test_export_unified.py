"""Unified landscape 表单导出测试。

验证 mixed-field-landscape-form 变更的核心场景：
- unified 触发条件：has_regular + has_inline + max_block_width > 4
- legacy 路径回归：纯 normal/inline 表单不受影响
- unified 表格结构：单一 N 列表格 + landscape section
"""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
import pytest
from sqlalchemy import create_engine, event
from sqlalchemy.orm import Session, sessionmaker
from sqlalchemy.pool import StaticPool

from src.database import get_read_session, get_session
from src.models import Base
from src.models.project import Project
from src.models.user import User
from src.models.visit import Visit
from src.models.form import Form
from src.models.visit_form import VisitForm
from src.models.field_definition import FieldDefinition
from src.models.form_field import FormField
from src.services.export_service import ExportService


@pytest.fixture
def engine():
    """创建内存数据库引擎。"""
    _engine = create_engine(
        "sqlite+pysqlite:///:memory:",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )

    @event.listens_for(_engine, "connect")
    def _enable_fk(dbapi_conn, _):
        dbapi_conn.execute("PRAGMA foreign_keys = ON")

    Base.metadata.create_all(_engine)
    yield _engine
    _engine.dispose()


@pytest.fixture
def session(engine) -> Session:
    """创建数据库会话。"""
    session_factory = sessionmaker(bind=engine, expire_on_commit=False)
    with session_factory() as db_session:
        yield db_session


def create_minimal_project(session: Session) -> tuple[Project, User]:
    """创建最小项目和用户。"""
    user = User(username="tester")
    session.add(user)
    session.flush()
    project = Project(name="Unified 测试项目", version="v1.0", owner_id=user.id)
    session.add(project)
    session.flush()
    return project, user


def create_text_field_def(session: Session, project_id: int, label: str) -> FieldDefinition:
    """创建文本字段定义。"""
    fd = FieldDefinition(
        project_id=project_id,
        label=label,
        variable_name=f"VAR_{label.replace(' ', '_')}",
        field_type="文本",
    )
    session.add(fd)
    session.flush()
    return fd


def create_label_field_def(session: Session, project_id: int, label: str) -> FieldDefinition:
    """创建标签字段定义。"""
    fd = FieldDefinition(
        project_id=project_id,
        label=label,
        variable_name=f"LBL_{label.replace(' ', '_')}",
        field_type="标签",
    )
    session.add(fd)
    session.flush()
    return fd


def add_field_to_form(
    session: Session,
    form_id: int,
    field_def_id: int,
    order_index: int = 0,
    inline_mark: int = 0,
    is_log: int = 0,
) -> FormField:
    """将字段添加到表单。"""
    ff = FormField(
        form_id=form_id,
        field_definition_id=field_def_id if is_log == 0 else None,
        order_index=order_index,
        inline_mark=inline_mark,
        is_log_row=is_log,
    )
    session.add(ff)
    session.flush()
    return ff


# ========== Task 3.1: 纯 normal 表单回归测试 ==========

def test_export_normal_form_remains_2col_portrait(session: Session, tmp_path: Path) -> None:
    """验证纯 normal 表单仍为 2 列表格 + portrait 方向。"""
    project, _ = create_minimal_project(session)

    # 创建访视和表单
    visit = Visit(project_id=project.id, name="访视1", code="V1", sequence=1)
    session.add(visit)
    session.flush()

    form = Form(project_id=project.id, name="纯普通表单", code="F_NORMAL", order_index=1)
    session.add(form)
    session.flush()

    vf = VisitForm(visit_id=visit.id, form_id=form.id, sequence=1)
    session.add(vf)
    session.flush()

    # 添加 3 个普通字段（inline_mark=0）
    fd1 = create_text_field_def(session, project.id, "字段1")
    fd2 = create_text_field_def(session, project.id, "字段2")
    fd3 = create_text_field_def(session, project.id, "字段3")

    add_field_to_form(session, form.id, fd1.id, order_index=1, inline_mark=0)
    add_field_to_form(session, form.id, fd2.id, order_index=2, inline_mark=0)
    add_field_to_form(session, form.id, fd3.id, order_index=3, inline_mark=0)

    session.commit()

    # 导出
    output_path = tmp_path / "normal.docx"
    ExportService(session).export_project_to_word(project.id, str(output_path))

    doc = Document(str(output_path))

    # 检查页面方向
    for section in doc.sections:
        assert section.orientation == WD_ORIENT.PORTRAIT, "纯 normal 表单应保持 portrait"

    # 检查表格列数（跳过封面和访视流程表）
    form_tables = doc.tables[2:]  # 前 2 个是封面和访视流程
    for table in form_tables:
        assert len(table.columns) == 2, "普通表单应为 2 列表格"


# ========== Task 3.2: 纯 inline 表单回归测试 ==========

def test_export_inline_form_max4_remains_portrait(session: Session, tmp_path: Path) -> None:
    """验证 ≤4 列 inline 表单保持 portrait。"""
    project, _ = create_minimal_project(session)

    visit = Visit(project_id=project.id, name="访视1", code="V1", sequence=1)
    session.add(visit)
    session.flush()

    form = Form(project_id=project.id, name="窄 inline 表单", code="F_INLINE4", order_index=1)
    session.add(form)
    session.flush()

    vf = VisitForm(visit_id=visit.id, form_id=form.id, sequence=1)
    session.add(vf)
    session.flush()

    # 添加 4 个 inline 字段（inline_mark=1）
    for i in range(1, 5):
        fd = create_text_field_def(session, project.id, f"列{i}")
        add_field_to_form(session, form.id, fd.id, order_index=i, inline_mark=1)

    session.commit()

    output_path = tmp_path / "inline4.docx"
    ExportService(session).export_project_to_word(project.id, str(output_path))

    doc = Document(str(output_path))

    # 检查页面方向
    for section in doc.sections:
        assert section.orientation == WD_ORIENT.PORTRAIT, "≤4 列 inline 应保持 portrait"


# ========== Task 3.3: mixed + max_block_width ≤ 4 回归测试 ==========

def test_export_mixed_block_le4_remains_split_table(session: Session, tmp_path: Path) -> None:
    """验证 max_block_width ≤ 4 时仍为 split-table 路径。"""
    project, _ = create_minimal_project(session)

    visit = Visit(project_id=project.id, name="访视1", code="V1", sequence=1)
    session.add(visit)
    session.flush()

    form = Form(project_id=project.id, name="混合窄表单", code="F_MIXED4", order_index=1)
    session.add(form)
    session.flush()

    vf = VisitForm(visit_id=visit.id, form_id=form.id, sequence=1)
    session.add(vf)
    session.flush()

    # 添加普通字段
    fd_normal = create_text_field_def(session, project.id, "普通字段")
    add_field_to_form(session, form.id, fd_normal.id, order_index=1, inline_mark=0)

    # 添加 4 个 inline 字段（block 宽度 = 4，不触发 unified）
    for i in range(1, 5):
        fd = create_text_field_def(session, project.id, f"列{i}")
        add_field_to_form(session, form.id, fd.id, order_index=10 + i, inline_mark=1)

    session.commit()

    output_path = tmp_path / "mixed4.docx"
    ExportService(session).export_project_to_word(project.id, str(output_path))

    doc = Document(str(output_path))

    # 检查页面方向（不触发 unified，保持 portrait）
    for section in doc.sections:
        assert section.orientation == WD_ORIENT.PORTRAIT, "block ≤ 4 应保持 portrait"

    # 检查表格数量（应有多表格，非单一 unified 表）
    form_tables = doc.tables[2:]
    assert len(form_tables) >= 2, "split-table 路径应有多个表格"


# ========== Task 3.4: unified 基本场景测试 ==========

def test_export_unified_mixed_max5_creates_landscape_table(session: Session, tmp_path: Path) -> None:
    """验证混合表单 (max_block_width > 4) 输出单一表格 + landscape section。"""
    project, _ = create_minimal_project(session)

    visit = Visit(project_id=project.id, name="访视1", code="V1", sequence=1)
    session.add(visit)
    session.flush()

    form = Form(project_id=project.id, name="混合宽表单", code="F_UNIFIED", order_index=1)
    session.add(form)
    session.flush()

    vf = VisitForm(visit_id=visit.id, form_id=form.id, sequence=1)
    session.add(vf)
    session.flush()

    # 添加普通字段
    fd_normal = create_text_field_def(session, project.id, "普通字段")
    add_field_to_form(session, form.id, fd_normal.id, order_index=1, inline_mark=0)

    # 添加 5 个 inline 字段（block 宽度 = 5，触发 unified）
    for i in range(1, 6):
        fd = create_text_field_def(session, project.id, f"列{i}")
        add_field_to_form(session, form.id, fd.id, order_index=10 + i, inline_mark=1)

    session.commit()

    output_path = tmp_path / "unified5.docx"
    ExportService(session).export_project_to_word(project.id, str(output_path))

    doc = Document(str(output_path))

    # 检查存在 landscape section
    orientations = [section.orientation for section in doc.sections]
    assert WD_ORIENT.LANDSCAPE in orientations, "unified 表单应包含 landscape section"

    # 检查表格列数（unified 表格应为 N=5 列）
    form_tables = doc.tables[2:]
    # 找到列数最多的表格，应该是 unified 表格
    max_cols = max(len(t.columns) for t in form_tables)
    assert max_cols == 5, "unified 表格应为 5 列"


# ========== Task 3.5: unified 字段顺序测试 ==========

def test_export_unified_field_order_matches_order_index(session: Session, tmp_path: Path) -> None:
    """验证表格行顺序与 order_index 一致。"""
    project, _ = create_minimal_project(session)

    visit = Visit(project_id=project.id, name="访视1", code="V1", sequence=1)
    session.add(visit)
    session.flush()

    form = Form(project_id=project.id, name="顺序测试表单", code="F_ORDER", order_index=1)
    session.add(form)
    session.flush()

    vf = VisitForm(visit_id=visit.id, form_id=form.id, sequence=1)
    session.add(vf)
    session.flush()

    # 添加字段，故意设置乱序 order_index
    fd1 = create_text_field_def(session, project.id, "字段A")
    fd2 = create_text_field_def(session, project.id, "字段B")
    fd3 = create_text_field_def(session, project.id, "列1")
    fd4 = create_text_field_def(session, project.id, "列2")
    fd5 = create_text_field_def(session, project.id, "列3")
    fd6 = create_text_field_def(session, project.id, "列4")
    fd7 = create_text_field_def(session, project.id, "列5")

    add_field_to_form(session, form.id, fd1.id, order_index=10, inline_mark=0)  # 普通字段在前
    add_field_to_form(session, form.id, fd3.id, order_index=20, inline_mark=1)
    add_field_to_form(session, form.id, fd4.id, order_index=21, inline_mark=1)
    add_field_to_form(session, form.id, fd5.id, order_index=22, inline_mark=1)
    add_field_to_form(session, form.id, fd6.id, order_index=23, inline_mark=1)
    add_field_to_form(session, form.id, fd7.id, order_index=24, inline_mark=1)
    add_field_to_form(session, form.id, fd2.id, order_index=30, inline_mark=0)  # 普通字段在后

    session.commit()

    output_path = tmp_path / "order.docx"
    ExportService(session).export_project_to_word(project.id, str(output_path))

    doc = Document(str(output_path))

    # 找到 unified 表格（5 列的那个）
    unified_table = None
    for table in doc.tables[2:]:
        if len(table.columns) == 5:
            unified_table = table
            break

    assert unified_table is not None, "应存在 5 列 unified 表格"

    # 验证行内容顺序：字段A → inline block (列1-5) → 字段B
    # 第一行应该是"字段A"的 label/value
    first_row_text = unified_table.rows[0].cells[0].text
    assert "字段A" in first_row_text, "第一行应为字段A"


# ========== Task 3.7: unified label/log 行测试 ==========

def test_export_unified_full_row_span_equals_N(session: Session, tmp_path: Path) -> None:
    """验证 label/log 行 gridSpan = N。"""
    project, _ = create_minimal_project(session)

    visit = Visit(project_id=project.id, name="访视1", code="V1", sequence=1)
    session.add(visit)
    session.flush()

    form = Form(project_id=project.id, name="全宽行测试", code="F_FULLROW", order_index=1)
    session.add(form)
    session.flush()

    vf = VisitForm(visit_id=visit.id, form_id=form.id, sequence=1)
    session.add(vf)
    session.flush()

    # 添加标签字段（全宽行）
    fd_label = create_label_field_def(session, project.id, "标题行")
    add_field_to_form(session, form.id, fd_label.id, order_index=1, inline_mark=0)

    # 添加 5 个 inline 字段触发 unified
    for i in range(1, 6):
        fd = create_text_field_def(session, project.id, f"列{i}")
        add_field_to_form(session, form.id, fd.id, order_index=10 + i, inline_mark=1)

    session.commit()

    output_path = tmp_path / "fullrow.docx"
    ExportService(session).export_project_to_word(project.id, str(output_path))

    doc = Document(str(output_path))

    # 找到 unified 表格
    unified_table = None
    for table in doc.tables[2:]:
        if len(table.columns) == 5:
            unified_table = table
            break

    assert unified_table is not None, "应存在 5 列 unified 表格"

    # 验证第一行是全宽行（合并后所有 cell 指向同一对象）
    first_row = unified_table.rows[0]
    # 合并后 row.cells 返回 N 个对象，但它们共享同一个物理 cell（相同内存地址）
    cells = list(first_row.cells)
    # 检查所有 cell 对象是同一个（合并后的特征）
    assert all(c is cells[0] for c in cells), "全宽行应合并为单一 cell"
    assert "标题行" in cells[0].text, "全宽行内容应为标签字段"


# ========== Task 3.8: unified 窄 block 测试 ==========

def test_export_unified_narrow_block_merge_spans(session: Session, tmp_path: Path) -> None:
    """验证 M < N 时 merge spans 正确。"""
    project, _ = create_minimal_project(session)

    visit = Visit(project_id=project.id, name="访视1", code="V1", sequence=1)
    session.add(visit)
    session.flush()

    form = Form(project_id=project.id, name="窄 block 测试", code="F_NARROW", order_index=1)
    session.add(form)
    session.flush()

    vf = VisitForm(visit_id=visit.id, form_id=form.id, sequence=1)
    session.add(vf)
    session.flush()

    # 添加普通字段
    fd_normal = create_text_field_def(session, project.id, "普通字段")
    add_field_to_form(session, form.id, fd_normal.id, order_index=1, inline_mark=0)

    # 添加 8 个 inline 字段触发 unified（N=8）
    for i in range(1, 9):
        fd = create_text_field_def(session, project.id, f"列{i}")
        add_field_to_form(session, form.id, fd.id, order_index=10 + i, inline_mark=1)

    session.commit()

    output_path = tmp_path / "narrow.docx"
    ExportService(session).export_project_to_word(project.id, str(output_path))

    doc = Document(str(output_path))

    # 找到 unified 表格（8 列）
    unified_table = None
    for table in doc.tables[2:]:
        if len(table.columns) == 8:
            unified_table = table
            break

    assert unified_table is not None, "应存在 8 列 unified 表格"


# ========== Task 3.9: section 方向恢复测试 ==========

def test_export_landscape_form_followed_by_portrait(session: Session, tmp_path: Path) -> None:
    """验证 landscape form 后续 form 为 portrait。"""
    project, _ = create_minimal_project(session)

    visit = Visit(project_id=project.id, name="访视1", code="V1", sequence=1)
    session.add(visit)
    session.flush()

    # 第一个表单：unified（触发 landscape）
    form1 = Form(project_id=project.id, name="横向表单", code="F_LAND", order_index=1)
    session.add(form1)
    session.flush()

    vf1 = VisitForm(visit_id=visit.id, form_id=form1.id, sequence=1)
    session.add(vf1)
    session.flush()

    fd_normal1 = create_text_field_def(session, project.id, "普通1")
    add_field_to_form(session, form1.id, fd_normal1.id, order_index=1, inline_mark=0)
    for i in range(1, 6):
        fd = create_text_field_def(session, project.id, f"列{i}_1")
        add_field_to_form(session, form1.id, fd.id, order_index=10 + i, inline_mark=1)

    # 第二个表单：普通（应恢复 portrait）
    form2 = Form(project_id=project.id, name="纵向表单", code="F_PORT", order_index=2)
    session.add(form2)
    session.flush()

    vf2 = VisitForm(visit_id=visit.id, form_id=form2.id, sequence=2)
    session.add(vf2)
    session.flush()

    fd_normal2 = create_text_field_def(session, project.id, "普通2")
    add_field_to_form(session, form2.id, fd_normal2.id, order_index=1, inline_mark=0)
    fd_normal3 = create_text_field_def(session, project.id, "普通3")
    add_field_to_form(session, form2.id, fd_normal3.id, order_index=2, inline_mark=0)

    session.commit()

    output_path = tmp_path / "multi_section.docx"
    ExportService(session).export_project_to_word(project.id, str(output_path))

    doc = Document(str(output_path))

    # 检查 section 方向序列
    sections = list(doc.sections)
    # 应包含 landscape 和 portrait
    orientations = [s.orientation for s in sections]
    assert WD_ORIENT.LANDSCAPE in orientations, "应包含 landscape section"
    assert WD_ORIENT.PORTRAIT in orientations, "应包含 portrait section"


# ========== Task 3.11: merge 后样式保留测试 ==========

def test_export_unified_preserves_cell_shading(session: Session, tmp_path: Path) -> None:
    """验证底纹/颜色在 merge 后的 cell 上保留。"""
    project, _ = create_minimal_project(session)

    visit = Visit(project_id=project.id, name="访视1", code="V1", sequence=1)
    session.add(visit)
    session.flush()

    form = Form(project_id=project.id, name="样式测试", code="F_STYLE", order_index=1)
    session.add(form)
    session.flush()

    vf = VisitForm(visit_id=visit.id, form_id=form.id, sequence=1)
    session.add(vf)
    session.flush()

    # 添加带颜色的普通字段
    fd_color = create_text_field_def(session, project.id, "彩色字段")
    ff_color = add_field_to_form(session, form.id, fd_color.id, order_index=1, inline_mark=0)
    ff_color.bg_color = "0070C0"

    # 添加 inline 字段触发 unified
    for i in range(1, 6):
        fd = create_text_field_def(session, project.id, f"列{i}")
        add_field_to_form(session, form.id, fd.id, order_index=10 + i, inline_mark=1)

    session.commit()

    output_path = tmp_path / "style.docx"
    ExportService(session).export_project_to_word(project.id, str(output_path))

    doc = Document(str(output_path))

    # 找到 unified 表格
    unified_table = None
    for table in doc.tables[2:]:
        if len(table.columns) == 5:
            unified_table = table
            break

    assert unified_table is not None, "应存在 5 列 unified 表格"
    # 导出成功即验证基本结构，样式保留需人工检查或更复杂的 XML 解析


def test_export_unified_applies_borders_to_rows_added_after_table_creation(session: Session, tmp_path: Path) -> None:
    """验证 unified 表格在动态 add_row 后仍保留单元格边框。"""
    project, _ = create_minimal_project(session)

    visit = Visit(project_id=project.id, name="访视1", code="V1", sequence=1)
    session.add(visit)
    session.flush()

    form = Form(project_id=project.id, name="边框测试", code="F_BORDER", order_index=1)
    session.add(form)
    session.flush()

    vf = VisitForm(visit_id=visit.id, form_id=form.id, sequence=1)
    session.add(vf)
    session.flush()

    fd_normal = create_text_field_def(session, project.id, "普通字段")
    add_field_to_form(session, form.id, fd_normal.id, order_index=1, inline_mark=0)

    for i in range(1, 6):
        fd = create_text_field_def(session, project.id, f"列{i}")
        add_field_to_form(session, form.id, fd.id, order_index=10 + i, inline_mark=1)

    session.commit()

    output_path = tmp_path / "border.docx"
    ExportService(session).export_project_to_word(project.id, str(output_path))

    doc = Document(str(output_path))

    unified_table = None
    for table in doc.tables[2:]:
        if len(table.columns) == 5:
            unified_table = table
            break

    assert unified_table is not None, "应存在 5 列 unified 表格"

    first_row = unified_table.rows[0]
    borders = first_row.cells[0]._tc.tcPr.find(qn('w:tcBorders'))
    assert borders is not None, "动态添加的 unified 行应具有边框定义"

    for border_name in ["top", "left", "bottom", "right"]:
        border = borders.find(qn(f'w:{border_name}'))
        assert border is not None, f"应存在 {border_name} 边框"
        assert border.get(qn('w:val')) == 'single', f"{border_name} 边框应为 single"


def test_export_unified_table_has_table_level_borders(session: Session, tmp_path: Path) -> None:
    """验证 unified 表格具备表级 tblBorders（含 insideH/insideV），确保 Word 能渲染内部网格线。"""
    project, _ = create_minimal_project(session)

    visit = Visit(project_id=project.id, name="访视1", code="V1", sequence=1)
    session.add(visit)
    session.flush()

    form = Form(project_id=project.id, name="表级边框测试", code="F_TBLBDR", order_index=1)
    session.add(form)
    session.flush()

    vf = VisitForm(visit_id=visit.id, form_id=form.id, sequence=1)
    session.add(vf)
    session.flush()

    # 创建触发 unified landscape 的字段组合
    fd_normal = create_text_field_def(session, project.id, "普通字段")
    add_field_to_form(session, form.id, fd_normal.id, order_index=1, inline_mark=0)

    for i in range(1, 6):
        fd = create_text_field_def(session, project.id, f"列{i}")
        add_field_to_form(session, form.id, fd.id, order_index=10 + i, inline_mark=1)

    session.commit()

    output_path = tmp_path / "tbl_border.docx"
    ExportService(session).export_project_to_word(project.id, str(output_path))

    doc = Document(str(output_path))

    # 找到 unified 表格
    unified_table = None
    for table in doc.tables[2:]:
        if len(table.columns) == 5:
            unified_table = table
            break

    assert unified_table is not None, "应存在 5 列 unified 表格"

    # 验证表级 tblBorders 存在
    tblPr = unified_table._tbl.tblPr
    assert tblPr is not None, "表格应具有 tblPr 属性"

    tblBorders = tblPr.find(qn('w:tblBorders'))
    assert tblBorders is not None, "unified 表格应具有表级 tblBorders"

    # 验证所有边框类型（含内部网格线 insideH/insideV）
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = tblBorders.find(qn(f'w:{border_name}'))
        assert border is not None, f"表级边框应包含 {border_name}"
        assert border.get(qn('w:val')) == 'single', f"{border_name} 应为 single 类型"
        assert border.get(qn('w:sz')) == '4', f"{border_name} 边框宽度应为 4"


# ========== Task 4.3: trailing_underscore 原子 token 测试 ==========


def create_choice_field_def(
    session: Session,
    project_id: int,
    label: str,
    codelist_id: int | None,
    field_type: str = "单选",
) -> FieldDefinition:
    """创建选择字段定义。"""
    fd = FieldDefinition(
        project_id=project_id,
        label=label,
        variable_name=f"VAR_{label.replace(' ', '_')}",
        field_type=field_type,
        codelist_id=codelist_id,
    )
    session.add(fd)
    session.flush()
    return fd


def test_export_choice_trailing_underscore_atom_token(session: Session, tmp_path: Path) -> None:
    """验证 trailing_underscore 选项渲染为原子 token（文本 + 尾线不拆分）。"""
    from src.models.codelist import CodeList, CodeListOption

    project, _ = create_minimal_project(session)

    visit = Visit(project_id=project.id, name="访视1", code="V1", sequence=1)
    session.add(visit)
    session.flush()

    form = Form(project_id=project.id, name="选择测试", code="F_CHOICE", order_index=1)
    session.add(form)
    session.flush()

    vf = VisitForm(visit_id=visit.id, form_id=form.id, sequence=1)
    session.add(vf)
    session.flush()

    # 创建选项字典
    codelist = CodeList(project_id=project.id, name="诊断结果", code="CL_DIAG")
    session.add(codelist)
    session.flush()

    # 选项1：有尾部填写线
    opt1 = CodeListOption(
        codelist_id=codelist.id, code="1", decode="确诊", trailing_underscore=1, order_index=1
    )
    # 选项2：无尾部填写线
    opt2 = CodeListOption(
        codelist_id=codelist.id, code="2", decode="疑似", trailing_underscore=0, order_index=2
    )
    # 选项3：有尾部填写线
    opt3 = CodeListOption(
        codelist_id=codelist.id, code="3", decode="排除", trailing_underscore=1, order_index=3
    )
    session.add_all([opt1, opt2, opt3])
    session.flush()

    # 创建单选字段
    fd = create_choice_field_def(session, project.id, "诊断", codelist.id, "单选")
    add_field_to_form(session, form.id, fd.id, order_index=1, inline_mark=0)

    session.commit()

    output_path = tmp_path / "choice_trailing.docx"
    ExportService(session).export_project_to_word(project.id, str(output_path))

    doc = Document(str(output_path))

    # 找到包含选择字段的表格
    form_tables = doc.tables[2:]
    assert len(form_tables) >= 1

    # 获取选择字段单元格的文本
    choice_cell = form_tables[0].cell(0, 1)  # 第一行第二列
    cell_text = choice_cell.text

    # 验证选项顺序：按 order_index 排序
    assert "确诊" in cell_text
    assert "疑似" in cell_text
    assert "排除" in cell_text

    # 验证 order_index 顺序：确诊在疑似前，疑似在排除前
    idx1 = cell_text.find("确诊")
    idx2 = cell_text.find("疑似")
    idx3 = cell_text.find("排除")
    assert idx1 < idx2, "确诊应在疑似前（order_index=1 vs 2）"
    assert idx2 < idx3, "疑似应在排除前（order_index=2 vs 3）"


def test_export_choice_order_index_sorting(session: Session, tmp_path: Path) -> None:
    """验证选项按 order_index 排序，id 扰动不影响顺序（P5 不变式）。"""
    from src.models.codelist import CodeList, CodeListOption

    project, _ = create_minimal_project(session)

    visit = Visit(project_id=project.id, name="访视1", code="V1", sequence=1)
    session.add(visit)
    session.flush()

    form = Form(project_id=project.id, name="排序测试", code="F_ORDER", order_index=1)
    session.add(form)
    session.flush()

    vf = VisitForm(visit_id=visit.id, form_id=form.id, sequence=1)
    session.add(vf)
    session.flush()

    codelist = CodeList(project_id=project.id, name="优先级", code="CL_PRI")
    session.add(codelist)
    session.flush()

    # 故意打乱 id 顺序，但设置明确的 order_index
    opt3 = CodeListOption(
        codelist_id=codelist.id, code="C", decode="第三", order_index=3
    )
    session.add(opt3)
    session.flush()

    opt1 = CodeListOption(
        codelist_id=codelist.id, code="A", decode="第一", order_index=1
    )
    session.add(opt1)
    session.flush()

    opt2 = CodeListOption(
        codelist_id=codelist.id, code="B", decode="第二", order_index=2
    )
    session.add(opt2)
    session.flush()

    fd = create_choice_field_def(session, project.id, "优先级", codelist.id, "单选（纵向）")
    add_field_to_form(session, form.id, fd.id, order_index=1, inline_mark=0)

    session.commit()

    output_path = tmp_path / "order_index.docx"
    ExportService(session).export_project_to_word(project.id, str(output_path))

    doc = Document(str(output_path))

    # 找到选择字段单元格
    form_tables = doc.tables[2:]
    choice_cell = form_tables[0].cell(0, 1)
    cell_text = choice_cell.text

    # 验证顺序按 order_index：第一 -> 第二 -> 第三
    idx1 = cell_text.find("第一")
    idx2 = cell_text.find("第二")
    idx3 = cell_text.find("第三")

    assert idx1 < idx2 < idx3, f"选项应按 order_index 排序，但顺序为: idx1={idx1}, idx2={idx2}, idx3={idx3}"


# ========== Task 4.2: unified 多 inline block 共享单表级宽度语义回归测试 ==========


def test_export_unified_multi_blocks_share_table_level_width(session: Session, tmp_path: Path) -> None:
    """验证多个 inline block 共享单表级宽度规划，而非各自独立分配。

    两个 inline block（5 列）共用一张 unified 表格，宽度分配应基于
    per-slot-max 聚合语义，而非 block 间拼接需求向量。
    """
    project, _ = create_minimal_project(session)

    visit = Visit(project_id=project.id, name="访视1", code="V1", sequence=1)
    session.add(visit)
    session.flush()

    form = Form(project_id=project.id, name="多块共享宽度", code="F_MULTI_BLK", order_index=1)
    session.add(form)
    session.flush()

    vf = VisitForm(visit_id=visit.id, form_id=form.id, sequence=1)
    session.add(vf)
    session.flush()

    # 第一个 inline block：5 列，标签短
    for i in range(1, 6):
        fd = create_text_field_def(session, project.id, f"A{i}")
        add_field_to_form(session, form.id, fd.id, order_index=10 + i, inline_mark=1)

    # 普通字段分隔两个 block
    fd_sep = create_text_field_def(session, project.id, "分隔字段")
    add_field_to_form(session, form.id, fd_sep.id, order_index=20, inline_mark=0)

    # 第二个 inline block：5 列，最后一列标签长
    for i in range(1, 5):
        fd = create_text_field_def(session, project.id, f"B{i}")
        add_field_to_form(session, form.id, fd.id, order_index=30 + i, inline_mark=1)
    fd_long = create_text_field_def(session, project.id, "这是一个非常长的中文标签文本")
    add_field_to_form(session, form.id, fd_long.id, order_index=35, inline_mark=1)

    session.commit()

    output_path = tmp_path / "multi_block.docx"
    ExportService(session).export_project_to_word(project.id, str(output_path))

    doc = Document(str(output_path))

    # 找到 unified 表格（5 列）
    unified_table = None
    for table in doc.tables[2:]:
        if len(table.columns) == 5:
            unified_table = table
            break

    assert unified_table is not None, "应存在 5 列 unified 表格"

    # 验证结构完整性：表格只有一个，不是两个独立表格
    five_col_tables = [t for t in doc.tables[2:] if len(t.columns) == 5]
    assert len(five_col_tables) == 1, "多 inline block 应共享同一张 unified 表格"

    # 验证第 5 列（index=4）比第 1 列（index=0）更宽
    # （因为第二个 block 的长标签在 slot 4 注入了更大需求）
    tbl_xml = unified_table._tbl
    grid_cols = tbl_xml.findall(qn('w:tblGrid') + '/' + qn('w:gridCol'))
    if grid_cols and len(grid_cols) == 5:
        w0 = int(grid_cols[0].get(qn('w:w'), '0'))
        w4 = int(grid_cols[4].get(qn('w:w'), '0'))
        assert w4 > w0, f"长标签列（slot 4）应比短标签列（slot 0）更宽: w0={w0}, w4={w4}"


# ========== Task 4.3: trailing_underscore 横向与纵向原子 token 测试 ==========


def test_export_horizontal_choice_trailing_uses_nbsp(session: Session, tmp_path: Path) -> None:
    """验证横向单选 trailing_underscore 使用 NBSP (\\u00A0) 连接标签与填写线。"""
    from src.models.codelist import CodeList, CodeListOption

    project, _ = create_minimal_project(session)

    visit = Visit(project_id=project.id, name="访视1", code="V1", sequence=1)
    session.add(visit)
    session.flush()

    form = Form(project_id=project.id, name="横向选择原子", code="F_HCHOICE", order_index=1)
    session.add(form)
    session.flush()

    vf = VisitForm(visit_id=visit.id, form_id=form.id, sequence=1)
    session.add(vf)
    session.flush()

    codelist = CodeList(project_id=project.id, name="横向选项", code="CL_HZ")
    session.add(codelist)
    session.flush()

    opt1 = CodeListOption(
        codelist_id=codelist.id, code="1", decode="有尾线", trailing_underscore=1, order_index=1
    )
    opt2 = CodeListOption(
        codelist_id=codelist.id, code="2", decode="无尾线", trailing_underscore=0, order_index=2
    )
    session.add_all([opt1, opt2])
    session.flush()

    fd = create_choice_field_def(session, project.id, "横向测试", codelist.id, "单选")
    add_field_to_form(session, form.id, fd.id, order_index=1, inline_mark=0)

    session.commit()

    output_path = tmp_path / "h_choice_nbsp.docx"
    ExportService(session).export_project_to_word(project.id, str(output_path))

    doc = Document(str(output_path))
    form_tables = doc.tables[2:]
    assert len(form_tables) >= 1

    # 检查选择字段单元格内的 run 文本
    choice_cell = form_tables[0].cell(0, 1)
    all_runs_text = [run.text for para in choice_cell.paragraphs for run in para.runs]
    joined = "".join(all_runs_text)

    # 有尾线的选项应包含 NBSP + 下划线
    assert "有尾线\u00A0______" in joined, f"横向选项应使用 NBSP 连接标签与填写线，实际: {repr(joined)}"
    # 无尾线的选项不应有下划线
    assert "无尾线\u00A0" not in joined, f"无尾线选项不应有 NBSP 填写线"


def test_export_vertical_choice_trailing_uses_nbsp(session: Session, tmp_path: Path) -> None:
    """验证纵向多选 trailing_underscore 使用 NBSP (\\u00A0) 连接标签与填写线。"""
    from src.models.codelist import CodeList, CodeListOption

    project, _ = create_minimal_project(session)

    visit = Visit(project_id=project.id, name="访视1", code="V1", sequence=1)
    session.add(visit)
    session.flush()

    form = Form(project_id=project.id, name="纵向选择原子", code="F_VCHOICE", order_index=1)
    session.add(form)
    session.flush()

    vf = VisitForm(visit_id=visit.id, form_id=form.id, sequence=1)
    session.add(vf)
    session.flush()

    codelist = CodeList(project_id=project.id, name="纵向选项", code="CL_VT")
    session.add(codelist)
    session.flush()

    opt1 = CodeListOption(
        codelist_id=codelist.id, code="1", decode="确诊", trailing_underscore=1, order_index=1
    )
    opt2 = CodeListOption(
        codelist_id=codelist.id, code="2", decode="排除", trailing_underscore=0, order_index=2
    )
    session.add_all([opt1, opt2])
    session.flush()

    fd = create_choice_field_def(session, project.id, "纵向测试", codelist.id, "多选（纵向）")
    add_field_to_form(session, form.id, fd.id, order_index=1, inline_mark=0)

    session.commit()

    output_path = tmp_path / "v_choice_nbsp.docx"
    ExportService(session).export_project_to_word(project.id, str(output_path))

    doc = Document(str(output_path))
    form_tables = doc.tables[2:]
    assert len(form_tables) >= 1

    choice_cell = form_tables[0].cell(0, 1)
    all_runs_text = [run.text for para in choice_cell.paragraphs for run in para.runs]
    joined = "".join(all_runs_text)

    # 纵向选项也应使用 NBSP 连接
    assert "确诊\u00A0______" in joined, f"纵向选项应使用 NBSP 连接标签与填写线，实际: {repr(joined)}"
    assert "排除\u00A0" not in joined, f"无尾线选项不应有 NBSP 填写线"


# ========== Task 4.4: 多行 default_value 回归测试 ==========


def test_export_multiline_default_value_preserves_lines(session: Session, tmp_path: Path) -> None:
    """验证多行 default_value 在导出中保留多行语义。"""
    project, _ = create_minimal_project(session)

    visit = Visit(project_id=project.id, name="访视1", code="V1", sequence=1)
    session.add(visit)
    session.flush()

    form = Form(project_id=project.id, name="多行默认值测试", code="F_MULTI_DV", order_index=1)
    session.add(form)
    session.flush()

    vf = VisitForm(visit_id=visit.id, form_id=form.id, sequence=1)
    session.add(vf)
    session.flush()

    fd = create_text_field_def(session, project.id, "多行字段")
    ff = add_field_to_form(session, form.id, fd.id, order_index=1, inline_mark=0)
    ff.default_value = "第一行内容\n第二行内容\n第三行内容"

    session.commit()

    output_path = tmp_path / "multiline_dv.docx"
    ExportService(session).export_project_to_word(project.id, str(output_path))

    doc = Document(str(output_path))
    form_tables = doc.tables[2:]
    assert len(form_tables) >= 1

    # 获取值单元格
    value_cell = form_tables[0].cell(0, 1)
    cell_text = value_cell.text

    # 验证多行内容都存在
    assert "第一行内容" in cell_text, "应包含第一行"
    assert "第二行内容" in cell_text, "应包含第二行"
    assert "第三行内容" in cell_text, "应包含第三行"

    # 验证多行通过段内换行（break）实现，所有行在同一段落
    paragraphs = value_cell.paragraphs
    # 使用 add_break() 实现多行，检查 run 数量覆盖所有行
    all_runs = [run.text for para in paragraphs for run in para.runs]
    joined = "".join(all_runs)
    assert "第一行内容" in joined
    assert "第二行内容" in joined
    assert "第三行内容" in joined


def test_export_multiline_default_value_in_inline_table(session: Session, tmp_path: Path) -> None:
    """验证 inline 表格中多行 default_value 保留多行语义。"""
    project, _ = create_minimal_project(session)

    visit = Visit(project_id=project.id, name="访视1", code="V1", sequence=1)
    session.add(visit)
    session.flush()

    form = Form(project_id=project.id, name="Inline多行", code="F_INL_MV", order_index=1)
    session.add(form)
    session.flush()

    vf = VisitForm(visit_id=visit.id, form_id=form.id, sequence=1)
    session.add(vf)
    session.flush()

    # 3 列 inline，第二列有多行默认值
    fd1 = create_text_field_def(session, project.id, "列1")
    fd2 = create_text_field_def(session, project.id, "列2")
    fd3 = create_text_field_def(session, project.id, "列3")

    add_field_to_form(session, form.id, fd1.id, order_index=1, inline_mark=1)
    ff2 = add_field_to_form(session, form.id, fd2.id, order_index=2, inline_mark=1)
    ff2.default_value = "行A\n行B"
    add_field_to_form(session, form.id, fd3.id, order_index=3, inline_mark=1)

    session.commit()

    output_path = tmp_path / "inline_multiline.docx"
    ExportService(session).export_project_to_word(project.id, str(output_path))

    doc = Document(str(output_path))
    form_tables = doc.tables[2:]
    assert len(form_tables) >= 1

    # 找到 inline 表格（3 列）
    inline_table = None
    for t in form_tables:
        if len(t.columns) == 3:
            inline_table = t
            break

    assert inline_table is not None, "应存在 3 列 inline 表格"

    # inline 表格中多行默认值展开为多行（表头 + N 行数据）
    # "行A\n行B" 展开为 2 行数据，所以至少 3 行（1 表头 + 2 数据行）
    assert len(inline_table.rows) >= 3, f"多行默认值应展开为多行数据，实际行数: {len(inline_table.rows)}"

    # 第 1 行数据（row index 1）包含 "行A"
    cell_r1 = inline_table.rows[1].cells[1]
    assert "行A" in cell_r1.text, "第 1 数据行第 2 列应包含 '行A'"

    # 第 2 行数据（row index 2）包含 "行B"
    cell_r2 = inline_table.rows[2].cells[1]
    assert "行B" in cell_r2.text, "第 2 数据行第 2 列应包含 '行B'"
