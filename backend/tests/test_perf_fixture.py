from __future__ import annotations

from collections import Counter
from pathlib import Path

from docx import Document
from sqlalchemy import create_engine, func, select
from sqlalchemy.orm import Session

from scripts.generate_perf_fixture import (
    CHOICE_FIELD_TYPES,
    DOCX_FORM_TABLE_COUNT,
    DOCX_SKIPPED_TABLE_COUNT,
    FIELD_TYPE_COUNTS,
    FIXTURE_ID,
    FIXTURE_SCHEMA_VERSION,
    FIXTURE_SEED,
    MAIN_PROJECT_NAME,
    OWNER_PASSWORD,
    OWNER_USERNAME,
    generate_heavy_fixture,
)
from src.models.codelist import CodeList, CodeListOption
from src.models.field_definition import FieldDefinition
from src.models.form import Form
from src.models.form_field import FormField
from src.models.project import Project
from src.models.unit import Unit
from src.models.user import User
from src.models.visit import Visit
from src.models.visit_form import VisitForm
from src.services.auth_service import verify_password
from src.services.docx_import_service import DocxImportService

EXPECTED_HOST_COUNTS = {
    "users": 1,
    "projects": 1,
    "visits": 10,
    "forms": 40,
    "visit_forms": 40,
    "field_definitions": 1520,
    "form_fields": 1600,
    "log_rows": 80,
    "codelists": 20,
    "codelist_options": 400,
    "units": 30,
}

EXPECTED_MERGE_COUNTS = {
    "users": 1,
    "projects": 5,
    "visits": 50,
    "forms": 200,
    "visit_forms": 200,
    "field_definitions": 7600,
    "form_fields": 8000,
    "log_rows": 400,
    "codelists": 100,
    "codelist_options": 2000,
    "units": 150,
}


def test_heavy_fixture_generation_is_deterministic_for_fixed_seed(tmp_path: Path) -> None:
    with generate_heavy_fixture(seed=FIXTURE_SEED, root_dir=tmp_path) as first, generate_heavy_fixture(
        seed=FIXTURE_SEED,
        root_dir=tmp_path,
    ) as second:
        assert first.fixture_id == second.fixture_id == FIXTURE_ID
        assert first.schema_version == second.schema_version == FIXTURE_SCHEMA_VERSION
        assert first.db_path != second.db_path
        assert first.merge_db_path != second.merge_db_path
        assert first.docx_path != second.docx_path
        assert first.upload_docx_path != second.upload_docx_path

        first_snapshot = _snapshot_fixture(first.db_path, first.merge_db_path, first.docx_path, first.upload_docx_path)
        second_snapshot = _snapshot_fixture(second.db_path, second.merge_db_path, second.docx_path, second.upload_docx_path)

        assert first.counts == second.counts == EXPECTED_HOST_COUNTS
        assert first.merge_counts == second.merge_counts == EXPECTED_MERGE_COUNTS
        assert first.docx_counts == second.docx_counts == {
            "physical_tables": DOCX_SKIPPED_TABLE_COUNT + DOCX_FORM_TABLE_COUNT,
            "skipped_tables": DOCX_SKIPPED_TABLE_COUNT,
            "form_tables": DOCX_FORM_TABLE_COUNT,
        }
        assert first_snapshot == second_snapshot

        first_paths = [first.db_path, first.merge_db_path, first.docx_path, first.upload_docx_path]
        second_paths = [second.db_path, second.merge_db_path, second.docx_path, second.upload_docx_path]
        assert all(path.exists() for path in first_paths + second_paths)

    assert all(not path.exists() for path in first_paths + second_paths)


def test_heavy_fixture_rejects_non_phase1_seed(tmp_path: Path) -> None:
    try:
        generate_heavy_fixture(seed=1, root_dir=tmp_path)
    except ValueError as exc:
        assert str(FIXTURE_SEED) in str(exc)
    else:
        raise AssertionError("generate_heavy_fixture should reject non-fixed seeds")


def _snapshot_fixture(db_path: Path, merge_db_path: Path, docx_path: Path, upload_docx_path: Path) -> dict:
    return {
        "host": _snapshot_host_database(db_path),
        "merge": _snapshot_merge_database(merge_db_path),
        "docx": _snapshot_docx(docx_path, upload_docx_path),
    }


def _snapshot_host_database(db_path: Path) -> dict:
    engine = create_engine(f"sqlite+pysqlite:///{db_path.as_posix()}")
    try:
        with Session(engine) as session:
            counts = _collect_counts(session)
            assert counts == EXPECTED_HOST_COUNTS

            owner = session.scalar(select(User).order_by(User.id))
            assert owner is not None
            assert owner.username == OWNER_USERNAME
            assert verify_password(OWNER_PASSWORD, owner.hashed_password)

            project = session.scalar(select(Project).order_by(Project.id))
            assert project is not None
            assert project.name == MAIN_PROJECT_NAME

            visits = session.scalars(select(Visit).where(Visit.project_id == project.id).order_by(Visit.sequence)).all()
            forms = session.scalars(select(Form).where(Form.project_id == project.id).order_by(Form.order_index)).all()
            assert [visit.sequence for visit in visits] == list(range(1, 11))
            assert [form.order_index for form in forms] == list(range(1, 41))

            field_definitions = session.scalars(
                select(FieldDefinition).where(FieldDefinition.project_id == project.id).order_by(FieldDefinition.order_index)
            ).all()
            field_definition_by_id = {field_definition.id: field_definition for field_definition in field_definitions}
            field_definition_ids = set(field_definition_by_id)
            codelist_ids = set(session.scalars(select(CodeList.id).where(CodeList.project_id == project.id)).all())
            unit_ids = set(session.scalars(select(Unit.id).where(Unit.project_id == project.id)).all())

            type_sequences = []
            for form in forms:
                form_fields = session.scalars(
                    select(FormField).where(FormField.form_id == form.id).order_by(FormField.order_index)
                ).all()
                assert [form_field.order_index for form_field in form_fields] == list(range(1, 41))

                field_types = []
                for form_field in form_fields:
                    if form_field.is_log_row:
                        assert form_field.field_definition_id is None
                        assert form_field.label_override is not None
                        field_types.append("日志行")
                        continue

                    assert form_field.field_definition_id in field_definition_ids
                    field_definition = field_definition_by_id[form_field.field_definition_id]
                    assert field_definition.project_id == project.id
                    field_types.append(field_definition.field_type)

                    if field_definition.field_type == "数值":
                        assert field_definition.unit_id in unit_ids
                        assert field_definition.integer_digits is not None
                        assert field_definition.decimal_digits is not None
                    if field_definition.field_type in CHOICE_FIELD_TYPES:
                        assert field_definition.codelist_id in codelist_ids
                    if field_definition.field_type == "日期":
                        assert field_definition.date_format == "yyyy-MM-dd"
                    if field_definition.field_type == "日期时间":
                        assert field_definition.date_format == "yyyy-MM-dd HH:mm"
                    if field_definition.field_type == "时间":
                        assert field_definition.date_format == "HH:mm"

                assert Counter(field_types) == FIELD_TYPE_COUNTS
                type_sequences.append(tuple(field_types))

            visit_form_pairs = session.execute(
                select(Visit.sequence, VisitForm.sequence, Form.order_index)
                .join(VisitForm, VisitForm.visit_id == Visit.id)
                .join(Form, Form.id == VisitForm.form_id)
                .order_by(Visit.sequence, VisitForm.sequence)
            ).all()
            assert len(visit_form_pairs) == 40
            assert [form_order for _, _, form_order in visit_form_pairs] == list(range(1, 41))

            return {
                "counts": counts,
                "owner_username": owner.username,
                "project_name": project.name,
                "visit_names": [visit.name for visit in visits],
                "form_names": [form.name for form in forms],
                "field_type_sequences": type_sequences,
                "visit_form_pairs": [tuple(row) for row in visit_form_pairs],
            }
    finally:
        engine.dispose()


def _snapshot_merge_database(merge_db_path: Path) -> dict:
    engine = create_engine(f"sqlite+pysqlite:///{merge_db_path.as_posix()}")
    try:
        with Session(engine) as session:
            counts = _collect_counts(session)
            assert counts == EXPECTED_MERGE_COUNTS
            project_names = session.scalars(select(Project.name).order_by(Project.order_index, Project.id)).all()
            assert len(project_names) == 5
            assert project_names[:2] == [MAIN_PROJECT_NAME, f"{MAIN_PROJECT_NAME}_导入"]
            return {"counts": counts, "project_names": project_names}
    finally:
        engine.dispose()


def _snapshot_docx(docx_path: Path, upload_docx_path: Path) -> dict:
    document = Document(str(docx_path))
    parsed_forms = DocxImportService.parse_full(str(upload_docx_path))
    assert len(document.tables) == DOCX_SKIPPED_TABLE_COUNT + DOCX_FORM_TABLE_COUNT
    assert len(parsed_forms) == DOCX_FORM_TABLE_COUNT
    assert all(len(form["fields"]) == 40 for form in parsed_forms)
    assert all(form["name"].startswith("PERF_表单_") for form in parsed_forms)

    expected_docx_field_type_counts = {
        "文本": FIELD_TYPE_COUNTS["文本"],
        "数值": FIELD_TYPE_COUNTS["数值"],
        "日期": FIELD_TYPE_COUNTS["日期"],
        "日期时间": FIELD_TYPE_COUNTS["日期时间"],
        "时间": FIELD_TYPE_COUNTS["时间"],
        "单选": FIELD_TYPE_COUNTS["单选"],
        "多选": FIELD_TYPE_COUNTS["多选"],
        "单选（纵向）": FIELD_TYPE_COUNTS["单选（纵向）"],
        "多选（纵向）": FIELD_TYPE_COUNTS["多选（纵向）"],
        "标签": FIELD_TYPE_COUNTS["标签"],
        "log_row": FIELD_TYPE_COUNTS["日志行"],
    }
    parsed_field_type_counts = [Counter(field.get("field_type", field.get("type", "未知")) for field in form["fields"]) for form in parsed_forms]
    assert all(counter == expected_docx_field_type_counts for counter in parsed_field_type_counts)

    return {
        "physical_tables": len(document.tables),
        "parsed_form_names": [form["name"] for form in parsed_forms],
        "parsed_field_counts": [len(form["fields"]) for form in parsed_forms],
        "parsed_field_type_counts": parsed_field_type_counts,
    }


def _collect_counts(session: Session) -> dict[str, int]:
    return {
        "users": _count(session, User),
        "projects": _count(session, Project),
        "visits": _count(session, Visit),
        "forms": _count(session, Form),
        "visit_forms": _count(session, VisitForm),
        "field_definitions": _count(session, FieldDefinition),
        "form_fields": _count(session, FormField),
        "log_rows": session.scalar(select(func.count()).select_from(FormField).where(FormField.is_log_row == 1)) or 0,
        "codelists": _count(session, CodeList),
        "codelist_options": _count(session, CodeListOption),
        "units": _count(session, Unit),
    }


def _count(session: Session, model) -> int:
    return session.scalar(select(func.count()).select_from(model)) or 0
