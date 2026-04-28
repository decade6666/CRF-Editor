from __future__ import annotations

import argparse
import json
import shutil
import sys
import tempfile
from collections import Counter
from dataclasses import dataclass, field
from pathlib import Path
from random import Random
from typing import Any

from docx import Document
from sqlalchemy import create_engine, event, func, select
from sqlalchemy.engine import Engine
from sqlalchemy.orm import Session, sessionmaker

BACKEND_ROOT = Path(__file__).resolve().parents[1]
if str(BACKEND_ROOT) not in sys.path:
    sys.path.insert(0, str(BACKEND_ROOT))

from src.models import Base
from src.models.codelist import CodeList, CodeListOption
from src.models.field_definition import FieldDefinition
from src.models.form import Form
from src.models.form_field import FormField
from src.models.project import Project
from src.models.unit import Unit
from src.models.user import User
from src.models.visit import Visit
from src.models.visit_form import VisitForm
from src.services.auth_service import hash_password

FIXTURE_SEED = 20260425
FIXTURE_ID = f"heavy-1600-seed-{FIXTURE_SEED}"
FIXTURE_SCHEMA_VERSION = 1
OWNER_USERNAME = f"PERF_owner_{FIXTURE_SEED}"
OWNER_PASSWORD = f"PerfPass-{FIXTURE_SEED}"
MAIN_PROJECT_NAME = f"PERF_{FIXTURE_ID}_Main_Project"
DOCX_SKIPPED_TABLE_COUNT = 2
DOCX_FORM_TABLE_COUNT = 40
FIELD_TYPE_SEQUENCE = (
    ["文本"] * 8
    + ["数值"] * 5
    + ["日期"] * 4
    + ["日期时间"] * 3
    + ["时间"] * 2
    + ["单选"] * 4
    + ["多选"] * 4
    + ["单选（纵向）"] * 3
    + ["多选（纵向）"] * 3
    + ["标签"] * 2
    + ["日志行"] * 2
)
FIELD_TYPE_COUNTS = dict(Counter(FIELD_TYPE_SEQUENCE))
CHOICE_FIELD_TYPES = {"单选", "多选", "单选（纵向）", "多选（纵向）"}


@dataclass
class GeneratedPerfFixture:
    fixture_id: str
    seed: int
    schema_version: int
    temp_dir: Path
    db_path: Path
    merge_db_path: Path
    docx_path: Path
    upload_docx_path: Path
    owner_username: str
    owner_password: str
    main_project_name: str
    counts: dict[str, int]
    merge_counts: dict[str, int]
    docx_counts: dict[str, int]
    _temp_dir_handle: tempfile.TemporaryDirectory[str] | None = field(default=None, repr=False)

    def cleanup(self) -> None:
        if self._temp_dir_handle is None:
            return
        self._temp_dir_handle.cleanup()
        self._temp_dir_handle = None

    def to_summary(self) -> dict[str, Any]:
        return {
            "fixture_id": self.fixture_id,
            "fixture_schema_version": self.schema_version,
            "seed": self.seed,
            "owner_username": self.owner_username,
            "main_project_name": self.main_project_name,
            "counts": self.counts,
            "merge_counts": self.merge_counts,
            "docx_counts": self.docx_counts,
        }

    def __enter__(self) -> "GeneratedPerfFixture":
        return self

    def __exit__(self, exc_type, exc, traceback) -> None:
        self.cleanup()


def generate_heavy_fixture(*, seed: int = FIXTURE_SEED, root_dir: str | Path | None = None) -> GeneratedPerfFixture:
    if seed != FIXTURE_SEED:
        raise ValueError(f"Phase 1 fixture seed is fixed at {FIXTURE_SEED}")

    parent_dir = Path(root_dir) if root_dir is not None else None
    if parent_dir is not None:
        parent_dir.mkdir(parents=True, exist_ok=True)

    temp_dir_handle = tempfile.TemporaryDirectory(prefix=f"{FIXTURE_ID}-", dir=parent_dir)
    temp_dir = Path(temp_dir_handle.name)
    db_path = temp_dir / "host.sqlite3"
    merge_db_path = temp_dir / "merge.sqlite3"
    docx_path = temp_dir / f"{FIXTURE_ID}.docx"
    upload_dir = temp_dir / "uploads"
    upload_dir.mkdir(parents=True, exist_ok=True)
    upload_docx_path = upload_dir / f"{FIXTURE_ID}-upload.docx"

    try:
        counts = _build_host_database(db_path)
        merge_counts = _build_merge_database(merge_db_path)
        docx_counts = _build_docx_fixture(docx_path, upload_docx_path)
    except Exception:
        temp_dir_handle.cleanup()
        raise

    return GeneratedPerfFixture(
        fixture_id=FIXTURE_ID,
        seed=seed,
        schema_version=FIXTURE_SCHEMA_VERSION,
        temp_dir=temp_dir,
        db_path=db_path,
        merge_db_path=merge_db_path,
        docx_path=docx_path,
        upload_docx_path=upload_docx_path,
        owner_username=OWNER_USERNAME,
        owner_password=OWNER_PASSWORD,
        main_project_name=MAIN_PROJECT_NAME,
        counts=counts,
        merge_counts=merge_counts,
        docx_counts=docx_counts,
        _temp_dir_handle=temp_dir_handle,
    )


def _create_engine(db_path: Path) -> Engine:
    engine = create_engine(f"sqlite+pysqlite:///{db_path.as_posix()}")

    @event.listens_for(engine, "connect")
    def _configure_sqlite(dbapi_conn, _connection_record):
        dbapi_conn.execute("PRAGMA foreign_keys = ON")
        dbapi_conn.execute("PRAGMA journal_mode=WAL")
        dbapi_conn.execute("PRAGMA busy_timeout=5000")
        dbapi_conn.execute("PRAGMA synchronous=NORMAL")

    Base.metadata.create_all(engine)
    return engine


def _build_host_database(db_path: Path) -> dict[str, int]:
    engine = _create_engine(db_path)
    session_factory = sessionmaker(bind=engine, expire_on_commit=False)
    try:
        with session_factory() as session:
            owner = _create_owner(session, OWNER_USERNAME)
            _populate_project_graph(session, owner_id=owner.id, name=MAIN_PROJECT_NAME, project_order=1)
            session.commit()
            return _collect_counts(session)
    finally:
        engine.dispose()


def _build_merge_database(db_path: Path) -> dict[str, int]:
    engine = _create_engine(db_path)
    session_factory = sessionmaker(bind=engine, expire_on_commit=False)
    merge_names = [
        MAIN_PROJECT_NAME,
        f"{MAIN_PROJECT_NAME}_导入",
        f"PERF_{FIXTURE_ID}_Merge_Project_03",
        f"PERF_{FIXTURE_ID}_Merge_Project_04",
        f"PERF_{FIXTURE_ID}_Merge_Project_05",
    ]
    try:
        with session_factory() as session:
            owner = _create_owner(session, f"PERF_merge_owner_{FIXTURE_SEED}")
            for project_order, name in enumerate(merge_names, start=1):
                _populate_project_graph(session, owner_id=owner.id, name=name, project_order=project_order)
            session.commit()
            return _collect_counts(session)
    finally:
        engine.dispose()


def _create_owner(session: Session, username: str) -> User:
    owner = User(
        username=username,
        hashed_password=hash_password(OWNER_PASSWORD),
        is_admin=False,
    )
    session.add(owner)
    session.flush()
    return owner


def _populate_project_graph(session: Session, *, owner_id: int, name: str, project_order: int) -> None:
    project = Project(
        name=name,
        version="PERF_v1",
        order_index=project_order,
        owner_id=owner_id,
        trial_name=f"PERF_合成研究_{project_order:02d}",
        crf_version="PERF_CRF_v1",
        protocol_number=f"PERF_PROTO_{FIXTURE_SEED}_{project_order:02d}",
        screening_number_format="PERF-SCR-###",
        sponsor="PERF_Synthetic_Sponsor",
        data_management_unit="PERF_合成数据管理单位",
    )
    session.add(project)
    session.flush()

    units = _create_units(session, project.id)
    codelists = _create_codelists(session, project.id)
    visits = _create_visits(session, project.id)
    forms = _create_forms(session, project.id)
    _create_fields(session, project.id, forms, units, codelists)
    _create_visit_forms(session, visits, forms)
    session.flush()


def _create_units(session: Session, project_id: int) -> list[Unit]:
    units: list[Unit] = []
    for index in range(1, 31):
        unit = Unit(
            project_id=project_id,
            symbol=f"PERF单位{index:02d}",
            code=f"PERF_UNIT_{index:02d}",
            order_index=index,
        )
        session.add(unit)
        units.append(unit)
    session.flush()
    return units


def _create_codelists(session: Session, project_id: int) -> list[CodeList]:
    codelists: list[CodeList] = []
    for codelist_index in range(1, 21):
        codelist = CodeList(
            project_id=project_id,
            name=f"PERF_字典_{codelist_index:02d}",
            code=f"PERF_CL_{codelist_index:02d}",
            description=f"PERF_合成选项字典_{codelist_index:02d}",
            order_index=codelist_index,
        )
        session.add(codelist)
        session.flush()
        codelists.append(codelist)
        for option_index in range(1, 21):
            session.add(
                CodeListOption(
                    codelist_id=codelist.id,
                    code=f"PERF_OPT_{option_index:02d}",
                    decode=f"PERF_选项_{codelist_index:02d}_{option_index:02d}",
                    trailing_underscore=1 if option_index % 5 == 0 else 0,
                    order_index=option_index,
                )
            )
    session.flush()
    return codelists


def _create_visits(session: Session, project_id: int) -> list[Visit]:
    visits: list[Visit] = []
    for index in range(1, 11):
        visit = Visit(
            project_id=project_id,
            name=f"PERF_访视_{index:02d}",
            code=f"PERF_VISIT_{index:02d}",
            sequence=index,
        )
        session.add(visit)
        visits.append(visit)
    session.flush()
    return visits


def _create_forms(session: Session, project_id: int) -> list[Form]:
    forms: list[Form] = []
    for index in range(1, 41):
        form = Form(
            project_id=project_id,
            name=f"PERF_表单_{index:02d}",
            code=f"PERF_FORM_{index:02d}",
            domain=f"PF{index:03d}",
            order_index=index,
            design_notes=f"PERF_合成表单说明_{index:02d}",
        )
        session.add(form)
        forms.append(form)
    session.flush()
    return forms


def _create_fields(
    session: Session,
    project_id: int,
    forms: list[Form],
    units: list[Unit],
    codelists: list[CodeList],
) -> None:
    label_rng = Random(FIXTURE_SEED)
    field_definition_order = 1
    for form_index, form in enumerate(forms, start=1):
        pending_field_links: list[tuple[int, str, FieldDefinition]] = []
        for field_index, field_type in enumerate(FIELD_TYPE_SEQUENCE, start=1):
            if field_type == "日志行":
                session.add(
                    FormField(
                        form_id=form.id,
                        field_definition_id=None,
                        is_log_row=1,
                        order_index=field_index,
                        label_override=f"PERF_以下为log行_{form_index:02d}_{field_index:02d}",
                    )
                )
                continue

            field_definition = FieldDefinition(
                project_id=project_id,
                variable_name=f"PERF_F{form_index:02d}_{field_index:02d}",
                label=_label_for(label_rng, field_type, form_index, field_index),
                field_type=field_type,
                integer_digits=_integer_digits_for(field_type, field_index),
                decimal_digits=_decimal_digits_for(field_type, field_index),
                date_format=_date_format_for(field_type),
                codelist_id=_codelist_id_for(field_type, codelists, form_index, field_index),
                unit_id=_unit_id_for(field_type, units, form_index, field_index),
                order_index=field_definition_order,
            )
            session.add(field_definition)
            pending_field_links.append((field_index, field_type, field_definition))
            field_definition_order += 1

        session.flush()
        for field_index, field_type, field_definition in pending_field_links:
            session.add(
                FormField(
                    form_id=form.id,
                    field_definition_id=field_definition.id,
                    order_index=field_index,
                    required=1 if field_index % 7 == 0 else 0,
                    default_value=_default_value_for(field_type, form_index, field_index),
                    inline_mark=1 if field_index % 4 == 0 else 0,
                    bg_color=_color_for(form_index, field_index) if field_index % 11 == 0 else None,
                    text_color=_color_for(field_index, form_index) if field_index % 13 == 0 else None,
                )
            )
    session.flush()


def _create_visit_forms(session: Session, visits: list[Visit], forms: list[Form]) -> None:
    for form_index, form in enumerate(forms):
        visit = visits[form_index // 4]
        session.add(
            VisitForm(
                visit_id=visit.id,
                form_id=form.id,
                sequence=form_index % 4 + 1,
            )
        )
    session.flush()


def _label_for(rng: Random, field_type: str, form_index: int, field_index: int) -> str:
    suffix = f"{form_index:02d}_{field_index:02d}"
    if field_type == "文本":
        variants = [
            f"PERF_TEXT_ASCII_{suffix}",
            f"PERF_合成文本字段_{suffix}",
            f"PERF_合成文本字段_{suffix}_LONG_ASCII_CJK_PLACEHOLDER",
        ]
        return variants[rng.randrange(len(variants))]
    if field_type == "标签":
        return f"PERF_结构标签_{suffix}_合成说明"
    return f"PERF_{field_type}_{suffix}"


def _integer_digits_for(field_type: str, field_index: int) -> int | None:
    if field_type != "数值":
        return None
    return 3 + field_index % 5


def _decimal_digits_for(field_type: str, field_index: int) -> int | None:
    if field_type != "数值":
        return None
    return field_index % 3


def _date_format_for(field_type: str) -> str | None:
    if field_type == "日期":
        return "yyyy-MM-dd"
    if field_type == "日期时间":
        return "yyyy-MM-dd HH:mm"
    if field_type == "时间":
        return "HH:mm"
    return None


def _codelist_id_for(field_type: str, codelists: list[CodeList], form_index: int, field_index: int) -> int | None:
    if field_type not in CHOICE_FIELD_TYPES:
        return None
    return codelists[(form_index + field_index) % len(codelists)].id


def _unit_id_for(field_type: str, units: list[Unit], form_index: int, field_index: int) -> int | None:
    if field_type != "数值":
        return None
    return units[(form_index + field_index) % len(units)].id


def _default_value_for(field_type: str, form_index: int, field_index: int) -> str | None:
    token = f"PERF_VALUE_{form_index:02d}_{field_index:02d}"
    if field_type == "标签":
        return f"{token}_合成标签正文"
    if field_type == "文本" and field_index % 3 == 0:
        return f"{token}_ASCII_CJK_占位"
    return None


def _color_for(left: int, right: int) -> str:
    return f"{(left * 4099 + right * 131) % 0xFFFFFF:06X}"


def _collect_counts(session: Session) -> dict[str, int]:
    return {
        "users": _count(session, User),
        "projects": _count(session, Project),
        "visits": _count(session, Visit),
        "forms": _count(session, Form),
        "visit_forms": _count(session, VisitForm),
        "field_definitions": _count(session, FieldDefinition),
        "form_fields": _count(session, FormField),
        "log_rows": session.scalar(select(func.count()).select_from(FormField).where(FormField.is_log_row == 1)) or 0,
        "codelists": _count(session, CodeList),
        "codelist_options": _count(session, CodeListOption),
        "units": _count(session, Unit),
    }


def _count(session: Session, model: type[Any]) -> int:
    return session.scalar(select(func.count()).select_from(model)) or 0


def _build_docx_fixture(docx_path: Path, upload_docx_path: Path) -> dict[str, int]:
    document = Document()
    _add_preface_table(document, "PERF_封面信息表")
    _add_preface_table(document, "PERF_访视分布图")

    for form_index in range(1, DOCX_FORM_TABLE_COUNT + 1):
        document.add_paragraph(f"PERF_表单_{form_index:02d}", style="List Paragraph")
        table = document.add_table(rows=0, cols=2)
        for field_index, field_type in enumerate(FIELD_TYPE_SEQUENCE, start=1):
            row = table.add_row()
            row.cells[0].text = _docx_label_for(field_type, form_index, field_index)
            row.cells[1].text = _docx_value_for(field_type, form_index, field_index)

    document.save(docx_path)
    shutil.copyfile(docx_path, upload_docx_path)
    return {
        "physical_tables": len(document.tables),
        "skipped_tables": DOCX_SKIPPED_TABLE_COUNT,
        "form_tables": DOCX_FORM_TABLE_COUNT,
    }


def _add_preface_table(document: Document, title: str) -> None:
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = title
    table.cell(0, 1).text = f"PERF_{FIXTURE_SEED}_SYNTHETIC"


def _docx_label_for(field_type: str, form_index: int, field_index: int) -> str:
    if field_type == "日志行":
        return f"PERF_说明行_{form_index:02d}_{field_index:02d}"
    return f"PERF_{field_type}_字段_{form_index:02d}_{field_index:02d}"


def _docx_value_for(field_type: str, form_index: int, field_index: int) -> str:
    suffix = f"{form_index:02d}_{field_index:02d}"
    if field_type == "文本":
        return ""
    if field_type == "数值":
        return "|__|__|.|__|PERF单位"
    if field_type == "日期":
        return "|__|__|__|__|年|__|__|月|__|__|日"
    if field_type == "日期时间":
        return "|__|__|__|__|年|__|__|月|__|__|日 |__|__|:|__|__| PERF_DT_ASCII"
    if field_type == "时间":
        return "|__|__|:|__|__|"
    if field_type == "单选":
        return f"○PERF_选项A_{suffix}  ○PERF_选项B_{suffix}"
    if field_type == "多选":
        return f"□PERF_选项A_{suffix}  □PERF_选项B_{suffix}"
    if field_type == "单选（纵向）":
        return f"○PERF_纵向选项A_{suffix}\n○PERF_纵向选项B_{suffix}"
    if field_type == "多选（纵向）":
        return f"□PERF_纵向选项A_{suffix}\n□PERF_纵向选项B_{suffix}"
    if field_type == "标签":
        return f"PERF_合成标签正文_{suffix}_ASCII_CJK"
    if field_type == "日志行":
        return "以下为log行"
    return f"PERF_PLACEHOLDER_{suffix}"


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Generate deterministic heavy-1600 performance fixtures.")
    parser.add_argument("--seed", type=int, default=FIXTURE_SEED)
    parser.add_argument("--root-dir", type=Path, default=None)
    args = parser.parse_args(argv)

    with generate_heavy_fixture(seed=args.seed, root_dir=args.root_dir) as fixture:
        print(json.dumps(fixture.to_summary(), ensure_ascii=False, sort_keys=True))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
